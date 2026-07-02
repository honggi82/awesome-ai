from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from generate_awesome_ai_survey_paper_ko_rich import (
    AWESOME_GITHUB,
    AWESOME_SITE,
    DATA_DIR,
    PAPER_DIR,
    REFERENCE_ARXIV,
    REFERENCE_GITHUB,
    REFERENCE_TITLE,
    add_bullets,
    add_figure,
    add_heading,
    add_numbered,
    add_p,
    add_table,
    clean,
    configure_doc,
    fmt_int,
    make_figures,
    set_korean_font,
    split_tags,
    top_papers_text,
)


OUT_PATH = PAPER_DIR / "awesome_ai_survey_paper_en.docx"


def add_title_page(doc: Document, stats: dict[str, object]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Awesome AI: Roadmap & User Guide")
    run.bold = True
    set_korean_font(run, 23)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("A Citation-Ranked Taxonomy Map of AI Research, 2000-2026")
    run.bold = True
    set_korean_font(run, 16)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "An English long-form survey draft based on honggi82/awesome-ai and structured after arXiv:2605.18661"
    )
    run.italic = True
    set_korean_font(run, 11)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"Date: 2026-06-26 | Selected papers: {fmt_int(stats['selected'])} | Candidate records: {fmt_int(stats['candidates'])} | Selected-set citations: {fmt_int(stats['citations'])}"
    )
    set_korean_font(run, 10)


def add_reference_structure_analysis(doc: Document) -> None:
    add_heading(doc, "0. How the Reference Paper Turns a GitHub Repository into a Survey", 1)
    add_p(
        doc,
        f"The reference paper, {REFERENCE_TITLE}, does not merely repeat the contents of its GitHub repository. It turns a living curated list into a scholarly framework. The paper begins by linking the project page and GitHub repository, then reorganizes the repository into a four-phase research lifecycle. Each phase is decomposed into stages, each stage receives subtopics, assessment criteria, findings, observations, and a transition into the next stage. The appendices preserve inventory-style material, but the main paper interprets the list as a roadmap.",
    )
    add_p(
        doc,
        "This English version applies the same move to honggi82/awesome-ai. The source repository is not treated as a flat awesome list. It is treated as a living citation map with a data-generation pipeline, candidate corpus, selected dataset, taxonomy dataset, period analysis, link audit, GitHub-link provenance, interactive website, and review artifacts. The paper therefore explains both the map and the method used to make the map.",
    )
    add_table(
        doc,
        ["Element in arXiv:2605.18661", "Function in the reference paper", "Adaptation for awesome-ai"],
        [
            [
                "Project links and abstract",
                "Declare a living project page and summarize the field-level tension.",
                "Link GitHub, website, CSV datasets, and define the citation-influence challenge.",
            ],
            [
                "Preliminaries",
                "Define lifecycle, methodological families, scope, and timeline.",
                "Define citation-map lifecycle, metadata-curation families, scope, and AI-history periods.",
            ],
            [
                "Phase chapters",
                "Analyze each lifecycle stage with submethods, assessment, and observations.",
                "Analyze corpus construction, historical roadmap, taxonomy analysis, validation, and dissemination.",
            ],
            [
                "Assessment sections",
                "Separate capability from scientific reliability.",
                "Separate citation influence from evidence quality, reproducibility, and freshness.",
            ],
            [
                "Appendix inventories",
                "Keep long tool lists while preventing the main paper from becoming a list dump.",
                "Keep repository inventory, top papers, yearly leaders, and data-source references.",
            ],
        ],
    )
    add_heading(doc, "0.1 What It Means to Explain GitHub Content Scientifically", 2)
    add_p(
        doc,
        "A repository-centered survey must do more than paraphrase README sections. It should answer why the repository is organized as it is, what evidence supports the organization, which risks remain hidden, and how a reader should use the resource. In the reference paper, the GitHub repository tracks papers and tools, while the paper turns that tracking into a lifecycle argument. The same principle guides this document.",
    )
    add_p(
        doc,
        "For awesome-ai, the central object is not an agentic research lifecycle but an evidence-backed citation map. The repository describes a taxonomy-first, citation-ranked map of AI research. The paper below translates that into four phases: corpus construction, historical roadmap, taxonomy analysis, and validation/dissemination. This structure lets the reader understand the data, the ranking logic, the historical interpretation, and the practical uses without mistaking the map for an exhaustive expert judgment.",
    )


def add_abstract(doc: Document, stats: dict[str, object], cat: pd.DataFrame) -> None:
    top_count = cat.iloc[0]
    top_citations = cat.sort_values("citations", ascending=False).iloc[0]
    add_heading(doc, "Abstract", 1)
    add_p(
        doc,
        f"Artificial intelligence research has moved from statistical learning, kernel methods, and data mining toward deep learning, benchmark-centered vision, transformers, foundation models, multimodal systems, generative AI, AI for science, and responsible AI. Yet researchers still lack a compact, auditable map that connects long-term influence with modern research frontiers. This paper analyzes the honggi82/awesome-ai repository as a citation-ranked map of AI research from {stats['years']}. The repository audits up to 1,000 AI-related candidate papers per year, selects the top 100 papers per year by citation count, and produces {fmt_int(stats['selected'])} selected papers from {fmt_int(stats['candidates'])} candidate records, with {fmt_int(stats['citations'])} citations in the selected set.",
    )
    add_p(
        doc,
        f"Following the roadmap-and-user-guide structure of {REFERENCE_TITLE}, this document reorganizes awesome-ai into four phases. Phase 1 describes corpus construction: query design, AI relevance filtering, deduplication, year-balanced ranking, deterministic enrichment, and output generation. Phase 2 presents a historical roadmap over four periods: 2000-2008, 2009-2016, 2017-2021, and 2022-2026. Phase 3 analyzes the ten-category taxonomy, including representative papers, category-level limitations, keyword tags, method tags, and category interactions. Phase 4 examines validation and dissemination through DOI, arXiv, PubMed, open-PDF, GitHub-link, README, GitHub Pages, and link-audit evidence.",
    )
    add_p(
        doc,
        f"The largest category by paper count is {top_count.category} with {fmt_int(top_count.papers)} papers, while the largest citation mass belongs to {top_citations.category} with {fmt_int(top_citations.citations)} citations. This difference shows that AI influence is not a single linear story. It is built through datasets, benchmarks, software infrastructure, architectures, pretraining recipes, model releases, safety frameworks, and domain validation. The paper argues that awesome-ai is best used not as a final authority, but as an auditable starting point for literature review, graduate seminars, benchmark search, research planning, and source-driven expert reading.",
    )
    add_heading(doc, "Project and Data Links", 1)
    add_bullets(
        doc,
        [
            f"Awesome AI GitHub repository: {AWESOME_GITHUB}",
            f"Interactive website: {AWESOME_SITE}",
            "Selected dataset: awesome-ai/data/papers_2000_2026.csv",
            "Taxonomy dataset: awesome-ai/data/papers_taxonomy_2000_2026.csv",
            "Candidate pool: awesome-ai/data/candidates_top1000_2000_2026.csv",
            f"Reference paper: {REFERENCE_ARXIV}",
            f"Reference GitHub repository: {REFERENCE_GITHUB}",
        ],
    )


def add_introduction_and_preliminaries(
    doc: Document,
    stats: dict[str, object],
    figures: dict[str, Path],
) -> None:
    add_heading(doc, "1. Introduction", 1)
    intro = [
        "The AI literature is now too large to read as a single stream. A researcher entering the field encounters foundational textbooks, software libraries, classical machine-learning methods, ImageNet-era vision architectures, reinforcement-learning breakthroughs, transformer-based language models, multimodal pretraining systems, generative models, AI-for-science systems, and responsible-AI frameworks. The problem is not scarcity of information. It is the absence of a compact, source-backed map that helps readers decide what to read first and how to interpret influence.",
        "Awesome-ai addresses this problem with citation-ranked metadata curation. It does not claim to replace expert reading. Instead, it constructs an auditable map: broad AI queries collect candidates; explicit relevance rules remove obvious noise; identifiers and normalized titles deduplicate records; each year contributes the top 100 papers by citation count; deterministic rules add categories, tags, key ideas, strengths, limitations, and source links.",
        "This design has an important advantage. A simple all-time citation ranking would overrepresent older papers and a few highly cited benchmark periods. A purely recent list would lose the historical infrastructure that made modern AI possible. The year-balanced design gives each year a voice while retaining citation influence as the primary ranking signal. It is not perfect, but it is a useful compromise between historical depth and contemporary coverage.",
        "The reference paper on AI auto-research frames its field through a capability-integrity challenge: AI systems can produce research-like artifacts faster than they can verify them. Awesome-ai faces a parallel citation-influence challenge: citation count identifies influence, but not necessarily correctness, reproducibility, fairness, safety, or current validity. This paper repeatedly returns to that distinction.",
    ]
    for paragraph in intro:
        add_p(doc, paragraph)
    add_heading(doc, "1.1 Contributions", 2)
    add_bullets(
        doc,
        [
            f"We turn the {AWESOME_GITHUB} repository into a long-form survey-style roadmap and user guide.",
            f"We analyze {fmt_int(stats['selected'])} selected papers from {fmt_int(stats['candidates'])} candidate records across {stats['years']}.",
            "We map the repository into four phases: corpus construction, historical roadmap, taxonomy analysis, and validation/dissemination.",
            "We provide category-level interpretations, period-level transitions, metadata coverage analysis, and practical reading protocols.",
            "We explicitly distinguish citation influence from scientific closure, reproducibility, and deployment readiness.",
        ],
    )
    add_figure(doc, figures["framework"], "Figure 1. Four-phase framework for converting awesome-ai into a roadmap and user guide.")

    add_heading(doc, "2. Preliminaries", 1)
    add_heading(doc, "2.1 Citation-Map Lifecycle", 2)
    add_p(
        doc,
        "The citation-map lifecycle has four phases. P1 Corpus Construction builds the candidate and selected datasets. P2 Historical Roadmap places selected papers on a timeline and interprets period transitions. P3 Taxonomy Analysis organizes papers into ten research categories and explains their representative contributions and limitations. P4 Validation and Dissemination audits links, code provenance, public artifacts, and user-facing outputs.",
    )
    add_table(
        doc,
        ["Phase", "Role in awesome-ai", "Representative artifacts", "Primary risk"],
        [
            [
                "P1 Corpus Construction",
                "Collect candidates, filter AI relevance, deduplicate records, and select top papers per year.",
                "scripts/build_awesome_ai.py; candidates_top1000_2000_2026.csv",
                "Query bias, metadata noise, and citation bias.",
            ],
            [
                "P2 Historical Roadmap",
                "Interpret selected papers across four historical periods.",
                "papers_2000_2026.csv; period_analysis_2000_2026.json",
                "Citation lag and overemphasis on benchmark-heavy periods.",
            ],
            [
                "P3 Taxonomy Analysis",
                "Group papers by AI research taxonomy and explain category-level evidence.",
                "papers_taxonomy_2000_2026.csv; README.md",
                "Oversimplification caused by assigning one main category.",
            ],
            [
                "P4 Validation and Dissemination",
                "Expose links, audits, GitHub provenance, website, and citation metadata.",
                "link_audit_2000_2026.json; github_links_2000_2026.json; docs/index.html",
                "Confusing link existence with full reproducibility.",
            ],
        ],
    )
    add_heading(doc, "2.2 Metadata-Curation Families", 2)
    add_p(
        doc,
        "The reference paper compares prompt engineering, retrieval-augmented generation, agentic systems, training-based systems, and hybrid approaches. For awesome-ai, the analogous objects are metadata-curation families. These families describe how the map is produced rather than how models are trained.",
    )
    add_table(
        doc,
        ["Curation family", "Purpose", "Implementation cue", "Evaluation question"],
        [
            ["Broad keyword search", "Cover AI broadly across eras.", "18 Semantic Scholar queries.", "Does the query set miss important subfields?"],
            ["Relevance filtering", "Remove obvious non-AI noise.", "AI_RELEVANCE_PATTERNS.", "Does filtering balance recall and precision?"],
            ["Identifier deduplication", "Merge duplicate records.", "DOI, arXiv, PubMed, CorpusId, paperId, normalized title.", "Does deduplication avoid false merges?"],
            ["Year-balanced ranking", "Select 100 papers per year.", "citationCount descending, with audit fields retained.", "Does the design balance old influence and recent coverage?"],
            ["Taxonomy enrichment", "Add category, keywords, method tags, ideas, strengths, and limitations.", "CATEGORIES and deterministic rules.", "Does the taxonomy aid reading without overclaiming?"],
            ["Provenance auditing", "Track paper links, PDFs, and official GitHub links.", "link_audit and github_links JSON.", "Which links are official, reachable, and useful?"],
            ["Dissemination", "Publish README, website, datasets, and drafts.", "README.md, docs/index.html, paper outputs.", "Can users inspect and reuse the map?"],
        ],
    )


def add_corpus_construction(doc: Document, df: pd.DataFrame, stats: dict[str, object], figures: dict[str, Path]) -> None:
    add_heading(doc, "3. Phase 1: Corpus Construction", 1)
    add_p(
        doc,
        "Corpus construction is the methodological foundation of awesome-ai. The final list of 2,700 papers is only the visible artifact. Behind it is a pipeline for search, filtering, deduplication, ranking, enrichment, and publication. This section makes that pipeline explicit so that readers understand what the map can and cannot support.",
    )
    add_heading(doc, "3.1 Search Scope and Candidate Collection", 2)
    add_p(
        doc,
        "The repository searches broadly across artificial intelligence, machine learning, deep learning, neural networks, support vector machines, data mining, pattern recognition, Bayesian networks, foundation models, large language models, natural language processing, computer vision, reinforcement learning, generative AI, diffusion models, graph neural networks, multimodal learning, AI safety, fairness, explainability, robustness, AI for science, healthcare, and robotics. This breadth is essential because AI history cannot be reduced to LLMs or any single modern subfield.",
    )
    add_p(
        doc,
        "The broad query design improves recall but introduces noise. Terms such as classification, clustering, graph, clinical, science, and robot can identify AI papers, but they can also retrieve papers that use adjacent vocabulary without making an AI contribution. The next stages therefore filter relevance and remove bad-title patterns such as editorials, proceedings front matter, and overly generic book titles.",
    )
    add_heading(doc, "3.2 Relevance Filtering and Deduplication", 2)
    add_p(
        doc,
        "AI relevance is determined from title, abstract, and venue metadata using explicit patterns. These patterns include machine learning, deep learning, transformer, language model, LLM, foundation model, computer vision, object detection, semantic segmentation, NLP, reinforcement learning, diffusion model, GAN, graph neural, self-supervised learning, explainable AI, robustness, fairness, privacy, prompt, and retrieval-augmented generation. This is transparent and reproducible, but it remains metadata-based rather than full-text expert reading.",
    )
    add_p(
        doc,
        "Deduplication uses strong identifiers first and weaker normalized-title matching later. DOI, arXiv, PubMed, CorpusId, and Semantic Scholar paperId reduce duplicate records across versions and metadata sources. Normalized title matching catches records without strong identifiers, but also carries a risk of false merges. The pipeline therefore reflects a practical large-scale metadata compromise rather than a perfect bibliographic authority file.",
    )
    add_heading(doc, "3.3 Ranking and Year Balancing", 2)
    add_p(
        doc,
        "The final selected set keeps the top 100 papers per publication year by citation count. This is the key design choice. It prevents older highly cited papers from fully dominating the map while preserving citation influence within each year. The result is a corpus that can show both long-term AI infrastructure and recent frontier movement.",
    )
    add_p(
        doc,
        "Year balancing is analogous to the reference paper's lifecycle balancing. The auto-research survey prevents a mature stage such as writing from swallowing less mature stages such as rebuttal or dissemination. Awesome-ai prevents historically dominant years from swallowing recent years. In both cases, the structure is an attempt to make heterogeneous evidence comparable without claiming that all evidence is equally mature.",
    )
    add_heading(doc, "3.4 Enrichment and Generated Outputs", 2)
    add_p(
        doc,
        "Each selected paper receives a category, method tags, keyword tags, key idea, strengths, limitations, source queries, importance reasons, and source links when available. The enrichment is deterministic and metadata-driven; no paid LLM, paid translation, or paid compute is used. This makes the outputs auditable and cheap to regenerate, but it also means that the generated key ideas should be treated as navigation aids rather than authoritative paper summaries.",
    )
    add_heading(doc, "3.5 Assessment: Corpus Quality", 2)
    add_p(
        doc,
        "Corpus quality can be assessed along four axes: coverage, precision, auditability, and interpretability. Coverage asks whether major AI streams are included. Precision asks whether non-AI or weakly related records are excluded. Auditability asks whether a user can trace a selected paper back to metadata fields and links. Interpretability asks whether the resulting categories and tags help a reader form useful questions. Awesome-ai is strongest on auditability and broad coverage; expert-level precision still requires human review.",
    )
    add_heading(doc, "3.6 Findings and Observations", 2)
    add_bullets(
        doc,
        [
            "Year-balanced citation ranking is a useful compromise between historical influence and recent coverage.",
            "Broad AI search is necessary for a cross-era map but inevitably creates metadata noise.",
            "Deterministic enrichment improves transparency but cannot replace full-paper expert review.",
            "Candidate files, selected files, taxonomy files, period analysis, GitHub links, and link audits form an evidence layer around the final README.",
        ],
    )
    add_figure(doc, figures["coverage"], "Figure 2. Evidence and provenance coverage in the selected dataset.")


def add_historical_roadmap(
    doc: Document,
    df: pd.DataFrame,
    periods: list[tuple[str, int, int, str]],
    figures: dict[str, Path],
) -> None:
    add_heading(doc, "4. Phase 2: Historical Roadmap", 1)
    add_p(
        doc,
        "The historical roadmap transforms the selected dataset into a story of field-level transitions. The question is not only which paper has the highest citation count. The stronger question is how research infrastructure, datasets, architectures, benchmarks, software libraries, and governance concerns changed over time.",
    )
    rows = []
    for label, start, end, description in periods:
        group = df[(df.year >= start) & (df.year <= end)]
        top_year = int(group.groupby("year").citationCount.sum().sort_values(ascending=False).index[0])
        lead_category = group.category.value_counts().index[0]
        top = group.sort_values("citationCount", ascending=False).iloc[0]
        rows.append(
            [
                label,
                fmt_int(len(group)),
                fmt_int(group.citationCount.sum()),
                top_year,
                lead_category,
                f"{top.title} ({int(top.year)}, {fmt_int(top.citationCount)} citations)",
                description,
            ]
        )
    add_table(doc, ["Period", "Papers", "Citations", "Peak year", "Largest category", "Leading paper", "Interpretation"], rows)
    add_figure(doc, figures["yearly"], "Figure 3. Citation mass by publication year.")
    add_figure(doc, figures["heatmap"], "Figure 4. Taxonomy coverage across four historical periods.")

    period_sections = {
        "2000-2008": [
            "This period is best read as the formation of the pre-deep-learning research language. Data mining, statistical learning, kernel methods, random forests, independent component analysis, clustering, early computer vision, recommender systems, and bioinformatics all appear as central reference points.",
            "The point is not that this period was primitive. It produced concepts and tools that still shape AI: generalization, feature representation, supervised and unsupervised learning, evaluation protocols, reusable software, and textbook-level consolidation. Many highly cited works in this period are books, surveys, or libraries because they became common infrastructure.",
            "A reader should therefore treat the period as an infrastructure layer. The citation map identifies which concepts became common language. It does not necessarily identify the most modern method for a current task.",
        ],
        "2009-2016": [
            "This period captures the deep-learning transition. ImageNet, AlexNet, VGG, Inception, ResNet, Faster R-CNN, YOLO, Batch Normalization, DQN, XGBoost, TensorFlow, and PyTorch-era infrastructure reshaped how AI research was performed.",
            "Vision dominates the citation mass because large-scale datasets, GPU training, and benchmark competitions made progress measurable and reusable. ResNet is not only a model paper; it is a training recipe that made very deep networks practical. ImageNet is not only a dataset; it is a governance mechanism for comparison.",
            "This period also shows that software infrastructure is research infrastructure. Scikit-learn, XGBoost, TensorFlow, and PyTorch changed the cost of experimentation and therefore changed what kinds of research were feasible.",
        ],
        "2017-2021": [
            "The 2017 transformer turn is the hinge. Attention Is All You Need begins as an NLP architecture paper but becomes a general-purpose design pattern. BERT, T5, GPT-style modeling, ViT, Swin Transformer, CLIP, and AlphaFold show how pretraining, transfer, scaling, and multimodal alignment spread across subfields.",
            "This period weakens old boundaries. NLP architectures move into vision. Vision-language models become a central interface. AI-for-science systems demonstrate that learned models can transform scientific workflows rather than merely automate a small prediction step.",
            "The assessment burden also increases. Benchmark scores become harder to interpret because model scale, data scale, pretraining corpora, prompt design, fine-tuning protocols, and leakage all affect results.",
        ],
        "2022-2026": [
            "This period is dominated by foundation models, instruction tuning, open model releases, system cards, reasoning prompts, multimodal LLMs, interpretability, and safety. InstructGPT, Chain-of-Thought prompting, LLaMA, Llama 2, Llama 3, Qwen, DeepSeek-R1, GPT-4o system-card style reporting, and mechanistic interpretability become visible reference points.",
            "Citation lag is strongest here. A 2025 or 2026 paper may be important but still have few citations. For this period, the map should be read as a frontier signal rather than a settled canon. Adoption, official code, model reuse, benchmark uptake, and expert judgment matter more than raw citation count alone.",
            "The period also shows that governance is no longer outside the technical story. Data provenance, model release policy, alignment, safety evaluation, interpretability, copyright, privacy, and misuse are now part of how foundation-model papers are interpreted.",
        ],
    }
    for index, (label, _, _, _) in enumerate(periods, 1):
        add_heading(doc, f"4.{index} {label}", 2)
        group = df[(df.year >= periods[index - 1][1]) & (df.year <= periods[index - 1][2])]
        for paragraph in period_sections[label]:
            add_p(doc, paragraph)
        add_p(doc, "Representative high-citation papers: " + top_papers_text(group, 6))
    add_heading(doc, "4.5 Assessment and Observations", 2)
    add_bullets(
        doc,
        [
            "The 2000-2008 period explains the statistical, software, and educational substrate of AI.",
            "The 2009-2016 period explains why dataset-centered deep learning became the dominant experimental style.",
            "The 2017-2021 period explains the transition from task-specific deep learning to transferable architectures and multimodal systems.",
            "The 2022-2026 period must be read with citation-lag awareness and governance awareness.",
        ],
    )


def add_taxonomy_analysis(doc: Document, df: pd.DataFrame, tax: pd.DataFrame, cat: pd.DataFrame, figures: dict[str, Path]) -> None:
    add_heading(doc, "5. Phase 3: Taxonomy Analysis", 1)
    add_p(
        doc,
        "The taxonomy is the central reading interface of awesome-ai. It is not a claim that AI research is cleanly separable into ten boxes. It is a pragmatic map. Papers such as Attention, CLIP, AlphaFold, PyTorch, and adversarial-examples work influence multiple categories, but a primary category helps readers navigate the dataset.",
    )
    rows = []
    for _, row in cat.iterrows():
        top = df[df.category == row.category].sort_values("citationCount", ascending=False).iloc[0]
        rows.append(
            [
                row.category,
                fmt_int(row.papers),
                fmt_int(row.citations),
                fmt_int(row.influential),
                int(round(row.median_year)),
                f"{top.title} ({int(top.year)}, {fmt_int(top.citationCount)})",
            ]
        )
    add_table(doc, ["Category", "Papers", "Citations", "Influential citations", "Median year", "Leading paper"], rows)
    add_figure(doc, figures["category"], "Figure 5. Selected papers by taxonomy category.")
    add_figure(doc, figures["keywords"], "Figure 6. Keyword-tag distribution.")
    add_figure(doc, figures["methods"], "Figure 7. Method-tag distribution.")

    category_guide = {
        "General AI Methods and Systems": "This category captures the shared substrate of AI: textbooks, surveys, toolkits, regularization methods, scalable learning systems, and classical algorithms. Its large size reflects the long citation life of infrastructure and educational standards.",
        "Foundation Models and Large Language Models": "This category captures the rapid shift toward general-purpose pretrained models, instruction tuning, alignment, retrieval augmentation, model releases, and system-level reporting. It has a very recent median year, so citation lag must be considered.",
        "Vision and Multimodal Learning": "This category has the largest citation mass. ImageNet, AlexNet, VGG, ResNet, Faster R-CNN, YOLO, ViT, CLIP, and Swin Transformer show how datasets, architectures, and representation learning mutually reinforced one another.",
        "AI for Science, Healthcare, and Robotics": "This category connects AI to external validation. AlphaFold, KITTI, medical-image analysis, biomedical modeling, robotics, and autonomous driving require domain constraints beyond benchmark scores.",
        "Graph Learning, Recommendation, and Core Methods": "This category covers non-Euclidean data, recommender systems, optimization, Bayesian methods, architecture search, and explanation methods. It is especially sensitive to split design, leakage, and comparison budgets.",
        "Reinforcement Learning and Agents": "This category moves from control and games toward human feedback, planning, tool use, and agentic systems. Reward misspecification, sample inefficiency, simulator bias, and oversight remain central risks.",
        "Natural Language Processing and Knowledge": "This category covers language modeling, translation, retrieval, question answering, summarization, dialogue, and knowledge-intensive reasoning. It is deeply connected to foundation models but retains distinct evaluation concerns.",
        "Trustworthy, Explainable, and Responsible AI": "This category provides cross-cutting evaluation language for robustness, fairness, privacy, interpretability, uncertainty, safety, and governance. Its influence is broader than its paper count.",
        "Generative Models and Synthetic Media": "This category moves from GAN-centered synthesis to diffusion, score-based modeling, controllable generation, and synthetic media. Evaluation must consider not only perceptual quality but also provenance, controllability, misuse, and factuality.",
        "Representation, Self-Supervised, and Transfer Learning": "This smaller category acts as connective tissue. Word embeddings, transfer-learning surveys, adversarial examples, contrastive learning, and self-supervision affect nearly every modern subfield.",
    }
    for index, category in enumerate(cat["category"], 1):
        group = tax[tax.category == category]
        overview = clean(group["categoryOverview"].dropna().iloc[0]) if group["categoryOverview"].notna().any() else ""
        limitations = clean(group["categoryLimitations"].dropna().iloc[0]) if group["categoryLimitations"].notna().any() else ""
        top5 = group.sort_values("citationCount", ascending=False).head(5)
        add_heading(doc, f"5.{index} {category}", 2)
        add_p(doc, category_guide.get(category, overview))
        if overview:
            add_p(doc, "Repository overview: " + overview)
        if limitations:
            add_p(doc, "Assessment limitations: " + limitations)
        add_p(
            doc,
            "Representative papers: "
            + "; ".join(f"{int(row.year)} {row.title} ({fmt_int(row.citationCount)} citations)" for _, row in top5.iterrows()),
        )
        add_p(
            doc,
            "Reading note: this category should be used as a primary navigation coordinate, not as a hard intellectual boundary. Many high-impact papers connect several categories at once.",
        )
    add_heading(doc, "5.11 Summary and Transition", 2)
    add_p(
        doc,
        "Across the ten categories, AI history appears as layered infrastructure rather than a single line of model progress. General methods and software enable experimentation. Vision benchmarks reshape representation learning. NLP architectures become foundation-model architectures. AI for science and trustworthy AI add external validation and governance constraints. The taxonomy is therefore most useful when read together with time, provenance, and assessment criteria.",
    )


def add_validation_crosscutting_and_guide(
    doc: Document,
    df: pd.DataFrame,
    stats: dict[str, object],
    link_audit: dict[str, object],
    github_links: dict[str, object],
    cat: pd.DataFrame,
) -> None:
    add_heading(doc, "6. Phase 4: Validation and Dissemination", 1)
    add_p(
        doc,
        "For awesome-ai, validation means metadata and provenance validation, not peer review of each selected paper. The question is whether paper links, PDF links, GitHub links, source datasets, generated files, and public interfaces can be inspected. Dissemination means turning the dataset into README tables, an interactive GitHub Pages site, review drafts, citation metadata, and reusable CSV/JSON files.",
    )
    coverage_rows = []
    for col, label in [
        ("doi", "DOI"),
        ("arxiv", "arXiv"),
        ("pubmed", "PubMed"),
        ("openAccessPdf", "Open-access PDF metadata"),
        ("githubUrl", "Official GitHub URL"),
        ("abstract", "Abstract"),
    ]:
        count = int(df[col].fillna("").astype(str).str.len().gt(0).sum())
        coverage_rows.append([label, fmt_int(count), f"{count / len(df):.1%}"])
    add_table(doc, ["Evidence field", "Papers", "Coverage"], coverage_rows)

    add_heading(doc, "6.1 Link Audit and GitHub Provenance", 2)
    add_p(
        doc,
        f"The GitHub-link file uses a conservative policy: only official links explicitly indicated by paper text or Papers with Code official paper-mentioned flags are displayed. Among {fmt_int(github_links.get('paperCount', stats['selected']))} selected papers, {fmt_int(github_links.get('githubLinkCount', 0))} GitHub links are recorded. This protects against noisy third-party implementation matches, but it also means users who want every possible implementation must search beyond the official-link layer.",
    )
    add_p(
        doc,
        "The link-audit file separates DOI normalization, arXiv canonicalization, Semantic Scholar links, PDF availability, content-title or content-DOI matches, missing PDFs, and GitHub reachability. This matters because a citation map without link provenance can look authoritative while hiding broken or ambiguous sources.",
    )
    top_audit = sorted((link_audit.get("summary") or {}).items(), key=lambda item: item[1], reverse=True)[:10]
    add_table(doc, ["Link-audit status", "Count"], [[key, fmt_int(value)] for key, value in top_audit])

    add_heading(doc, "6.2 Dissemination Artifacts", 2)
    add_p(
        doc,
        f"The README is the static entry point. The website at {AWESOME_SITE} is the interactive entry point. The CSV and JSON files are the analytical entry points. The curation-method document explains scope, data source, ranking, enrichment, and provenance. The CITATION.cff file makes the curated dataset citeable. Together, these artifacts make awesome-ai closer to a living dataset than a one-time blog post.",
    )

    add_heading(doc, "7. Cross-Cutting Analysis", 1)
    cross = [
        ("7.1 Citation influence is not scientific closure", "Citation count measures reuse and attention, not truth. Highly cited papers can be methods, datasets, toolkits, surveys, standards, or even objects of criticism. A citation map should guide reading, not replace judgment."),
        ("7.2 Infrastructure papers shape the field quietly", "Scikit-learn, PyTorch, TensorFlow, XGBoost, LIBSVM, and benchmark datasets change the cost of experimentation. Their influence is often deeper than a single model result."),
        ("7.3 Benchmarks create progress and blind spots", "ImageNet-style benchmarks make comparison possible, but they also produce benchmark-centric incentives. Deployment robustness, rare cases, and distribution shift require additional evidence."),
        ("7.4 Foundation models collapse older boundaries", "Transformers, pretraining, retrieval, multimodal alignment, and agentic tool use connect NLP, vision, RL, software engineering, and AI for science. Taxonomy labels should be read as coordinates, not walls."),
        ("7.5 Governance is now technical interpretation", "Data provenance, copyright, privacy, release policy, alignment, safety, interpretability, and evaluation contamination are no longer peripheral issues. They shape how modern AI papers should be read."),
    ]
    for heading, text in cross:
        add_heading(doc, heading, 2)
        add_p(doc, text)

    add_heading(doc, "8. Practitioner-Oriented User Guide", 1)
    add_table(
        doc,
        ["User", "Primary question", "Recommended path", "Verification point"],
        [
            ["New researcher", "Where should I start?", "README overview -> taxonomy -> top category papers -> original papers.", "Correct citation lag for recent papers."],
            ["Graduate seminar", "How should I build a reading list?", "Combine four periods with ten categories.", "Add connecting and critical papers."],
            ["Survey writer", "How should I structure related work?", "Use taxonomy overview and limitations as an outline.", "Verify every claim against original papers."],
            ["Benchmark seeker", "Which baselines and datasets matter?", "Search category leaders, method tags, and GitHub links.", "Check code officialness, license, and environment."],
            ["Research planner", "Which areas combine influence and freshness?", "Compare period heatmap, yearly citation mass, and recent frontier rows.", "Use expert review beyond citation counts."],
        ],
    )
    add_heading(doc, "8.1 Recommended Reading Protocol", 2)
    add_numbered(
        doc,
        [
            "Select a taxonomy category and read its overview and limitations.",
            "Read the top five high-citation papers and five recent papers separately.",
            "Check DOI, arXiv, open-PDF, and official GitHub fields before using a paper as a baseline.",
            "Compare citationCount, influentialCitationCount, venue, abstract availability, and source queries.",
            "Use generated key ideas as navigation aids, not as final scholarly claims.",
        ],
    )

    add_heading(doc, "9. Open Challenges and Future Directions", 1)
    challenges = [
        ("Citation-lag correction", "Recent papers need complementary signals such as adoption, official code, benchmark use, and expert annotation."),
        ("Reproducibility-aware ranking", "Future maps should combine citations with code execution, data availability, license clarity, and artifact evaluation."),
        ("Multilingual and regional coverage", "English-language and large-venue metadata is overrepresented in public databases."),
        ("Category-overlap modeling", "A multi-label graph would better represent papers that influence several areas."),
        ("Human expert review loop", "Expert annotations could improve deterministic metadata enrichment and correct taxonomy errors."),
        ("Negative evidence layer", "Critical follow-up papers, failed replications, and retired assumptions would make the map more mature."),
    ]
    for heading, text in challenges:
        add_heading(doc, heading, 2)
        add_p(doc, text)

    add_heading(doc, "10. Detailed Discussion", 1)
    discussion = [
        "The strongest lesson from the reference paper is that a repository becomes a survey only when the list is transformed into an argument. For awesome-ai, the argument is that AI influence can be mapped through a transparent metadata pipeline, but the resulting map must be read with explicit awareness of citation bias, citation lag, provenance limits, and category overlap.",
        "A reader should not ask only which paper is ranked highest. The more productive questions are: what kind of artifact is this paper, why did it become a shared reference, what infrastructure did it create, what assumptions did it normalize, what later work corrected or extended it, and what evidence would be required before using it in a new project?",
        "The repository's evidence layers make this possible. Candidate files show the search space. Selected files show final choices. Taxonomy files add interpretive structure. Period analysis supports interactive exploration. GitHub-link files separate official code signals from noisy implementation matches. Link audits expose missing or questionable evidence. This is why the repository is more valuable than a static ranked list.",
        "The category structure also reveals interdependence. General methods provide the experimental language. Vision provides benchmark culture and representation-learning pressure. NLP provides transformer architectures and language-model evaluation. Reinforcement learning contributes feedback, agents, and control. AI for science tests models against external reality. Trustworthy AI asks whether systems should be trusted, not only whether they score well.",
        "The most important practical rule is to separate map use from claim use. It is reasonable to use awesome-ai to decide what to read, how to plan a seminar, or where to look for baselines. It is not reasonable to cite the generated key idea as if it were a verified expert summary. The original paper remains the authority for method, evidence, and limitations.",
    ]
    for paragraph in discussion:
        add_p(doc, paragraph)

    add_heading(doc, "10.1 Category Reading Guides", 2)
    guides = {
        "General AI Methods and Systems": "Read this category as the shared substrate of AI experimentation: algorithms, regularization, toolkits, surveys, textbooks, and software practices.",
        "Foundation Models and Large Language Models": "Read this category with freshness awareness. Adoption and ecosystem impact may matter before citations fully accumulate.",
        "Vision and Multimodal Learning": "Read this category through datasets, benchmarks, architectures, and distribution-shift concerns.",
        "AI for Science, Healthcare, and Robotics": "Read this category through external validation, safety, domain expertise, and regulatory evidence.",
        "Graph Learning, Recommendation, and Core Methods": "Read this category through data splits, leakage, popularity bias, and hyperparameter budgets.",
        "Reinforcement Learning and Agents": "Read this category through reward design, exploration risk, simulator bias, and oversight.",
        "Natural Language Processing and Knowledge": "Read this category through annotation protocols, retrieval settings, prompt sensitivity, and language coverage.",
        "Trustworthy, Explainable, and Responsible AI": "Read this category as the evaluation language that cuts across all other categories.",
        "Generative Models and Synthetic Media": "Read this category through quality, controllability, provenance, copyright, misuse, and factuality.",
        "Representation, Self-Supervised, and Transfer Learning": "Read this category as connective tissue across modern AI.",
    }
    for category, text in guides.items():
        add_heading(doc, category, 3)
        add_p(doc, text)

    add_heading(doc, "10.2 Interpretive Q&A", 2)
    qa = [
        ("Is the top-ranked paper the most important AI paper?", "No. It is the most cited paper under this snapshot and selection rule. Importance depends on the research question."),
        ("Does a low rank mean a recent paper is unimportant?", "No. Recent papers suffer from citation lag. Adoption and expert judgment are needed."),
        ("Does a GitHub link mean the work is reproducible?", "No. A link is a starting point. Environment, data, checkpoints, license, and execution must be checked."),
        ("Does one category label fully describe a paper?", "No. The label is a navigation coordinate. Many papers influence several areas."),
        ("Can the generated key idea be cited directly?", "It should not be used as an authoritative claim without reading the original paper."),
        ("What is the best use of awesome-ai?", "It is best used as an auditable starting point for reading, teaching, benchmark search, and research planning."),
    ]
    for question, answer in qa:
        add_p(doc, "Q. " + question, style="List Bullet")
        add_p(doc, "A. " + answer)

    add_heading(doc, "10.3 Practical Scenarios", 2)
    scenarios = [
        (
            "Scenario A: entering a new AI subfield.",
            "A new researcher should begin with the category overview, then separate historical anchors from current frontier papers. For example, a reader entering vision should not only read ResNet and ImageNet, but should also inspect ViT, CLIP, Segment Anything, robustness work, and dataset-bias discussions. The purpose is to avoid mistaking the most cited historical anchor for the current research frontier. Awesome-ai gives the starting coordinates, while the reader supplies the expert synthesis.",
        ),
        (
            "Scenario B: writing a related-work section.",
            "A related-work section should not simply list the top cited papers. It should explain how a problem evolved, which baselines became standard, which assumptions later changed, and which limitations remain unresolved. The taxonomy dataset can provide a first outline: category overview for paragraph structure, representative papers for anchors, and limitations for discussion. The author should then return to the original papers and write claims from primary evidence.",
        ),
        (
            "Scenario C: choosing baselines for experiments.",
            "A benchmark seeker should use citation rank only as one signal. A strong baseline should be relevant to the exact task, have usable code or enough implementation detail, use comparable data, and remain meaningful under current evaluation standards. A highly cited paper may be historically important but no longer an appropriate baseline. Conversely, a recent paper may be a strong baseline even before citation counts mature. GitHub-link provenance and link audits help, but do not replace local execution checks.",
        ),
        (
            "Scenario D: designing a graduate seminar.",
            "A seminar can use the four periods as modules and the ten categories as reading tracks. Each week should pair one historical anchor with one modern extension and one critical or evaluative reading. For instance, a foundation-model week might pair Attention, BERT or GPT-style scaling, an open-model technical report, and a safety or evaluation-contamination paper. This structure makes the map pedagogical rather than merely encyclopedic.",
        ),
        (
            "Scenario E: planning a research portfolio.",
            "A lab or program manager can use the map to identify areas where long-term influence, current momentum, and infrastructure readiness overlap. However, portfolio decisions should never be made from citation counts alone. Compute requirements, data access, regulatory risk, openness of tooling, community saturation, and local expertise matter. Awesome-ai is useful because it makes the first layer of evidence visible, not because it automates strategic judgment.",
        ),
    ]
    for heading, text in scenarios:
        add_p(doc, heading, style="List Bullet")
        add_p(doc, text)

    add_heading(doc, "11. Conclusion", 1)
    add_p(
        doc,
        "Awesome-ai is not a complete encyclopedia of AI research. It is more usefully understood as a transparent, citation-ranked, taxonomy-first map. It helps readers identify historical anchors, category leaders, recent frontier signals, source links, and metadata gaps. Its value comes from being inspectable, regenerable, and practical.",
    )
    add_p(
        doc,
        "The central message is that AI influence is built through many kinds of artifacts: algorithms, datasets, software libraries, benchmarks, architectures, pretraining recipes, model releases, safety methods, domain validations, and public repositories. A good roadmap should expose all of these layers while reminding the reader that influence is not the same as truth. Used carefully, awesome-ai can be a strong starting point for expert reading and research design.",
    )


def add_appendices_and_references(doc: Document, df: pd.DataFrame) -> None:
    add_heading(doc, "Appendix A. Repository and Data Inventory", 1)
    add_table(
        doc,
        ["Artifact", "Role"],
        [
            ["README.md", "Static GitHub entry point for the taxonomy-first citation map."],
            ["data/papers_2000_2026.csv", "Selected dataset of 2,700 papers."],
            ["data/papers_taxonomy_2000_2026.csv", "Taxonomy dataset with overviews, limitations, key ideas, and paper-level fields."],
            ["data/candidates_top1000_2000_2026.csv", "Audited candidate pool of 27,000 records."],
            ["data/period_analysis_2000_2026.json", "Period-level summaries for the interactive site."],
            ["data/github_links_2000_2026.json", "Official GitHub link matching and provenance."],
            ["data/link_audit_2000_2026.json", "Paper, PDF, and GitHub link-audit results."],
            ["docs/index.html", "Interactive GitHub Pages website."],
            ["paper/curation_method.md", "Scope, source, ranking, enrichment, and provenance statement."],
            ["CITATION.cff", "Citation metadata for the curated dataset."],
        ],
    )

    add_heading(doc, "Appendix B. Top Papers by Citation Count", 1)
    top_rows = []
    for _, row in df.sort_values("citationCount", ascending=False).head(30).iterrows():
        top_rows.append([int(row["rank"]), int(row.year), row.title, row.category, fmt_int(row.citationCount)])
    add_table(doc, ["Rank", "Year", "Title", "Category", "Citations"], top_rows)

    add_heading(doc, "Appendix C. Yearly Leading Papers", 1)
    yearly_rows = []
    for year, group in df.groupby("year"):
        top = group.sort_values("citationCount", ascending=False).iloc[0]
        yearly_rows.append([int(year), top.title, top.category, fmt_int(top.citationCount)])
    add_table(doc, ["Year", "Leading selected paper", "Category", "Citations"], yearly_rows)

    add_heading(doc, "References and Source Links", 1)
    refs = [
        ("honggi82/awesome-ai GitHub repository", AWESOME_GITHUB),
        ("Awesome AI interactive website", AWESOME_SITE),
        (f"Kong et al., {REFERENCE_TITLE}, arXiv:2605.18661", REFERENCE_ARXIV),
        ("worldbench/awesome-ai-auto-research GitHub repository", REFERENCE_GITHUB),
        ("Semantic Scholar Academic Graph API", "https://www.semanticscholar.org/product/api"),
        ("Papers with Code links-between-paper-and-code archive", "https://huggingface.co/datasets/pwc-archive/links-between-paper-and-code"),
    ]
    for _, row in df.sort_values("citationCount", ascending=False).head(35).iterrows():
        authors = clean(row.authors)
        if len(authors) > 115:
            authors = authors[:112] + "..."
        venue = clean(row.venue) or "Unknown venue"
        link = clean(row.url) or clean(row.semanticScholarUrl) or (f"https://doi.org/{clean(row.doi)}" if clean(row.doi) else "")
        refs.append((f"{authors}. {row.title}. {venue}, {int(row.year)}. Snapshot citations: {fmt_int(row.citationCount)}.", link))
    for index, (title, link) in enumerate(refs, 1):
        add_p(doc, f"[{index}] {title} {link}")


def main() -> None:
    df = pd.read_csv(DATA_DIR / "papers_2000_2026.csv")
    tax = pd.read_csv(DATA_DIR / "papers_taxonomy_2000_2026.csv")
    link_audit = json.loads((DATA_DIR / "link_audit_2000_2026.json").read_text(encoding="utf-8"))
    github_links = json.loads((DATA_DIR / "github_links_2000_2026.json").read_text(encoding="utf-8"))
    cat = (
        df.groupby("category")
        .agg(
            papers=("title", "count"),
            citations=("citationCount", "sum"),
            influential=("influentialCitationCount", "sum"),
            median_year=("year", "median"),
        )
        .reset_index()
        .sort_values(["papers", "citations"], ascending=False)
    )
    stats = {
        "years": f"{int(df.year.min())}-{int(df.year.max())}",
        "selected": len(df),
        "candidates": 27000,
        "citations": int(df.citationCount.fillna(0).sum()),
        "categories": df.category.nunique(),
    }
    periods = [
        ("2000-2008", 2000, 2008, "Statistical learning, data mining, kernel methods, and classical ML infrastructure."),
        ("2009-2016", 2009, 2016, "ImageNet, CNNs, representation learning, scalable infrastructure, and deep RL."),
        ("2017-2021", 2017, 2021, "Transformers, BERT, ViT, CLIP, AlphaFold, and multimodal foundation transitions."),
        ("2022-2026", 2022, 2026, "Instruction tuning, LLMs, interpretability, safety, AI4Science, and governance."),
    ]
    figures = make_figures(df, cat, periods)

    doc = configure_doc()
    add_title_page(doc, stats)
    add_reference_structure_analysis(doc)
    add_abstract(doc, stats, cat)
    add_introduction_and_preliminaries(doc, stats, figures)
    add_corpus_construction(doc, df, stats, figures)
    add_historical_roadmap(doc, df, periods, figures)
    add_taxonomy_analysis(doc, df, tax, cat, figures)
    add_validation_crosscutting_and_guide(doc, df, stats, link_audit, github_links, cat)
    add_appendices_and_references(doc, df)

    props = doc.core_properties
    props.title = "Awesome AI: Roadmap & User Guide"
    props.subject = "English survey-style Word paper based on honggi82/awesome-ai"
    props.author = "Codex, based on honggi82/awesome-ai public metadata"
    props.keywords = "AI survey, awesome-ai, citation map, taxonomy, roadmap, user guide"
    props.comments = "Generated from local awesome-ai data snapshot dated 2026-06-26."
    doc.save(OUT_PATH)

    check = Document(OUT_PATH)
    text = "\n".join(paragraph.text for paragraph in check.paragraphs)
    with ZipFile(OUT_PATH) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    if len(text) < 50000 or len(text.split()) < 7500:
        raise RuntimeError(f"English document is too thin: chars={len(text)}, words={len(text.split())}")
    if len(check.tables) < 10:
        raise RuntimeError(f"Expected at least 10 tables, found {len(check.tables)}")
    if len(media) < 7:
        raise RuntimeError(f"Expected at least 7 embedded images, found {len(media)}")
    for needle in [AWESOME_GITHUB, REFERENCE_ARXIV, REFERENCE_GITHUB, "Phase 1: Corpus Construction", "Interpretive Q&A"]:
        if needle not in text:
            raise RuntimeError(f"Missing required text: {needle}")
    print(OUT_PATH.resolve())
    print(f"docx_size={OUT_PATH.stat().st_size}")
    print(f"chars={len(text)}")
    print(f"words={len(text.split())}")
    print(f"paragraphs={len(check.paragraphs)}")
    print(f"tables={len(check.tables)}")
    print(f"media={len(media)}")


if __name__ == "__main__":
    main()
