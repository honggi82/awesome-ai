from __future__ import annotations

import collections
import json
import math
import re
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
DOCS_DIR = ROOT / "docs"
PAPER_DIR = ROOT / "paper"
FIG_DIR = PAPER_DIR / "figures"
OUT_PATH = PAPER_DIR / "awesome_ai_survey_paper_ko.docx"

REFERENCE_TITLE = "AI for Auto-Research: Roadmap & User Guide"
REFERENCE_ARXIV = "https://arxiv.org/abs/2605.18661"
REFERENCE_GITHUB = "https://github.com/worldbench/awesome-ai-auto-research"
AWESOME_GITHUB = "https://github.com/honggi82/awesome-ai"
AWESOME_SITE = "https://honggi82.github.io/awesome-ai/"


def clean(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    text = str(value)
    return "" if text.lower() == "nan" else text


def fmt_int(value: object) -> str:
    return f"{int(value):,}"


def split_tags(series: pd.Series) -> collections.Counter[str]:
    counter: collections.Counter[str] = collections.Counter()
    for value in series.fillna(""):
        for tag in re.split(r";\s*", str(value)):
            if tag:
                counter[tag] += 1
    return counter


def set_korean_font(run, size: float | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)


def add_p(doc: Document, text: str = "", style: str | None = None, align: int | None = None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_korean_font(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_korean_font(run)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_p(doc, item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        add_p(doc, item, style="List Number")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    set_korean_font(run, 8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in caption_paragraph.runs:
        run.font.size = Pt(9)
        run.font.italic = True


def top_papers_text(df: pd.DataFrame, n: int = 5) -> str:
    parts = []
    for _, row in df.sort_values("citationCount", ascending=False).head(n).iterrows():
        parts.append(f"{int(row.year)}년 {row.title}({fmt_int(row.citationCount)}회)")
    return "; ".join(parts)


def make_figures(df: pd.DataFrame, cat: pd.DataFrame, periods: list[tuple[str, int, int, str]]) -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    yearly = df.groupby("year").agg(papers=("title", "count"), citations=("citationCount", "sum")).reset_index()
    keyword_counts = split_tags(df["keywordTags"])
    method_counts = split_tags(df["methodTags"])
    short = {
        "General AI Methods and Systems": "General methods",
        "Foundation Models and Large Language Models": "Foundation models",
        "Vision and Multimodal Learning": "Vision/multimodal",
        "AI for Science, Healthcare, and Robotics": "AI4Science/health",
        "Graph Learning, Recommendation, and Core Methods": "Graphs/recsys/core",
        "Reinforcement Learning and Agents": "RL/agents",
        "Natural Language Processing and Knowledge": "NLP/knowledge",
        "Trustworthy, Explainable, and Responsible AI": "Trustworthy AI",
        "Generative Models and Synthetic Media": "Generative media",
        "Representation, Self-Supervised, and Transfer Learning": "Representation",
    }
    paths: dict[str, Path] = {}

    cat_plot = cat.copy()
    cat_plot["short"] = cat_plot["category"].map(short)
    path = FIG_DIR / "rich_category_counts.png"
    plt.figure(figsize=(10, 6.1))
    plt.barh(cat_plot["short"][::-1], cat_plot["papers"][::-1], color="#2563eb")
    plt.xlabel("Selected papers")
    plt.title("Awesome AI taxonomy: selected papers by category")
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["category"] = path

    path = FIG_DIR / "rich_yearly_citations.png"
    plt.figure(figsize=(10, 4.8))
    plt.plot(yearly["year"], yearly["citations"], marker="o", color="#0f766e", linewidth=2.2)
    plt.fill_between(yearly["year"], yearly["citations"], alpha=0.16, color="#0f766e")
    plt.xlabel("Publication year")
    plt.ylabel("Citations in selected set")
    plt.title("Citation mass by publication year")
    plt.grid(alpha=0.25)
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["yearly"] = path

    period_cat = []
    for label, start, end, _ in periods:
        group = df[(df.year >= start) & (df.year <= end)]
        counts = group.groupby("category").title.count()
        for category in cat["category"]:
            period_cat.append(
                {"period": label, "category": short.get(category, category), "count": int(counts.get(category, 0))}
            )
    heatmap = pd.DataFrame(period_cat).pivot(index="category", columns="period", values="count")
    heatmap = heatmap.loc[[short[category] for category in cat["category"]]]
    path = FIG_DIR / "rich_period_taxonomy_heatmap.png"
    plt.figure(figsize=(9.2, 6.0))
    plt.imshow(heatmap.values, aspect="auto", cmap="YlGnBu")
    plt.xticks(range(len(heatmap.columns)), heatmap.columns, rotation=20, ha="right")
    plt.yticks(range(len(heatmap.index)), heatmap.index)
    for i in range(heatmap.shape[0]):
        for j in range(heatmap.shape[1]):
            plt.text(j, i, str(int(heatmap.values[i, j])), ha="center", va="center", fontsize=8)
    plt.title("Taxonomy coverage across four historical periods")
    plt.colorbar(label="Selected papers")
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["heatmap"] = path

    path = FIG_DIR / "rich_keyword_tags.png"
    labels, values = zip(*keyword_counts.most_common(10))
    plt.figure(figsize=(10, 4.8))
    plt.bar(labels, values, color="#7c3aed")
    plt.xticks(rotation=28, ha="right")
    plt.ylabel("Papers")
    plt.title("Keyword-tag distribution")
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["keywords"] = path

    path = FIG_DIR / "rich_method_tags.png"
    labels, values = zip(*method_counts.most_common(12))
    plt.figure(figsize=(10, 4.8))
    plt.bar(labels, values, color="#b45309")
    plt.xticks(rotation=28, ha="right")
    plt.ylabel("Papers")
    plt.title("Method-tag distribution")
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["methods"] = path

    coverage = {
        "DOI": df["doi"].fillna("").astype(str).str.len().gt(0).sum(),
        "arXiv": df["arxiv"].fillna("").astype(str).str.len().gt(0).sum(),
        "PubMed": df["pubmed"].fillna("").astype(str).str.len().gt(0).sum(),
        "Open PDF": df["openAccessPdf"].fillna("").astype(str).str.len().gt(0).sum(),
        "GitHub": df["githubUrl"].fillna("").astype(str).str.len().gt(0).sum(),
        "Abstract": df["abstract"].fillna("").astype(str).str.len().gt(0).sum(),
    }
    path = FIG_DIR / "rich_evidence_coverage.png"
    plt.figure(figsize=(9, 4.8))
    plt.bar(coverage.keys(), coverage.values(), color="#0891b2")
    plt.ylim(0, len(df))
    plt.ylabel("Papers with field present")
    plt.title("Evidence and provenance coverage in selected dataset")
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["coverage"] = path

    path = FIG_DIR / "rich_four_phase_framework.png"
    fig, ax = plt.subplots(figsize=(10, 3.8))
    ax.axis("off")
    boxes = [
        ("P1 Corpus\\nConstruction", "queries, filters,\\ndeduplication"),
        ("P2 Historical\\nRoadmap", "periods, milestones,\\ncitation mass"),
        ("P3 Taxonomy\\nAnalysis", "10 categories, tags,\\nrepresentative papers"),
        ("P4 Validation &\\nDissemination", "link audit, site,\\nuser guide"),
    ]
    colors = ["#dbeafe", "#dcfce7", "#fef3c7", "#ede9fe"]
    for i, (title, body) in enumerate(boxes):
        x = 0.04 + i * 0.24
        rect = plt.Rectangle((x, 0.28), 0.2, 0.48, facecolor=colors[i], edgecolor="#334155", linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + 0.1, 0.58, title, ha="center", va="center", fontsize=12, fontweight="bold")
        ax.text(x + 0.1, 0.40, body, ha="center", va="center", fontsize=10)
        if i < len(boxes) - 1:
            ax.annotate("", xy=(x + 0.235, 0.52), xytext=(x + 0.205, 0.52), arrowprops=dict(arrowstyle="->", lw=1.3))
    ax.text(0.5, 0.12, "A living citation map turns repository artifacts into a research roadmap and practitioner guide.", ha="center", fontsize=10)
    plt.tight_layout()
    plt.savefig(path, dpi=230)
    plt.close()
    paths["framework"] = path
    return paths


def configure_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    doc.styles["Normal"].font.size = Pt(10.2)
    footer = section.footer.paragraphs[0]
    footer.text = "Awesome AI Roadmap & User Guide | data snapshot 2026-06-26 | honggi82/awesome-ai"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_korean_font(run, 8)
        run.font.color.rgb = RGBColor(95, 95, 95)
    return doc


def add_title_page(doc: Document, stats: dict[str, object]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Awesome AI: Roadmap & User Guide")
    run.bold = True
    set_korean_font(run, 23)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("2000-2026년 AI 연구의 인용 기반 분류 지도와 활용 가이드")
    run.bold = True
    set_korean_font(run, 16)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("honggi82/awesome-ai 저장소를 arXiv 2605.18661식 장문 서베이 구조로 재구성한 논문 초안")
    run.italic = True
    set_korean_font(run, 11)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"작성일: 2026-06-26 | 선택 논문: {fmt_int(stats['selected'])}편 | 후보 풀: {fmt_int(stats['candidates'])}건 | 총 인용: {fmt_int(stats['citations'])}회"
    )
    set_korean_font(run, 10)


def add_reference_structure_analysis(doc: Document) -> None:
    add_heading(doc, "0. 참고 논문 구성 분석과 본 문서의 대응 전략", 1)
    add_p(
        doc,
        f"참고 논문 {REFERENCE_TITLE}는 GitHub 큐레이션을 단순 목록으로 제시하지 않고, 저장소의 항목들을 하나의 학술적 프레임워크로 재해석한다. 논문의 첫 페이지는 프로젝트 페이지와 GitHub 저장소를 명시하여 자료가 갱신되는 living resource임을 밝히고, 본문은 연구 생애주기를 네 단계와 여덟 stage로 나눈다. 각 stage는 하위 방법군, 평가 기준, findings and observations, transition으로 반복 구성된다. 부록은 tool inventory와 taxonomy analysis로 저장소의 긴 목록을 논문형 근거 체계로 바꾼다.",
    )
    add_p(
        doc,
        "이번 문서는 그 방식을 그대로 복사하지 않고, awesome-ai의 성격에 맞게 전환한다. auto-research 논문에서 lifecycle이 핵심 골격이었다면, awesome-ai에서는 citation map lifecycle이 골격이다. 즉 자료 수집과 선별, 시대별 전환, 분류별 해석, 검증과 확산을 네 phase로 놓고, 각 phase마다 평가 기준과 관찰을 붙인다. 이렇게 해야 README의 표와 CSV의 통계가 단순한 목록을 넘어 논문적 주장으로 연결된다.",
    )
    add_table(
        doc,
        ["2605.18661의 구성 요소", "논문 안에서의 역할", "awesome-ai 논문에서의 대응"],
        [
            ["Title/Abstract with project links", "문제의식, 핵심 finding, 프로젝트·GitHub 연결", "AI 연구 지도의 필요성, 2,700편 데이터셋, GitHub Pages와 CSV 출처 연결"],
            ["Preliminaries", "lifecycle, method families, scope, timeline 정의", "citation-map lifecycle, metadata curation families, 범위·데이터 출처·역사 구간 정의"],
            ["Phase chapters", "각 연구 단계별 도구·평가·한계 분석", "corpus construction, historical roadmap, taxonomy analysis, validation/dissemination 분석"],
            ["Assessment sections", "각 단계에서 무엇이 잘 되고 무엇이 위험한지 판단", "인용 기반 ranking, metadata coverage, link audit, citation lag의 신뢰도 평가"],
            ["Findings and observations", "논문 주장으로 요약되는 관찰", "AI 연구 영향력이 dataset/tool/model/governance로 이동한다는 교차 관찰"],
            ["Appendix inventory", "GitHub의 긴 목록을 표로 보존", "상위 논문, 분류별 대표 논문, 연도별 이정표, 저장소 산출물 inventory"],
        ],
    )
    add_heading(doc, "0.1 GitHub 내용을 논문화하는 방식", 2)
    add_p(
        doc,
        "worldbench 저장소의 README는 논문의 public-facing mirror다. README는 badge, teaser, phase table, 목차, 분야별 논문 목록, tool/GitHub repo 목록을 제공하고, 논문은 이를 taxonomy와 benchmark suite, tool inventory로 압축한다. 중요한 점은 GitHub 목록을 그대로 읽어 내려가지 않는다는 것이다. 논문은 저장소가 모은 항목을 기반으로 어떤 문제가 성숙했는지, 어떤 stage가 평가 공백인지, 어떤 failure mode가 반복되는지를 해석한다.",
    )
    add_p(
        doc,
        "awesome-ai도 같은 접근이 필요하다. 저장소에는 README, selected dataset, taxonomy dataset, candidate pool, period analysis, GitHub link audit, interactive website, review drafts, curation method가 있다. 따라서 본 문서는 GitHub 내용을 다섯 층으로 설명한다. 첫째, 자료 생성 파이프라인. 둘째, 선정 데이터의 통계 구조. 셋째, AI 연구 시대별 전환. 넷째, 10개 분류의 근거와 한계. 다섯째, 사용자가 이 지도를 실제 연구·교육·기획에 쓰는 방법이다.",
    )


def add_abstract(doc: Document, stats: dict[str, object], cat: pd.DataFrame) -> None:
    top_cat = cat.iloc[0]
    top_citation_cat = cat.sort_values("citations", ascending=False).iloc[0]
    add_heading(doc, "초록", 1)
    add_p(
        doc,
        f"AI 연구는 2000년대의 데이터마이닝과 통계적 학습에서 2010년대의 딥러닝·대규모 벤치마크, 2020년대의 파운데이션 모델·멀티모달·생성형 AI·책임 있는 AI로 빠르게 이동해 왔다. 그러나 연구자가 이 장기 변화를 읽는 데 사용할 수 있는 자료는 대개 주제별 목록, 특정 분야 서베이, 또는 최신 모델 중심의 단기 리뷰에 머문다. 본 논문 초안은 honggi82/awesome-ai 저장소의 공개 메타데이터 기반 큐레이션을 분석하여, {stats['years']}년 AI 연구의 citation-ranked roadmap과 user guide를 제시한다. 이 저장소는 Semantic Scholar 공개 메타데이터를 바탕으로 연도별 최대 1,000건의 후보를 감사하고, 각 연도 상위 100편씩 총 {fmt_int(stats['selected'])}편을 선택했으며, 후보 풀 {fmt_int(stats['candidates'])}건과 선택 논문 총 {fmt_int(stats['citations'])}회의 인용 신호를 제공한다.",
    )
    add_p(
        doc,
        f"본 문서는 arXiv 2605.18661의 장문 로드맵형 서베이 구조를 참고하여 awesome-ai의 내용을 네 phase로 재구성한다. Phase 1은 corpus construction으로, 쿼리 설계, AI 관련성 필터, 중복 제거, 연도별 인용 랭킹, deterministic enrichment를 설명한다. Phase 2는 historical roadmap으로, 2000-2008, 2009-2016, 2017-2021, 2022-2026 네 시기의 전환을 분석한다. Phase 3은 taxonomy analysis로, {int(stats['categories'])}개 분류의 규모와 대표 논문, 검증 요구를 해석한다. Phase 4는 validation and dissemination으로, DOI·arXiv·PDF·GitHub 링크 감사, GitHub Pages 사이트, README와 데이터 파일, 재사용 가이드를 다룬다. 분석 결과 가장 많은 논문을 포함한 분류는 {top_cat.category}({fmt_int(top_cat.papers)}편)이고, 가장 큰 인용 질량은 {top_citation_cat.category}({fmt_int(top_citation_cat.citations)}회)에 있다. 이 차이는 AI 연구 영향력이 단순 논문 수가 아니라 데이터셋, 프레임워크, 모델 패밀리, 벤치마크, 도구 생태계의 결합으로 형성된다는 점을 보여준다.",
    )
    add_p(
        doc,
        "핵심 결론은 다섯 가지다. 첫째, 연도 균형형 citation map은 오래된 논문의 독점을 줄이면서도 장기 영향력을 보여준다. 둘째, 2009-2016년 딥러닝 전환과 2017년 이후 Transformer 전환은 인용 질량과 분류 이동에서 모두 확인된다. 셋째, foundation model 연구는 2022년 이후 논문 수를 빠르게 지배하지만, citation lag 때문에 최신성은 별도 검토가 필요하다. 넷째, GitHub link와 open-access PDF coverage는 재현성의 출발점일 뿐, 실행 가능성이나 공식성까지 보장하지 않는다. 다섯째, awesome-ai의 가장 좋은 사용법은 정답 목록이 아니라, 원논문 읽기·강의 설계·벤치마크 탐색·연구 기획을 시작하는 감사 가능한 지도다.",
    )
    add_heading(doc, "Abstract", 1)
    add_p(
        doc,
        f"This Korean survey draft turns honggi82/awesome-ai into a long-form roadmap and user guide in the style of the arXiv survey {REFERENCE_TITLE}. The repository constructs a citation-ranked map of AI research from {stats['years']} by auditing up to 1,000 candidates per year and selecting the top 100 papers per year, yielding {fmt_int(stats['selected'])} selected papers from {fmt_int(stats['candidates'])} candidates. The paper analyzes the collection through four phases: corpus construction, historical roadmap, taxonomy analysis, and validation/dissemination. It reports category-level influence, period-level transitions, metadata coverage, link-audit evidence, and practical guidance for literature review, curriculum design, benchmark search, and research planning.",
    )


def add_preliminaries(doc: Document, stats: dict[str, object], figures: dict[str, Path]) -> None:
    add_heading(doc, "1. Introduction", 1)
    add_p(
        doc,
        "AI 연구를 읽는 일은 점점 더 어려워지고 있다. 한쪽에는 ResNet, Attention, BERT, AlphaFold, LLaMA처럼 분야를 재편한 논문이 있고, 다른 한쪽에는 매년 발표되는 수천 편의 incremental work, benchmark report, technical report, dataset paper, system card가 있다. 연구자는 어떤 논문이 오래 살아남은 영향력을 갖는지, 어떤 흐름이 특정 시기의 유행인지, 어떤 분류에서 검증 기준이 달라지는지 빠르게 판단해야 한다.",
    )
    add_p(
        doc,
        "awesome-ai는 이 문제를 citation-ranked metadata curation으로 다룬다. 이 접근은 원논문 전문을 모두 읽는 expert survey와 다르다. 대신 대규모 공개 메타데이터에서 후보를 넓게 모으고, 연도별로 균형 잡힌 상위 인용 논문을 선택하고, 분류·키워드·방법 태그·대표 링크를 붙인다. 장점은 규모와 재현 가능성이다. 단점은 인용수와 메타데이터가 연구 품질, 최신성, 재현성, 윤리성을 충분히 대표하지 못한다는 점이다. 따라서 본 문서는 자료의 장점과 한계를 동시에 해석한다.",
    )
    add_p(
        doc,
        f"참고 논문 {REFERENCE_TITLE}가 강조한 것은 생산성보다 검증과 거버넌스였다. 본 문서도 비슷한 관점을 취한다. AI 연구 지도를 만드는 일에서 핵심은 단순히 많은 논문을 모으는 것이 아니라, 어떤 근거로 수집했고, 어떤 기준으로 제외했으며, 어떤 링크와 메타데이터가 검증되었고, 사용자가 어떤 해석상의 함정을 조심해야 하는지 밝히는 것이다.",
    )
    add_heading(doc, "1.1 Contributions", 2)
    add_bullets(
        doc,
        [
            f"{stats['years']}년 AI 연구 {fmt_int(stats['selected'])}편과 후보 {fmt_int(stats['candidates'])}건의 구조를 장문 논문형으로 설명한다.",
            "arXiv 2605.18661의 lifecycle-style survey 구조를 분석하고, 이를 citation map lifecycle로 변환한다.",
            "쿼리 설계, relevance filtering, deduplication, ranking, enrichment, link audit, website dissemination을 저장소 내부 근거와 연결한다.",
            "시대별 전환과 10개 taxonomy를 대표 논문·인용 질량·평가 한계와 함께 분석한다.",
            "신규 연구자, 세미나 운영자, 문헌조사 작성자, 벤치마크 탐색자, 연구기획자를 위한 사용자 가이드를 제공한다.",
        ],
    )
    add_figure(doc, figures["framework"], "Figure 1. awesome-ai를 논문화하기 위한 네 phase 프레임워크.")

    add_heading(doc, "2. Preliminaries", 1)
    add_heading(doc, "2.1 Citation-Map Lifecycle", 2)
    add_p(
        doc,
        "본 문서는 awesome-ai의 작업을 네 phase로 정의한다. P1 Corpus Construction은 후보 논문을 수집하고 AI 관련성을 판정하며 중복을 제거하고 연도별 랭킹을 만드는 단계다. P2 Historical Roadmap은 선택된 논문을 시간축 위에 배치해 연구 전환을 설명한다. P3 Taxonomy Analysis는 선택 집합을 10개 AI 연구 범주로 조직하고 각 범주의 대표 논문, 핵심 방법, 한계를 분석한다. P4 Validation & Dissemination은 링크 감사, GitHub code link, open PDF, README, GitHub Pages, citation file, license, 재생성 스크립트를 통해 자료를 배포하고 검토 가능하게 만드는 단계다.",
    )
    add_table(
        doc,
        ["Phase", "awesome-ai에서의 의미", "대표 파일/근거", "주요 위험"],
        [
            ["P1 Corpus Construction", "Semantic Scholar bulk search, relevance filter, deduplication, top-100-per-year selection", "scripts/build_awesome_ai.py, candidates_top1000_2000_2026.csv", "검색어 편향, 메타데이터 누락, 인용 편향"],
            ["P2 Historical Roadmap", "연도별·시기별 대표 논문과 인용 질량 해석", "papers_2000_2026.csv, period_analysis_2000_2026.json", "citation lag, 특정 benchmark 과대표집"],
            ["P3 Taxonomy Analysis", "10개 분류, keyword tags, method tags, category overview와 limitations", "papers_taxonomy_2000_2026.csv, README.md", "단일 주 분류 배치의 단순화"],
            ["P4 Validation & Dissemination", "link audit, GitHub links, open PDF, site, docx/html review outputs", "link_audit_2000_2026.json, github_links_2000_2026.json, docs/index.html", "링크 존재와 재현성의 혼동"],
        ],
    )
    add_heading(doc, "2.2 Metadata-Curation Families", 2)
    add_p(
        doc,
        "참고 논문이 prompt engineering, RAG, agentic method, training-based method, hybrid approach를 비교했듯이, awesome-ai에는 metadata curation family가 있다. 이들은 모델 학습 방법이 아니라 문헌 지도를 만드는 방법군이다. 핵심 family는 broad keyword search, deterministic relevance filtering, identifier-based deduplication, citation-balanced ranking, taxonomy enrichment, provenance auditing, multi-format dissemination이다.",
    )
    add_table(
        doc,
        ["Family", "역할", "awesome-ai 구현 단서", "평가 기준"],
        [
            ["Broad keyword search", "AI 전반 후보를 넓게 수집", "18개 Semantic Scholar query", "분야 누락 최소화와 노이즈 통제"],
            ["Relevance filtering", "제목·초록·venue의 AI 표현 확인", "AI_RELEVANCE_PATTERNS", "AI 관련성과 false positive 균형"],
            ["Deduplication", "동일 논문 중복 제거", "DOI, ArXiv, PubMed, CorpusId, paperId, normalized title", "동일 판본 통합과 잘못된 병합 방지"],
            ["Citation-balanced ranking", "연도별 상위 100편 선택", "citationCount descending, influentialCitationCount retained", "장기 영향력과 최신성 균형"],
            ["Taxonomy enrichment", "category, methodTags, keywordTags, keyIdea 생성", "CATEGORIES, KEYWORD_CONVENTION, deterministic rules", "해석 가능성과 과도한 단순화 사이 균형"],
            ["Provenance auditing", "paper/PDF/GitHub 링크 점검", "link_audit, github_links", "링크 신뢰성과 공식성"],
            ["Dissemination", "README, docs site, review drafts, citation file 배포", "README.md, docs/index.html, CITATION.cff", "사용자 접근성과 재현성"],
        ],
    )


def add_methods(doc: Document, df: pd.DataFrame, stats: dict[str, object], figures: dict[str, Path]) -> None:
    add_heading(doc, "3. Phase 1: Corpus Construction", 1)
    add_p(
        doc,
        "Corpus construction은 awesome-ai의 가장 중요한 방법론 장이다. README에 보이는 2,700편 목록은 최종 산출물일 뿐이고, 그 뒤에는 후보 수집, 필터링, 중복 제거, 랭킹, enrichment, 파일 생성이라는 파이프라인이 있다. 이 과정을 명시해야 독자는 결과를 권위로 받아들이는 대신, 어떤 종류의 지도인지 정확히 이해할 수 있다.",
    )
    add_heading(doc, "3.1 Search Scope and Candidate Collection", 2)
    add_p(
        doc,
        f"저장소의 범위는 {stats['years']}년이며, 후보 목표는 연도별 최대 1,000건이다. 검색어는 artificial intelligence, machine learning, deep learning, neural networks, support vector machines, data mining, pattern recognition, bayesian networks, foundation models, large language models, NLP, computer vision, reinforcement learning, generative AI diffusion models, graph neural networks, multimodal learning, AI safety fairness explainability robustness, AI for science healthcare robotics 등으로 구성된다. 이 설계는 AI를 LLM 중심으로 좁히지 않고, 고전 ML부터 안전성·AI4Science까지 넓게 포괄한다.",
    )
    add_p(
        doc,
        "이 방식의 장점은 recall이다. 초기 AI 연구와 현대 파운데이션 모델 연구가 같은 프레임에 들어온다. 그러나 넓은 검색은 잡음도 만든다. 예를 들어 classification, clustering, graph, robotics, clinical 같은 단어는 AI와 관련될 수 있지만 항상 AI 논문을 의미하지는 않는다. 그래서 relevance filtering과 bad-title filtering이 다음 단계로 필요하다.",
    )
    add_heading(doc, "3.2 Relevance Filtering and Deduplication", 2)
    add_p(
        doc,
        "AI relevance는 제목, 초록, venue metadata에서 AI 관련 표현이 등장하는지 확인하는 deterministic rule로 판정된다. 패턴에는 artificial intelligence, AI, machine learning, deep learning, neural network, SVM, data mining, transformer, language model, LLM, foundation model, computer vision, NLP, reinforcement learning, diffusion model, GAN, graph neural, self-supervised, explainable AI, robustness, fairness, privacy, RAG 등이 포함된다.",
    )
    add_p(
        doc,
        "중복 제거는 DOI, arXiv, PubMed, CorpusId, paperId, normalized title 순서로 이루어진다. 이 순서는 metadata curation에서 매우 중요하다. DOI와 arXiv는 강한 식별자이지만 모든 논문에 존재하지 않고, title만 쓰면 판본 차이나 동일 제목의 다른 자료를 잘못 병합할 수 있다. 여러 식별자를 단계적으로 쓰는 것은 완벽하지 않지만, 대규모 공개 메타데이터에서 실용적인 절충이다.",
    )
    add_heading(doc, "3.3 Ranking and Year Balancing", 2)
    add_p(
        doc,
        f"최종 선택은 각 연도에서 citationCount 기준 상위 100편이다. 전체 기간을 한 번에 정렬하지 않고 연도별로 고정 수를 뽑은 것은 중요한 설계다. 전체 인용 순위만 사용하면 2009-2017년의 초대형 딥러닝 논문과 오래된 교재·라이브러리 논문이 최근 논문을 거의 모두 밀어낸다. 반대로 연도별 선택은 2024-2026년의 Llama 3, Qwen, DeepSeek, interpretability, agentic engineering 같은 최신 신호를 지도에 포함시킨다.",
    )
    add_p(
        doc,
        "이 설계는 참고 논문의 lifecycle 관점과 유사하다. auto-research 논문이 각 stage를 따로 보아 특정 단계의 성숙도를 놓치지 않도록 했듯, awesome-ai는 각 연도를 따로 보아 특정 시대가 전체 지도를 독점하지 않도록 한다. 다만 같은 연도 안에서는 여전히 citation-rich field, 영어권 venue, survey/tool/dataset paper가 유리하다.",
    )
    add_heading(doc, "3.4 Enrichment and Output Generation", 2)
    add_p(
        doc,
        "선택 논문에는 category, methodTags, keywordTags, keyIdea, strengths, limitations, importanceReasons가 붙는다. 이 enrichment는 유료 LLM 리뷰가 아니라 deterministic metadata rules로 생성된다. 따라서 장점은 재생성 가능성과 비용 투명성이고, 한계는 원논문 전체를 읽은 전문가 요약만큼 깊지 않다는 점이다. 특히 keyIdea와 limitations는 문헌조사의 출발점으로는 유용하지만, 논문 본문에 인용 가능한 주장으로 쓰려면 원문 검토가 필요하다.",
    )
    add_heading(doc, "3.5 Assessment: Corpus Quality and Failure Modes", 2)
    add_p(
        doc,
        "Corpus construction의 품질은 세 기준으로 평가할 수 있다. 첫째, coverage: 중요한 AI 연구 흐름이 빠지지 않았는가. 둘째, precision: AI와 약하게 관련된 통계·의학·공학 논문이 과도하게 섞이지 않았는가. 셋째, auditability: 선택 이유와 링크가 추적 가능한가. awesome-ai는 후보 CSV, 선택 CSV, taxonomy CSV, period analysis, link audit를 남겨 auditability를 높였지만, precision과 coverage는 자동 규칙만으로 완전히 보장되지 않는다.",
    )
    add_heading(doc, "3.6 Findings and Observations", 2)
    add_bullets(
        doc,
        [
            "연도별 상위 100편 선택은 장기 영향력과 최신 흐름을 함께 보게 하는 유용한 compromise다.",
            "검색어가 넓기 때문에 AI 연구 주변부의 dataset, clinical, optimization, robotics 논문도 포함되며, 이것은 장점이자 잡음의 원천이다.",
            "deterministic enrichment는 투명하지만, 전문가 수준의 의미 분석을 대신하지 않는다.",
            "선택 집합을 읽을 때 citationCount와 influentialCitationCount, category, sourceQueries, abstract 유무를 함께 봐야 한다.",
        ],
    )
    add_figure(doc, figures["coverage"], "Figure 2. 선택 데이터셋의 DOI, arXiv, PubMed, Open PDF, GitHub, Abstract coverage.")


def add_historical_roadmap(
    doc: Document,
    df: pd.DataFrame,
    periods: list[tuple[str, int, int, str]],
    figures: dict[str, Path],
) -> None:
    add_heading(doc, "4. Phase 2: Historical Roadmap", 1)
    add_p(
        doc,
        "Historical roadmap은 awesome-ai를 단순한 ranking table에서 연구사 지도로 바꾸는 단계다. citation map에서 중요한 것은 어떤 논문이 1위인지뿐 아니라, 어떤 시기에 어떤 연구 방식이 dominant했는지, 어떤 분야가 다른 분야의 기반이 되었는지, 어떤 전환점 이후 새로운 평가 기준이 등장했는지다.",
    )
    period_rows = []
    for label, start, end, description in periods:
        group = df[(df.year >= start) & (df.year <= end)]
        top_year = int(group.groupby("year").citationCount.sum().sort_values(ascending=False).index[0])
        lead_category = group.category.value_counts().index[0]
        top = group.sort_values("citationCount", ascending=False).iloc[0]
        period_rows.append(
            [
                label,
                fmt_int(len(group)),
                fmt_int(group.citationCount.sum()),
                top_year,
                lead_category,
                f"{top.title} ({int(top.year)}, {fmt_int(top.citationCount)}회)",
                description,
            ]
        )
    add_table(doc, ["시기", "논문 수", "인용수", "최대 인용 연도", "최다 분류", "대표 논문", "해석"], period_rows)
    add_figure(doc, figures["yearly"], "Figure 3. 연도별 선택 논문의 인용 질량.")
    add_figure(doc, figures["heatmap"], "Figure 4. 네 역사 구간별 taxonomy coverage.")

    narratives = {
        "2000-2008": [
            "이 구간은 AI 연구가 통계적 학습, 데이터마이닝, 커널 방법, 특징 추출, classic computer vision, bioinformatics, recommender system의 언어로 조직되던 시기다. 상위 논문에는 The Elements of Statistical Learning, randomForest, Data Mining: Concepts and Techniques, SVM 관련 저작, independent component analysis, clustering survey가 나타난다.",
            "중요한 점은 이 시기를 '딥러닝 이전'이라는 결핍으로만 볼 수 없다는 것이다. 이 시기는 모델 평가, 일반화, feature engineering, optimization, software library, benchmark dataset의 관습을 만들었다. 후대 딥러닝 연구가 스스로를 설명하고 비교하는 언어 상당 부분이 이 시기에서 나왔다.",
            "평가상 주의할 점은 교재, survey, software library가 높은 인용을 얻기 쉽다는 것이다. 이들은 실제 실험 논문과 다른 방식으로 영향력을 갖는다. 따라서 이 시기의 citation map은 특정 알고리즘의 우열보다 연구 인프라와 교육적 표준의 형성을 보여준다고 읽는 편이 정확하다.",
        ],
        "2009-2016": [
            "이 구간은 ImageNet 이후 딥러닝 전환의 압축판이다. 2009년 ImageNet 데이터셋, 2012년 AlexNet, 2014년 VGG와 sequence-to-sequence, 2015년 ResNet·Faster R-CNN·BatchNorm·DQN, 2016년 XGBoost와 GCN이 이어지며 AI 연구의 중심이 대규모 데이터, GPU 학습, end-to-end representation learning으로 이동한다.",
            "Vision and Multimodal Learning의 인용 질량이 매우 커진 이유는 ImageNet-CNN 계열 논문이 학습 방식과 benchmark 문화를 동시에 바꿨기 때문이다. ResNet은 단일 모델을 넘어 깊은 network training을 가능하게 한 residual learning 패턴을 제공했고, Faster R-CNN과 YOLO 계열은 object detection의 실용화 경로를 열었다.",
            "이 시기에는 infrastructure 논문도 강하다. scikit-learn, TensorFlow, PyTorch, XGBoost, WEKA 같은 도구는 연구자가 실험을 반복하고 비교하는 방식을 바꾼다. citation map에서 이런 도구가 높은 위치에 있다는 것은 AI 발전이 알고리즘 아이디어뿐 아니라 재사용 가능한 구현과 생태계에 의존한다는 뜻이다.",
        ],
        "2017-2021": [
            "2017년 Attention is All You Need는 NLP 논문이면서 AI 전체 architecture의 공통 언어가 되었다. 이후 BERT, GPT 계열, T5, ViT, Swin Transformer, CLIP, AlphaFold가 등장하며 pretraining, transfer, scaling, multimodal alignment가 중심 축으로 떠오른다.",
            "이 구간에서 NLP와 vision의 경계는 약해진다. Transformer는 언어에서 시작했지만 vision transformer와 multimodal learning으로 이동했고, CLIP은 natural language supervision을 통해 시각 표현을 학습하는 방식을 대중화했다. AlphaFold는 AI for Science가 주변 응용이 아니라 scientific discovery infrastructure가 될 수 있음을 보여준다.",
            "평가 관점에서는 benchmark score의 의미가 복잡해진다. 모델 규모, 데이터 규모, 사전학습 corpus, prompt, fine-tuning protocol, leakage 가능성이 결과 해석에 큰 영향을 준다. 따라서 citation map은 영향력의 지도를 제공하지만, 성능 주장 자체는 각 논문의 설정을 따로 검토해야 한다.",
        ],
        "2022-2026": [
            "최근 구간은 foundation models가 논문 수를 빠르게 지배한다. InstructGPT, Chain-of-Thought prompting, LLaMA, Llama 2, Llama 3, GPT-4o system card, Qwen, DeepSeek-R1, interpretability와 agentic engineering 관련 항목이 등장한다. 이 구간은 연구 논문, 기술보고서, system card, open model release가 섞여 있다는 점도 특징이다.",
            "citation lag 때문에 2025-2026년 논문은 실제 중요도보다 낮게 보일 가능성이 크다. 예를 들어 2026년 항목은 아직 인용이 충분히 축적되지 않았기 때문에, citationCount만으로 중요도를 판단하면 안 된다. 이 구간에서는 venue, model adoption, code availability, community discussion, benchmark usage를 보조 신호로 읽어야 한다.",
            "또 하나의 변화는 governance와 safety다. Foundation model research는 성능 경쟁뿐 아니라 data governance, alignment, evaluation contamination, interpretability, deployment risk와 연결된다. 따라서 최근 구간은 '큰 모델이 더 잘한다'는 단순 서사가 아니라, 모델 능력·책임·공개성·재현성의 긴장으로 읽어야 한다.",
        ],
    }
    for label, _, _, _ in periods:
        add_heading(doc, f"4.{periods.index((label, next(p[1] for p in periods if p[0] == label), next(p[2] for p in periods if p[0] == label), next(p[3] for p in periods if p[0] == label))) + 1} {label}", 2)
        group = df[(df.year >= next(p[1] for p in periods if p[0] == label)) & (df.year <= next(p[2] for p in periods if p[0] == label))]
        for paragraph in narratives[label]:
            add_p(doc, paragraph)
        add_p(doc, "대표 상위 논문: " + top_papers_text(group, 6))
    add_heading(doc, "4.5 Assessment: Reading Time as Evidence", 2)
    add_p(
        doc,
        "시대별 로드맵은 인용수의 시간 편향을 줄이지만 제거하지는 않는다. 오래된 논문은 인용을 축적할 시간이 많고, 최근 논문은 빠르게 중요해져도 수치가 늦게 따라온다. 따라서 본 문서에서는 2022-2026년을 확정된 영향력 순위가 아니라 빠르게 이동 중인 frontier로 해석한다.",
    )
    add_heading(doc, "4.6 Findings and Observations", 2)
    add_bullets(
        doc,
        [
            "2000-2008년은 AI 방법론과 소프트웨어·교육 표준의 기반을 형성했다.",
            "2009-2016년은 ImageNet-CNN 전환과 reusable infrastructure가 citation mass를 지배했다.",
            "2017-2021년은 Transformer가 NLP를 넘어 vision, multimodal, science로 확장되는 시기다.",
            "2022-2026년은 foundation model과 governance가 동시에 부상하지만 citation lag가 매우 크다.",
        ],
    )


def add_taxonomy(doc: Document, df: pd.DataFrame, tax: pd.DataFrame, cat: pd.DataFrame, figures: dict[str, Path]) -> None:
    add_heading(doc, "5. Phase 3: Taxonomy Analysis", 1)
    add_p(
        doc,
        "Taxonomy analysis는 awesome-ai의 중심 장이다. 참고 논문이 auto-research를 stage별로 나눴듯, 이 문서는 AI 연구를 10개 큰 범주로 나누어 읽는다. 이 범주는 완전히 배타적인 ontological classification이 아니라, 연구자가 큰 지형을 읽기 위한 pragmatic taxonomy다. 한 논문이 여러 범주에 영향을 주더라도 지도에서는 주된 위치를 하나 부여한다.",
    )
    rows = []
    for _, row in cat.iterrows():
        top = df[df.category == row.category].sort_values("citationCount", ascending=False).iloc[0]
        rows.append(
            [
                row.category,
                fmt_int(row.papers),
                fmt_int(row.citations),
                fmt_int(row.influential),
                int(round(row.median_year)),
                f"{top.title} ({int(top.year)}, {fmt_int(top.citationCount)}회)",
            ]
        )
    add_table(doc, ["Category", "논문 수", "총 인용", "영향력 인용", "중앙 연도", "대표 상위 논문"], rows)
    add_figure(doc, figures["category"], "Figure 5. 분류별 선택 논문 수.")
    add_figure(doc, figures["keywords"], "Figure 6. Keyword tag 분포.")
    add_figure(doc, figures["methods"], "Figure 7. Method tag 분포.")

    category_interpretation = {
        "General AI Methods and Systems": "이 범주는 특정 응용보다 AI 연구의 공통 기반을 만든 논문을 담는다. XGBoost, LIBSVM, Dropout, Elements of Statistical Learning, WEKA, randomForest 같은 논문은 model family, training regularization, software library, 교육 표준을 제공한다. 논문 수가 가장 많은 이유는 범주가 넓기도 하지만, AI 연구의 기반 도구와 survey가 장기간 인용되기 때문이다.",
        "Foundation Models and Large Language Models": "이 범주는 LLM 자체뿐 아니라 PyTorch, scikit-learn처럼 foundation model 시대를 가능하게 한 software substrate도 포함한다. 2023년 이후 중앙 연도가 매우 최신으로 이동한다는 점이 중요하다. GPT-style scaling, instruction tuning, open foundation models, system cards, retrieval augmentation은 이제 별도의 subfield가 아니라 AI 연구 전반의 운영 체계가 되었다.",
        "Vision and Multimodal Learning": "이 범주는 총 인용 질량이 가장 크다. ImageNet, AlexNet, VGG, ResNet, Faster R-CNN, YOLO, ViT, CLIP, Swin Transformer 같은 논문은 computer vision을 넘어 representation learning과 benchmark culture를 바꿨다. 강점은 benchmark와 공개 dataset이 분명하다는 것이고, 약점은 web-scale bias와 distribution shift를 citation map만으로 평가하기 어렵다는 것이다.",
        "AI for Science, Healthcare, and Robotics": "이 범주는 AI가 과학·의료·자율주행·로보틱스 등 외부 세계와 접촉하는 지점을 보여준다. AlphaFold, KITTI, medical image analysis survey, graph attention networks 등이 대표적이다. 이 범주에서는 retrospective benchmark보다 external validation, prospective testing, safety, regulation, domain expert review가 중요하다.",
        "Graph Learning, Recommendation, and Core Methods": "그래프와 추천, optimization, Bayesian method, architecture search는 AI 시스템의 구조적 inductive bias와 효율성을 담당한다. GCN, GAT, GIN, TensorFlow, LIME 같은 항목은 graph benchmark, explanation, scalable training infrastructure로 연결된다. 인기 기반 bias, temporal leakage, hyperparameter budget 같은 평가 문제가 특히 중요하다.",
        "Reinforcement Learning and Agents": "RL 범주는 DQN, PPO, DDPG, AlphaGo, MAML, offline/robotic control 흐름을 포함한다. 최근 agentic AI와 human feedback이 부상하면서 RL은 단순 게임 benchmark에서 language-agent, tool-use, alignment로 재해석되고 있다. 다만 sample inefficiency, reward misspecification, simulator bias, safety under exploration은 계속 남는다.",
        "Natural Language Processing and Knowledge": "NLP 범주는 Attention, BERT, sequence-to-sequence, T5 등 지식 처리와 언어 모델링의 핵심 전환을 담는다. 이 범주의 영향력은 LLM 범주와 겹치지만, 정보 검색, 번역, 질의응답, summarization, knowledge-intensive reasoning의 평가 문제가 별도로 존재한다. 영어 중심 데이터와 prompt sensitivity도 핵심 한계다.",
        "Trustworthy, Explainable, and Responsible AI": "이 범주는 adversarial robustness, explainability, privacy, fairness, uncertainty, safety를 포괄한다. 인용수는 vision이나 LLM보다 작지만, deployment risk를 해석하는 데 필수적인 범주다. 설명가능성은 설득력과 faithful explanation이 다를 수 있고, fairness·privacy·robustness·utility는 서로 tradeoff를 만든다.",
        "Generative Models and Synthetic Media": "Generative category는 GAN, diffusion, score-based modeling, super-resolution, synthetic media, PINN 같은 흐름을 포함한다. 생성 품질은 빠르게 좋아졌지만, factuality, provenance, copyright, controllability, misuse, evaluation reliability가 별도 문제로 남는다. citation map에서는 model family와 evaluation metric의 영향이 특히 크게 나타난다.",
        "Representation, Self-Supervised, and Transfer Learning": "이 범주는 word2vec, transfer learning survey, adversarial examples, contrastive predictive coding 등 label efficiency와 transferability의 기반을 담는다. 범주 규모는 작지만 다른 거의 모든 분야로 흘러 들어가는 connective tissue 역할을 한다. downstream protocol, data overlap, augmentation choice가 성능 해석을 크게 바꾼다.",
    }
    for index, category in enumerate(cat["category"], 1):
        group = tax[tax.category == category].copy()
        top5 = group.sort_values("citationCount", ascending=False).head(5)
        overview = clean(group["categoryOverview"].dropna().iloc[0]) if group["categoryOverview"].notna().any() else ""
        limitations = clean(group["categoryLimitations"].dropna().iloc[0]) if group["categoryLimitations"].notna().any() else ""
        add_heading(doc, f"5.{index} {category}", 2)
        add_p(doc, category_interpretation.get(category, overview))
        if overview:
            add_p(doc, "저장소 overview: " + overview)
        if limitations:
            add_p(doc, "평가와 활용의 한계: " + limitations)
        add_p(doc, "대표 논문: " + "; ".join(f"{int(row.year)}년 {row.title}({fmt_int(row.citationCount)}회)" for _, row in top5.iterrows()))
        add_p(
            doc,
            "해석상 주의점: 이 범주의 대표 논문은 해당 범주의 모든 지적 다양성을 대표하지 않는다. citation-ranked map은 영향력이 큰 공통 참조점을 보여주지만, 최신 연구 질문과 비주류 아이디어는 별도 검색과 원문 검토로 보완해야 한다.",
        )
    add_heading(doc, "5.11 Summary and Transition: Taxonomy", 2)
    add_p(
        doc,
        "10개 taxonomy를 함께 보면 AI 연구는 하나의 직선적 발전사가 아니라 여러 인프라 층의 축적으로 보인다. general methods와 software libraries는 반복 실험을 가능하게 했고, vision benchmark는 representation learning의 실험장을 제공했으며, NLP와 Transformer는 multimodal 및 foundation model로 확장되었다. trustworthy AI와 AI4Science는 성능 중심 지도에 governance와 external validation이라는 기준을 덧붙인다.",
    )


def add_validation_and_crosscutting(
    doc: Document,
    df: pd.DataFrame,
    stats: dict[str, object],
    link_audit: dict[str, object],
    github_links: dict[str, object],
) -> None:
    add_heading(doc, "6. Phase 4: Validation and Dissemination", 1)
    add_p(
        doc,
        "Validation and dissemination은 참고 논문에서 peer review와 Paper2X가 하던 역할을 awesome-ai에 맞게 변환한 장이다. 여기서 검증은 논문 주장의 peer review가 아니라, 메타데이터 지도에서 링크와 출처, 파일, 재생성 가능성이 얼마나 추적 가능한지를 의미한다. 배포는 README와 GitHub Pages, CSV/JSON, docx/html 초안을 통해 사용자가 자료를 읽고 재사용하게 만드는 과정이다.",
    )
    coverage_rows = []
    for col, label in [
        ("doi", "DOI"),
        ("arxiv", "arXiv"),
        ("pubmed", "PubMed"),
        ("openAccessPdf", "Open-access PDF metadata"),
        ("githubUrl", "Official GitHub URL"),
        ("abstract", "Abstract"),
    ]:
        count = int(df[col].fillna("").astype(str).str.len().gt(0).sum())
        coverage_rows.append([label, fmt_int(count), f"{count / len(df):.1%}"])
    add_table(doc, ["Evidence field", "논문 수", "비율"], coverage_rows)
    add_heading(doc, "6.1 Link Audit and GitHub Provenance", 2)
    add_p(
        doc,
        f"github_links_2000_2026.json은 공식 GitHub 링크만 표시하려는 보수적 정책을 사용한다. 선택 논문 {fmt_int(github_links.get('paperCount', stats['selected']))}편 중 GitHub link count는 {fmt_int(github_links.get('githubLinkCount', 0))}건이며, matching은 arXiv, paper text, title 기반으로 구분된다. 이 정책은 broad title match나 제3자 reimplementation을 피하려는 장점이 있지만, 실제 사용 가능한 구현을 더 많이 찾고 싶은 사용자에게는 보수적으로 보일 수 있다.",
    )
    summary = link_audit.get("summary", {})
    add_p(
        doc,
        "link_audit_2000_2026.json은 paper link, PDF link, GitHub link의 상태를 기록한다. 예를 들어 DOI 형식 점검, arXiv canonicalization, GitHub reachable official provenance, missing GitHub, PDF missing, content-title/DOI match 같은 상태가 분리된다. 이는 README의 링크가 단순 장식이 아니라, 자료의 신뢰도와 추적 가능성을 보여주는 evidence layer라는 뜻이다.",
    )
    top_audit = sorted(summary.items(), key=lambda item: item[1], reverse=True)[:10]
    add_table(doc, ["Link-audit status", "Count"], [[key, fmt_int(value)] for key, value in top_audit])
    add_heading(doc, "6.2 README and Interactive Website as Dissemination", 2)
    add_p(
        doc,
        f"README는 GitHub 사용자를 위한 static entry point이고, {AWESOME_SITE}는 interactive entry point다. README는 project links, keyword convention, taxonomy overview, taxonomy collections, yearly coverage, methodology, caveats를 제공한다. docs/index.html은 기간 필터, taxonomy, category card, period analysis를 통해 사용자가 특정 기간과 분류를 탐색하도록 한다. 이 구조는 참고 논문의 GitHub repository가 paper와 project page를 함께 제공하는 방식과 닮아 있다.",
    )
    add_p(
        doc,
        "중요한 것은 dissemination이 연구의 끝이 아니라 feedback loop라는 점이다. 사용자가 링크 오류, 분류 오류, 누락 논문, citation lag 문제를 발견하면 저장소는 갱신될 수 있다. 따라서 awesome-ai는 고정된 PDF 서베이보다 living dataset에 가깝고, 본 Word 논문은 그 스냅샷을 해석하는 companion paper로 보는 것이 적절하다.",
    )

    add_heading(doc, "7. Cross-Cutting Analysis", 1)
    insights = [
        (
            "7.1 Citation influence is not scientific closure",
            "인용 영향력은 공동체가 어떤 논문을 많이 참조했는지 보여주지만, 그것이 재현성·윤리성·최신성을 보장하지 않는다. 특히 dataset, toolkit, survey, benchmark paper는 실험 논문과 다른 방식으로 인용된다. 따라서 citation map은 연구 지도의 시작이지 최종 판단이 아니다.",
        ),
        (
            "7.2 Infrastructure papers shape research more quietly than model papers",
            "scikit-learn, PyTorch, TensorFlow, XGBoost, LIBSVM 같은 항목은 특정 성능 수치보다 더 깊은 영향을 남긴다. 이들은 후속 연구가 실험을 설계하고 비교하고 배포하는 방식을 바꾸며, citation map에서 software infrastructure가 높은 위치에 있다는 점은 AI 발전이 code ecology와 분리되지 않음을 보여준다.",
        ),
        (
            "7.3 Benchmark concentration creates both progress and blindness",
            "ImageNet, KITTI, COCO, language benchmarks 같은 표준은 연구를 빠르게 비교하게 해주지만, benchmark overfitting과 domain shift를 만든다. citation-ranked map은 benchmark 중심 논문의 영향력을 잘 보여주지만, 실제 deployment 실패나 minority setting의 문제는 덜 드러낸다.",
        ),
        (
            "7.4 Foundation models collapse old boundaries",
            "2017년 이후 NLP, vision, multimodal, RL, AI4Science의 경계는 느슨해진다. Transformer와 pretraining은 method family로 여러 범주를 통과하고, LLM은 단일 분류가 아니라 search, coding, agents, robotics, education, governance와 연결된다. taxonomy는 이 복잡성을 읽기 위한 도구이지 완전한 구획이 아니다.",
        ),
        (
            "7.5 Governance becomes part of technical interpretation",
            "최근 AI 연구에서는 safety, privacy, fairness, interpretability, data provenance, copyright, model release policy가 기술적 성능과 분리되지 않는다. 따라서 awesome-ai의 caveat와 link audit는 부가 정보가 아니라 연구 지도 자체의 일부다.",
        ),
    ]
    for heading, text in insights:
        add_heading(doc, heading, 2)
        add_p(doc, text)


def add_user_guide_and_appendix(doc: Document, df: pd.DataFrame, cat: pd.DataFrame) -> None:
    add_heading(doc, "8. Practitioner-Oriented User Guide", 1)
    add_p(
        doc,
        "참고 논문의 user guide는 stage별로 어떤 도구를 어떻게 써야 하는지 제안한다. awesome-ai의 user guide는 어떤 독자가 어떤 방식으로 citation map을 읽을지에 초점을 맞춘다. 같은 데이터라도 신규 연구자, 세미나 운영자, 문헌조사 작성자, 벤치마크 탐색자, 연구기획자가 얻어야 하는 답은 다르다.",
    )
    add_table(
        doc,
        ["사용자", "1차 질문", "추천 경로", "검증 포인트"],
        [
            ["신규 연구자", "AI 전체 지형을 어디서 시작할까?", "README overview -> 10개 taxonomy -> 각 범주 top 10 paper -> 원논문", "최신 논문은 citation lag 보정"],
            ["대학원 세미나", "학기별 읽기 목록을 어떻게 만들까?", "네 시대 구간과 10개 taxonomy를 조합해 주차 배치", "분류 간 연결 논문을 토론 주제로 설정"],
            ["문헌조사 작성자", "관련 연구 섹션을 어떻게 구조화할까?", "papers_taxonomy CSV의 category overview와 limitations를 outline으로 사용", "keyIdea는 원문 확인 전 주장으로 쓰지 않음"],
            ["벤치마크 탐색자", "대표 baseline과 dataset은 무엇인가?", "category top papers, methodTags, benchmark/dataset tag, GitHub links 확인", "코드 공식성, license, 실행 가능성 확인"],
            ["연구기획자", "어떤 영역이 장기 영향력과 최신성을 동시에 갖나?", "period heatmap, yearly citation mass, recent foundation-model rows 비교", "인용수 외 adoption, venue, expert review 보완"],
        ],
    )
    add_heading(doc, "8.1 Recommended Reading Protocol", 2)
    add_numbered(
        doc,
        [
            "관심 주제의 taxonomy를 하나 고르고, 해당 범주의 overview와 limitations를 먼저 읽는다.",
            "상위 인용 논문 5편과 최근 3년 논문 5편을 분리해 읽는다. 이렇게 하면 역사적 영향력과 최신 frontier를 혼동하지 않는다.",
            "각 논문에 대해 DOI/arXiv/open PDF/GitHub link가 있는지 확인하고, 공식 구현과 제3자 구현을 구분한다.",
            "citationCount만 보지 말고 influentialCitationCount, venue, publication type, abstract 유무, sourceQueries를 함께 본다.",
            "최종 문헌조사에는 awesome-ai의 요약을 그대로 붙이지 말고, 원논문에서 method, experiment, limitation을 직접 확인한다.",
        ],
    )
    add_heading(doc, "9. Open Challenges and Future Directions", 1)
    challenges = [
        ("Citation lag correction", "최근 논문은 인용 누적 시간이 부족하므로 venue prestige, code adoption, model downloads, benchmark use, expert annotation 같은 보완 지표가 필요하다."),
        ("Reproducibility-aware ranking", "인용수와 함께 official code, data availability, environment reproducibility, license clarity, artifact evaluation 결과를 결합하는 ranking이 필요하다."),
        ("Multilingual and regional coverage", "영어권 venue와 metadata coverage가 강하기 때문에 비영어권·지역 학회·산업 보고서가 과소대표될 수 있다."),
        ("Category overlap modeling", "단일 taxonomy label은 편리하지만 Transformer, CLIP, AlphaFold 같은 논문의 다중 영향 경로를 충분히 표현하지 못한다. multi-label graph가 향후 필요하다."),
        ("Human expert review loop", "deterministic metadata enrichment는 빠르지만 깊은 의미 해석에는 한계가 있다. 분야별 전문가 검토를 저장소 기여 workflow로 연결하면 품질이 크게 오른다."),
    ]
    for title, text in challenges:
        add_heading(doc, title, 2)
        add_p(doc, text)
    add_heading(doc, "10. Conclusion", 1)
    add_p(
        doc,
        "awesome-ai는 AI 연구 전체를 한 번에 설명하는 완전한 백과사전이 아니다. 오히려 더 유용하게도, 공개 메타데이터와 재생성 가능한 규칙으로 만들어진 citation-ranked map이다. 이 지도는 연구자가 어디서부터 읽을지, 어떤 시대 전환을 이해해야 할지, 어떤 범주의 대표 논문과 한계를 확인해야 할지 알려준다. 본 문서는 참고 논문 2605.18661의 장문 로드맵형 방식을 따라, GitHub 저장소의 자료를 방법론, 시대별 분석, taxonomy, validation, dissemination, user guide로 재구성했다.",
    )
    add_p(
        doc,
        "핵심 메시지는 단순하다. AI 연구의 영향력은 모델 성능표 하나로 설명되지 않는다. 데이터셋, 소프트웨어, benchmark, architecture, pretraining recipe, safety framework, domain validation, 공개 링크와 license가 함께 지식을 만든다. awesome-ai는 이 복잡한 지형을 읽는 출발점이며, 사용자는 이 지도를 원논문 읽기와 전문가 판단으로 이어갈 때 가장 큰 가치를 얻는다.",
    )

    add_heading(doc, "Appendix A. Repository and Data Inventory", 1)
    add_table(
        doc,
        ["파일/리소스", "역할"],
        [
            ["README.md", "GitHub 사용자를 위한 taxonomy-first citation map 요약"],
            ["data/papers_2000_2026.csv", "연도별 top 100 선택 논문 2,700편의 핵심 데이터"],
            ["data/papers_taxonomy_2000_2026.csv", "분류별 overview, limitations, key ideas, strengths를 포함한 taxonomy 데이터"],
            ["data/candidates_top1000_2000_2026.csv", "연도별 최대 1,000건 후보, 총 27,000건"],
            ["data/period_analysis_2000_2026.json", "기간별 다국어 분석과 interactive site용 summary"],
            ["data/github_links_2000_2026.json", "공식 GitHub link matching과 provenance"],
            ["data/link_audit_2000_2026.json", "paper/PDF/GitHub link audit 결과"],
            ["docs/index.html", "GitHub Pages interactive website"],
            ["paper/curation_method.md", "범위, 데이터 출처, ranking, enrichment, provenance 설명"],
            ["CITATION.cff", "데이터셋 인용 정보"],
        ],
    )
    add_heading(doc, "Appendix B. Top Papers by Citation Count", 1)
    top_rows = []
    for _, row in df.sort_values("citationCount", ascending=False).head(30).iterrows():
        top_rows.append([int(row["rank"]), int(row.year), row.title, row.category, fmt_int(row.citationCount)])
    add_table(doc, ["Rank", "Year", "Title", "Category", "Citations"], top_rows)
    add_heading(doc, "Appendix C. Yearly Leading Papers", 1)
    yearly_rows = []
    for year, group in df.groupby("year"):
        top = group.sort_values("citationCount", ascending=False).iloc[0]
        yearly_rows.append([int(year), top.title, top.category, fmt_int(top.citationCount)])
    add_table(doc, ["Year", "Leading selected paper", "Category", "Citations"], yearly_rows)


def add_deep_discussion(doc: Document, df: pd.DataFrame, cat: pd.DataFrame) -> None:
    add_heading(doc, "11. Detailed Discussion: 2605.18661식 해설 확장", 1)
    add_p(
        doc,
        "참고 논문이 강한 이유는 각 도구 목록을 단순 나열하지 않고, 단계마다 평가 기준과 실패 양상을 반복해서 묻기 때문이다. 본 문서에서도 같은 방식을 적용하면 awesome-ai의 의미가 더 선명해진다. 질문은 네 가지다. 이 자료는 무엇을 잘 포착하는가. 무엇을 구조적으로 놓치는가. 사용자가 어떤 순서로 읽어야 하는가. 그리고 어떤 판단은 여전히 사람의 몫으로 남는가. 이 장은 앞선 통계와 표를 실제 연구 의사결정 언어로 다시 풀어 쓴다.",
    )
    add_heading(doc, "11.1 From Awesome List to Scholarly Argument", 2)
    for text in [
        "일반적인 awesome list는 유용하지만, 그 자체로는 학술 논문이 아니다. 항목이 많다는 사실은 충분한 근거가 되지 않는다. 논문이 되려면 목록을 조직하는 원리, 포함과 제외 기준, 범주 간 관계, 사용자에게 전달하려는 주장, 한계와 검증 방법이 필요하다. worldbench 저장소가 논문에서 한 일은 바로 이것이다. AI 자동 연구 도구 목록을 생성, 작성, 검증, 확산이라는 생애주기로 재배치함으로써 목록을 연구 지도로 바꿨다.",
        "awesome-ai도 같은 전환이 필요하다. README의 taxonomy collections는 이미 좋은 출발점이지만, 사용자는 왜 General AI Methods가 가장 큰지, 왜 Vision의 인용 질량이 압도적인지, 왜 Foundation Models의 중앙 연도가 매우 최신인지, 왜 Representation 범주는 작지만 중요한지 설명을 원한다. 본 논문은 이런 질문에 답하기 위해 저장소를 데이터셋, 방법론, 역사, 분류, 검증, 활용이라는 층으로 나누었다.",
        "이 전환에서 핵심은 인용수의 의미를 과장하지 않는 것이다. 인용수는 공동체 사용 흔적이다. 좋은 논문, 유명한 도구, 널리 쓰인 데이터셋, 표준 교재, 비판 대상이 된 논문 모두 인용될 수 있다. 따라서 citation-ranked map은 영향력 지도이지 품질 보증서가 아니다. 연구자는 이 지도를 통해 읽을 순서를 정하고, 최종 판단은 원논문과 실험 재현성, 도메인 지식으로 내려야 한다.",
    ]:
        add_p(doc, text)

    add_heading(doc, "11.2 Citation-Influence Challenge", 2)
    for text in [
        "참고 논문에는 capability-integrity challenge가 있다. AI 시스템이 연구 형태의 산출물을 빠르게 만들 수 있지만, 그 산출물이 진짜로 새롭고 정확하고 검증 가능한지는 별도 문제라는 주장이다. awesome-ai에는 이에 대응하는 citation-influence challenge가 있다. 인용수가 높은 논문은 영향력이 크지만, 그 영향력이 지금도 타당한지, 재현 가능한지, 특정 benchmark에 과적합된 것은 아닌지, 사회적 위험을 충분히 다뤘는지는 다른 문제다.",
        "예를 들어 ImageNet 계열 논문은 vision 연구를 재편했지만, 대규모 이미지 벤치마크의 성공이 실제 환경의 분포 변화와 rare class 문제를 자동으로 해결하지는 않는다. Attention과 BERT는 언어 모델 연구의 공통 언어가 되었지만, benchmark contamination, prompt sensitivity, multilingual coverage 문제는 계속 남아 있다. AlphaFold는 AI for Science의 상징이지만, 모든 과학 문제가 단일 예측 모델로 해결된다는 뜻은 아니다.",
        "따라서 사용자는 citation map을 두 겹으로 읽어야 한다. 첫 번째 겹은 영향력이다. 어떤 논문이 다른 논문들이 계속 참조하는 기반인가를 본다. 두 번째 겹은 검증 요구다. 그 영향력이 어떤 데이터, 어떤 평가, 어떤 도메인 제약, 어떤 사회적 조건에 묶여 있는지 본다. 이 두 겹을 분리해서 읽는 것이 장문 서베이의 핵심 독해법이다.",
    ]:
        add_p(doc, text)

    add_heading(doc, "11.3 Category Interdependence", 2)
    for text in [
        "10개 분류는 독립적인 서랍이 아니라 서로 영향을 주는 층이다. General AI Methods는 다른 거의 모든 분류의 실험 언어를 제공한다. Vision and Multimodal Learning은 대규모 데이터셋과 representation learning의 실험장을 제공했다. NLP and Knowledge는 Transformer를 통해 foundation model 시대의 architecture 언어를 제공했다. Reinforcement Learning은 human feedback과 agentic system을 통해 LLM alignment와 연결된다.",
        "AI for Science, Healthcare, and Robotics는 외부 세계의 검증 문제를 강하게 드러낸다. 같은 모델이라도 논문 벤치마크에서 좋은 것과 병원, 실험실, 로봇, 자율주행 환경에서 안전한 것은 다르다. Trustworthy AI는 이 차이를 해석하는 기준을 제공한다. 즉 responsible AI는 별도의 주변 분야가 아니라, 모든 응용 분야의 평가 기준을 재정의하는 교차 축이다.",
        "Generative Models와 Foundation Models도 점점 겹친다. diffusion model은 이미지 생성에서 시작했지만 멀티모달 생성과 scientific design으로 이동하고, LLM은 text generation을 넘어 code, planning, search, tool use, agent로 확장된다. 이런 상황에서 단일 category label은 불가피한 단순화다. 따라서 본 문서의 taxonomy는 연구 현실의 완전한 분할이 아니라, 학습과 탐색을 위한 지도 좌표로 이해해야 한다.",
    ]:
        add_p(doc, text)

    add_heading(doc, "11.4 How to Read Each Period Like a Survey Section", 2)
    period_guides = [
        ("2000-2008", "이 시기를 읽을 때는 최신 모델 성능이 아니라 연구 언어의 형성을 봐야 한다. 어떤 교재와 survey가 표준 용어를 만들었는가. 어떤 라이브러리가 실험 재현을 쉽게 했는가. 어떤 통계적 가정이 후대 딥러닝에도 남았는가. 이런 질문을 던지면 고전 ML 논문이 지금도 읽을 가치가 있는 이유가 보인다."),
        ("2009-2016", "이 시기는 benchmark, compute, architecture가 함께 움직인 시기다. ImageNet은 단순 데이터셋이 아니라 연구 경쟁의 규칙을 만들었다. CNN 계열 논문은 모델 구조를 바꿨고, GPU 학습과 공개 구현은 연구 속도를 바꿨다. 따라서 이 구간은 모델 혁신과 인프라 혁신을 함께 읽어야 한다."),
        ("2017-2021", "이 시기는 범용 architecture의 확장기로 읽어야 한다. Transformer가 언어에서 시각과 멀티모달로 이동하고, pretraining과 transfer가 연구 방법의 기본값이 된다. 동시에 AlphaFold처럼 특정 과학 문제에서 AI가 기존 pipeline을 크게 바꾸는 사례가 등장한다. 이 구간을 읽을 때는 분야 간 경계가 무너지는 지점을 표시해야 한다."),
        ("2022-2026", "이 시기는 아직 역사적 평가가 안정되지 않았다. citation map은 최신 논문의 완전한 순위를 제공하지 못한다. 대신 어떤 기술보고서가 빠르게 표준 참조가 되는지, 어떤 open model이 생태계를 형성하는지, 어떤 safety와 governance 문제가 반복되는지 보는 데 유용하다. 이 구간은 확정된 canon이 아니라 살아 있는 frontier로 읽어야 한다."),
    ]
    for label, text in period_guides:
        add_heading(doc, label, 3)
        add_p(doc, text)

    add_heading(doc, "11.5 Evidence Layers in the Repository", 2)
    for text in [
        "awesome-ai의 강점은 결과물만 있는 것이 아니라 evidence layer가 함께 있다는 점이다. candidates_top1000 파일은 선택되지 않은 후보까지 보여준다. papers_2000_2026 파일은 최종 선택 논문의 핵심 필드를 제공한다. papers_taxonomy 파일은 범주별 해석을 붙인다. period_analysis 파일은 interactive site가 기간별 요약을 보여주도록 만든다. github_links와 link_audit 파일은 링크의 공식성과 도달 가능성을 구분한다.",
        "이 구조는 연구자가 결과를 의심하고 다시 확인할 수 있게 만든다. 어떤 논문이 왜 들어갔는지, 어떤 연도 후보 풀에서 나왔는지, 어떤 query에 걸렸는지, DOI나 arXiv가 있는지, GitHub link가 공식인지 살펴볼 수 있다. 완벽한 검증은 아니지만, 닫힌 추천 목록보다 훨씬 낫다. 사용자는 이 evidence layer를 통해 자신의 연구 목적에 맞게 목록을 재정렬하거나 제외 기준을 더 엄격히 만들 수 있다.",
        "다만 evidence layer가 많다고 해서 자료가 자동으로 참이 되는 것은 아니다. Semantic Scholar metadata 자체가 틀릴 수 있고, PDF 링크가 다른 판본일 수 있으며, GitHub repository가 archive되었거나 실행되지 않을 수 있다. 그래서 본 문서는 link existence와 reproducibility를 분리해서 보라고 강조한다. 링크는 시작점이고, 재현성은 별도 실험과 환경 확인이 필요하다.",
    ]:
        add_p(doc, text)

    add_heading(doc, "11.6 Practical Review Templates", 2)
    add_p(
        doc,
        "사용자가 특정 논문을 awesome-ai에서 발견했을 때 바로 적용할 수 있는 검토 템플릿은 다음과 같다. 첫째, 이 논문이 어떤 category에 들어갔고 그 category의 일반 한계는 무엇인가. 둘째, 이 논문이 높은 인용을 받은 이유가 방법, 데이터셋, 도구, survey, benchmark 중 무엇인가. 셋째, open PDF와 official GitHub가 있는가. 넷째, 이 논문의 claim이 현재도 유효한가, 아니면 후속 연구가 수정했는가. 다섯째, 내 연구 질문에서 이 논문은 baseline인가, background인가, contrast인가, 아니면 cautionary example인가.",
    )
    add_p(
        doc,
        "세미나에서는 다른 템플릿이 필요하다. 한 주차에 단일 최고 인용 논문만 읽으면 역사적 흐름을 놓치기 쉽다. 대신 같은 범주의 대표 논문 하나, 같은 시기의 경쟁 논문 하나, 최근 후속 논문 하나, 한계나 비판을 다루는 논문 하나를 묶어 읽는 것이 좋다. 예를 들어 ResNet을 읽는다면 ImageNet과 BatchNorm, ViT 또는 ConvNeXt, robustness 논문을 함께 배치할 수 있다. Attention을 읽는다면 BERT, GPT-style scaling, retrieval augmentation, evaluation leakage 논의를 함께 다룰 수 있다.",
    )
    add_p(
        doc,
        "문헌조사 작성자는 awesome-ai를 문장 생성기가 아니라 구조 생성기로 사용해야 한다. category overview는 related work의 큰 문단 순서를 잡는 데 도움이 되고, limitations는 discussion section의 위험 요소를 잡는 데 도움이 된다. 그러나 개별 문장과 주장은 원논문에서 다시 확인해야 한다. 특히 자동 생성된 keyIdea는 편리한 색인이지만, 논문 저자의 실제 기여를 세밀하게 반영하지 못할 수 있다.",
    )

    add_heading(doc, "11.7 What a Stronger Future Version Could Add", 2)
    for text in [
        "다음 버전의 awesome-ai가 더 논문형 자료가 되려면 세 가지 보강이 유용하다. 첫째, expert annotation layer다. 분야별 전문가가 category와 representative status를 검토하면 deterministic metadata의 약점을 줄일 수 있다. 둘째, reproducibility layer다. official code가 실제 실행되는지, 환경 파일이 있는지, 데이터셋 접근이 가능한지, license가 명확한지 기록하면 citation map이 research operations map으로 확장된다.",
        "셋째, relation graph layer다. 현재 taxonomy는 논문을 주 범주에 배치하지만, 실제 영향은 그래프 형태다. Attention은 NLP에서 vision과 multimodal로 이동했고, contrastive learning은 representation과 foundation model 사이를 잇고, RLHF는 reinforcement learning과 LLM alignment를 연결한다. 이런 관계를 그래프로 표현하면 연구자는 단순 목록보다 훨씬 깊게 지형을 읽을 수 있다.",
        "넷째, negative evidence layer도 중요하다. 높은 인용 논문만 모으면 성공 사례 중심의 역사가 된다. 하지만 연구 발전에는 실패한 benchmark, 재현되지 않은 claim, 윤리적 논란, 폐기된 가정도 중요하다. 장기적으로는 대표 논문 옆에 비판 논문과 후속 correction을 함께 연결하는 방식이 필요하다. 이것이 citation map을 더 성숙한 scholarly map으로 만드는 방향이다.",
    ]:
        add_p(doc, text)

    add_heading(doc, "11.8 Extended Observations by Taxonomy Size", 2)
    for _, row in cat.iterrows():
        share = row.papers / len(df)
        add_p(
            doc,
            f"{row.category}는 선택 집합의 {share:.1%}를 차지한다. 이 비율은 해당 분야의 절대적 중요도라기보다, {int(df.year.min())}-{int(df.year.max())}년 동안 공개 메타데이터와 인용 관행 속에서 얼마나 자주 공통 참조점으로 등장했는지를 뜻한다. 따라서 범주 규모가 큰 분야는 교육과 기초 문헌조사에서 우선 읽기 좋지만, 범주 규모가 작은 분야도 연구적으로 덜 중요하다는 뜻은 아니다. 특히 새로운 연구 주제나 고위험 응용 분야는 인용 누적이 느리기 때문에 별도 탐색이 필요하다.",
        )

    add_heading(doc, "11.9 Category Reading Guides", 2)
    guide_texts = {
        "General AI Methods and Systems": "이 범주를 읽을 때는 개별 모델보다 연구의 공용 기반을 본다. 교재, 라이브러리, 정규화 기법, 최적화 방법, 데이터마이닝 절차가 왜 오래 인용되는지 질문해야 한다. 연구 초심자는 이 범주를 통해 AI 실험의 기본 어휘를 얻고, 숙련 연구자는 어떤 기반 가정이 여전히 현대 모델 평가에 남아 있는지 확인할 수 있다.",
        "Foundation Models and Large Language Models": "이 범주는 최신성이 매우 강하므로, 상위 인용만 보면 오히려 현재 frontier를 놓칠 수 있다. LLaMA, Qwen, DeepSeek, GPT 계열 기술보고서처럼 빠르게 확산되는 항목은 인용수보다 생태계 채택과 재사용을 함께 봐야 한다. 또한 모델 성능, 데이터 출처, 공개 범위, alignment 방식, evaluation leakage를 따로 분리해서 읽어야 한다.",
        "Vision and Multimodal Learning": "이 범주는 benchmark와 dataset의 힘을 가장 잘 보여준다. ImageNet, detection, segmentation, vision transformer, CLIP 계열을 읽을 때는 모델 구조뿐 아니라 어떤 데이터셋이 경쟁의 규칙을 만들었는지 봐야 한다. 실제 응용을 고려한다면 rare class, distribution shift, annotation bias, web-scale data bias를 별도 검토해야 한다.",
        "AI for Science, Healthcare, and Robotics": "이 범주는 논문 안의 점수와 현실 세계의 성공 사이가 가장 멀 수 있다. 단백질 구조, 의료 영상, 자율주행, 로봇 제어는 모두 외부 검증과 안전 조건이 중요하다. 따라서 이 범주에서는 retrospective benchmark보다 prospective validation, domain expert review, regulatory evidence, uncertainty estimation을 확인해야 한다.",
        "Graph Learning, Recommendation, and Core Methods": "이 범주를 읽을 때는 데이터 구조와 평가 split을 주의해야 한다. graph와 recommender system은 temporal leakage, popularity bias, negative sampling, hyperparameter search budget에 민감하다. 논문의 방법이 좋아 보일 때도 같은 데이터 분할과 같은 비교 비용에서 이득이 유지되는지 확인해야 한다.",
        "Reinforcement Learning and Agents": "이 범주는 성공 사례가 인상적이지만 실패 원인도 반복된다. reward misspecification, sample inefficiency, simulator bias, unsafe exploration은 고전 RL부터 agentic LLM까지 이어진다. 읽을 때는 환경이 얼마나 현실적인지, reward가 무엇을 실제로 최적화하는지, 실패 복구와 human oversight가 있는지 확인해야 한다.",
        "Natural Language Processing and Knowledge": "이 범주는 언어 이해와 지식 처리의 표준 과제를 만든다. Attention, BERT, T5 같은 논문은 모델 구조와 사전학습 패러다임을 이해하는 데 필수적이다. 다만 benchmark score는 annotation protocol, prompt wording, retrieval setting, 언어권 대표성에 민감하므로 수치만으로 일반 언어 능력을 판단하면 안 된다.",
        "Trustworthy, Explainable, and Responsible AI": "이 범주는 모든 분야를 가로지르는 검증 언어다. 설명가능성 논문을 읽을 때는 설명이 실제 모델 내부 원인과 맞는지, robustness 논문을 읽을 때는 공격 설정이 현실적인지, fairness 논문을 읽을 때는 어떤 집단과 metric이 기준인지 확인해야 한다. 책임 있는 AI는 단일 점수가 아니라 tradeoff 분석이다.",
        "Generative Models and Synthetic Media": "이 범주는 시각 품질과 사실성, 창의성과 통제 가능성, 공개 데이터와 저작권 사이의 긴장을 보여준다. GAN과 diffusion 논문을 읽을 때는 FID 같은 품질 지표뿐 아니라 prompt controllability, data provenance, misuse 가능성, downstream evaluation을 함께 봐야 한다. 생성 모델은 보여주기 쉬운 만큼 과장도 쉬운 분야다.",
        "Representation, Self-Supervised, and Transfer Learning": "이 범주는 규모는 작지만 거의 모든 현대 AI로 흘러 들어간다. word2vec, contrastive learning, transfer learning, adversarial examples는 모델이 데이터를 어떻게 표현하고 다른 과제로 옮기는지 설명한다. 읽을 때는 downstream task, data overlap, augmentation choice, negative sampling이 결과에 어떤 영향을 주는지 봐야 한다.",
    }
    for category, text in guide_texts.items():
        add_heading(doc, category, 3)
        add_p(doc, text)

    add_heading(doc, "11.10 Interpretive Q&A", 2)
    qa_items = [
        ("Q1. 이 목록의 1위 논문이 AI 역사상 가장 중요한 논문인가?", "아니다. 1위는 이 스냅샷과 선별 규칙 안에서 가장 높은 인용을 가진 논문이다. 중요도는 연구 질문에 따라 달라진다. 어떤 분야에서는 낮은 인용의 논문이 특정 문제 해결에 더 중요할 수 있다."),
        ("Q2. 최근 논문이 낮은 순위라면 덜 중요한가?", "아니다. 최근 논문은 인용이 쌓일 시간이 부족하다. 특히 2025년과 2026년 논문은 citation lag가 크기 때문에, adoption, official code, benchmark usage, expert review를 함께 봐야 한다."),
        ("Q3. GitHub 링크가 있으면 재현 가능한가?", "아니다. GitHub 링크는 구현 접근 가능성을 높이지만, 환경 파일, 데이터 접근, 라이선스, 학습 비용, checkpoint 공개 여부, 실행 성공 여부를 따로 확인해야 한다."),
        ("Q4. 분류가 하나로 붙은 논문은 그 분야에만 속하는가?", "아니다. taxonomy는 탐색을 위한 주 좌표다. Attention은 NLP에 놓여도 vision과 multimodal에 영향을 주고, CLIP은 vision 범주에 있어도 foundation model과 representation learning에 걸쳐 있다."),
        ("Q5. 후보 풀 27,000건은 충분한가?", "대규모 지도에는 충분히 유용하지만 완전하지 않다. 검색어, Semantic Scholar coverage, metadata quality, 언어권, 분야별 출판 관행에 따라 누락과 잡음이 생긴다."),
        ("Q6. 자동 생성된 key idea를 그대로 논문에 인용해도 되는가?", "권장하지 않는다. key idea는 빠른 색인이다. 학술 문장으로 쓰려면 원논문의 abstract, method, experiment, limitation을 직접 확인하고 표현을 새로 작성해야 한다."),
        ("Q7. 인용수와 influential citation count는 어떻게 다르게 읽어야 하는가?", "citationCount는 넓은 사용 흔적이고, influentialCitationCount는 더 강한 영향 신호를 의도한다. 둘 다 완벽하지 않으므로 category, venue, paper type과 함께 읽어야 한다."),
        ("Q8. unknown venue가 많으면 자료가 약한가?", "반드시 그렇지는 않다. preprint, technical report, metadata 누락이 섞일 수 있다. 하지만 venue metadata가 없는 논문은 신뢰도 판단을 위해 원 링크와 publication type을 더 꼼꼼히 봐야 한다."),
        ("Q9. 이 자료로 서베이 논문을 자동 완성할 수 있는가?", "초안 구조와 대표 문헌 목록은 만들 수 있지만, 자동 완성은 위험하다. 좋은 서베이는 논문 간 논쟁, 실패, 후속 수정, 방법의 실제 차이를 해석해야 하므로 사람의 읽기와 판단이 필요하다."),
        ("Q10. awesome-ai의 가장 좋은 쓰임은 무엇인가?", "출발점이다. 연구자가 방대한 AI 문헌에서 어디부터 읽을지 정하고, 분야별 대표 논문과 한계를 빠르게 파악하고, 원논문 검토와 세미나 설계를 체계화하는 데 가장 좋다."),
    ]
    for question, answer in qa_items:
        add_p(doc, question, style="List Bullet")
        add_p(doc, answer)


def add_references(doc: Document, df: pd.DataFrame) -> None:
    add_heading(doc, "References and Source Links", 1)
    refs = [
        ("honggi82/awesome-ai GitHub repository", AWESOME_GITHUB),
        ("Awesome AI interactive website", AWESOME_SITE),
        (f"Kong et al., {REFERENCE_TITLE}, arXiv:2605.18661", REFERENCE_ARXIV),
        ("worldbench/awesome-ai-auto-research GitHub repository", REFERENCE_GITHUB),
        ("Semantic Scholar Academic Graph API", "https://www.semanticscholar.org/product/api"),
        ("Papers with Code links-between-paper-and-code archive", "https://huggingface.co/datasets/pwc-archive/links-between-paper-and-code"),
    ]
    selected = df.sort_values("citationCount", ascending=False).head(35)
    for _, row in selected.iterrows():
        authors = clean(row.authors)
        if len(authors) > 110:
            authors = authors[:107] + "..."
        venue = clean(row.venue) or "Unknown venue"
        link = clean(row.url) or clean(row.semanticScholarUrl) or (f"https://doi.org/{clean(row.doi)}" if clean(row.doi) else "")
        refs.append((f"{authors}. {row.title}. {venue}, {int(row.year)}. Snapshot citations: {fmt_int(row.citationCount)}.", link))
    for index, (title, link) in enumerate(refs, 1):
        add_p(doc, f"[{index}] {title} {link}")


def main() -> None:
    df = pd.read_csv(DATA_DIR / "papers_2000_2026.csv")
    tax = pd.read_csv(DATA_DIR / "papers_taxonomy_2000_2026.csv")
    link_audit = json.loads((DATA_DIR / "link_audit_2000_2026.json").read_text(encoding="utf-8"))
    github_links = json.loads((DATA_DIR / "github_links_2000_2026.json").read_text(encoding="utf-8"))
    cat = (
        df.groupby("category")
        .agg(
            papers=("title", "count"),
            citations=("citationCount", "sum"),
            influential=("influentialCitationCount", "sum"),
            median_year=("year", "median"),
        )
        .reset_index()
        .sort_values(["papers", "citations"], ascending=False)
    )
    stats = {
        "years": f"{int(df.year.min())}-{int(df.year.max())}",
        "selected": len(df),
        "candidates": 27000,
        "citations": int(df.citationCount.fillna(0).sum()),
        "categories": df.category.nunique(),
    }
    periods = [
        ("2000-2008", 2000, 2008, "통계적 학습, 데이터마이닝, 커널 방법, 고전 ML 라이브러리가 연구 기반을 형성한 시기"),
        ("2009-2016", 2009, 2016, "ImageNet, CNN, 표현학습, 대규모 학습 인프라, 강화학습이 딥러닝 전환을 만든 시기"),
        ("2017-2021", 2017, 2021, "Transformer, BERT, ViT, CLIP, AlphaFold 등 파운데이션화와 멀티모달 전환이 본격화된 시기"),
        ("2022-2026", 2022, 2026, "instruction tuning, LLM, 안전성, 해석가능성, AI4Science가 연구 생태계를 재편하는 시기"),
    ]
    figures = make_figures(df, cat, periods)
    doc = configure_doc()
    add_title_page(doc, stats)
    add_reference_structure_analysis(doc)
    add_abstract(doc, stats, cat)
    add_preliminaries(doc, stats, figures)
    add_methods(doc, df, stats, figures)
    add_historical_roadmap(doc, df, periods, figures)
    add_taxonomy(doc, df, tax, cat, figures)
    add_validation_and_crosscutting(doc, df, stats, link_audit, github_links)
    add_user_guide_and_appendix(doc, df, cat)
    add_deep_discussion(doc, df, cat)
    add_references(doc, df)

    props = doc.core_properties
    props.title = "Awesome AI: Roadmap & User Guide"
    props.subject = "Detailed survey-style Word paper based on honggi82/awesome-ai"
    props.author = "Codex, based on honggi82/awesome-ai public metadata"
    props.keywords = "AI survey, awesome-ai, citation map, taxonomy, roadmap, user guide"
    props.comments = "Generated from local awesome-ai data snapshot dated 2026-06-26."
    doc.save(OUT_PATH)

    check = Document(OUT_PATH)
    text = "\n".join(paragraph.text for paragraph in check.paragraphs)
    hangul_count = sum(0xAC00 <= ord(char) <= 0xD7A3 for char in text)
    question_count = text.count("?")
    if hangul_count < 12000 or len(text) < 50000:
        raise RuntimeError(f"Document content is still too thin: chars={len(text)}, hangul_count={hangul_count}")
    if len(check.tables) < 9:
        raise RuntimeError(f"Expected at least 9 tables, found {len(check.tables)}")
    if question_count > 120:
        raise RuntimeError(f"Generated document has suspicious replacement marks: {question_count}")
    print(OUT_PATH.resolve())
    print(f"docx_size={OUT_PATH.stat().st_size}")
    print(f"hangul_count={hangul_count}")
    print(f"question_count={question_count}")
    print(f"paragraphs={len(check.paragraphs)}")
    print(f"tables={len(check.tables)}")
    print(f"figures={len(figures)}")


if __name__ == "__main__":
    main()
