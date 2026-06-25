import csv
import hashlib
import html
import json
import math
import re
import shutil
import sys
import time
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

import requests
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
DOCS_DIR = ROOT / "docs"
PAPER_DIR = ROOT / "paper"
CACHE_DIR = DATA_DIR / "cache"

START_YEAR = 2020
END_YEAR = 2026
YEARS = list(range(START_YEAR, END_YEAR + 1))
YEAR_RANGE_TEXT = f"{START_YEAR}-{END_YEAR}"
YEAR_FILE_STEM = f"{START_YEAR}_{END_YEAR}"

PAPERS_JSON = f"papers_{YEAR_FILE_STEM}.json"
PAPERS_CSV = f"papers_{YEAR_FILE_STEM}.csv"
CANDIDATES_JSON = f"candidates_top1000_{YEAR_FILE_STEM}.json"
CANDIDATES_CSV = f"candidates_top1000_{YEAR_FILE_STEM}.csv"
TAXONOMY_CSV = f"papers_taxonomy_{YEAR_FILE_STEM}.csv"
PERIOD_ANALYSIS_JSON = f"period_analysis_{YEAR_FILE_STEM}.json"

CANDIDATES_PER_YEAR = 1000
TARGET_TOTAL = 100
S2_BULK_URL = "https://api.semanticscholar.org/graph/v1/paper/search/bulk"
REQUEST_DELAY = 0.05

S2_FIELDS = ",".join(
    [
        "paperId",
        "title",
        "year",
        "authors",
        "venue",
        "publicationVenue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "abstract",
        "url",
        "externalIds",
        "openAccessPdf",
        "s2FieldsOfStudy",
        "publicationTypes",
    ]
)

QUERIES = [
    "artificial intelligence",
    "machine learning",
    "deep learning",
    "foundation models",
    "large language models",
    "natural language processing",
    "computer vision",
    "reinforcement learning",
    "generative AI diffusion models",
    "graph neural networks",
    "multimodal learning",
    "AI safety fairness explainability robustness",
    "AI for science healthcare robotics",
]

AI_RELEVANCE_PATTERNS = [
    r"\bartificial intelligence\b",
    r"\bAI\b",
    r"\bmachine learning\b",
    r"\bdeep learning\b",
    r"\bneural network",
    r"\btransformer",
    r"\blanguage model",
    r"\blarge language model",
    r"\bLLM\b",
    r"\bfoundation model",
    r"\bcomputer vision\b",
    r"\bimage recognition\b",
    r"\bimage classification\b",
    r"\bobject detection\b",
    r"\bsemantic segmentation\b",
    r"\bnatural language processing\b",
    r"\bNLP\b",
    r"\breinforcement learning\b",
    r"\bpolicy gradient\b",
    r"\bgenerative model",
    r"\bdiffusion model",
    r"\bgenerative adversarial",
    r"\bGAN\b",
    r"\bgraph neural",
    r"\bself-supervised\b",
    r"\bcontrastive learning\b",
    r"\brepresentation learning\b",
    r"\bfederated learning\b",
    r"\bexplainable AI\b",
    r"\bXAI\b",
    r"\badversarial\b",
    r"\brobustness\b",
    r"\bfairness\b",
    r"\bprivacy\b",
    r"\bspeech recognition\b",
    r"\brecommendation",
    r"\bfew-shot\b",
    r"\bzero-shot\b",
    r"\bprompt",
    r"\bRAG\b",
    r"\bretrieval-augmented\b",
]

BAD_TITLE_PATTERNS = [
    r"^group$",
    r"^editorial\b",
    r"^preface\b",
    r"^front matter\b",
    r"^proceedings of\b",
    r"^international journal of\b",
    r"^advances in neural information processing systems \d+$",
    r"^the sciences of the artificial$",
    r"^statistical learning theory$",
    r"^generative adversarial networks$",
    r"^foundations of machine learning$",
    r"^language models$",
    r"^random forests$",
    r"^deep learning$",
    r"^machine learning$",
    r"^pattern recognition and machine learning$",
    r"^artificial intelligence a modern approach$",
]

IMPORTANT_VENUES = [
    "nature",
    "science",
    "cell",
    "neurips",
    "nips",
    "icml",
    "iclr",
    "aaai",
    "ijcai",
    "cvpr",
    "iccv",
    "eccv",
    "acl",
    "emnlp",
    "naacl",
    "kdd",
    "sigir",
    "www",
    "the web conference",
    "ieee transactions",
    "jmlr",
    "transactions on machine learning research",
    "pattern recognition",
    "artificial intelligence",
    "machine learning",
    "nature machine intelligence",
    "nature methods",
    "nature biomedical engineering",
]

CATEGORIES = [
    (
        "Foundation Models and Large Language Models",
        [
            "large language model",
            "language model",
            "foundation model",
            "gpt",
            "bert",
            "pretrain",
            "prompt",
            "retrieval-augmented",
            "rag",
            "scaling law",
            "instruction",
        ],
    ),
    (
        "Generative Models and Synthetic Media",
        [
            "diffusion",
            "generative",
            "gan",
            "adversarial network",
            "image synthesis",
            "text-to-image",
            "vae",
            "score-based",
            "synthetic",
        ],
    ),
    (
        "Vision and Multimodal Learning",
        [
            "computer vision",
            "image",
            "video",
            "vision transformer",
            "object detection",
            "segmentation",
            "clip",
            "multimodal",
            "vision-language",
            "visual",
            "point cloud",
        ],
    ),
    (
        "Natural Language Processing and Knowledge",
        [
            "natural language",
            "nlp",
            "translation",
            "question answering",
            "dialogue",
            "summarization",
            "information retrieval",
            "knowledge graph",
            "text",
            "speech recognition",
        ],
    ),
    (
        "Reinforcement Learning and Agents",
        [
            "reinforcement",
            "policy",
            "agent",
            "markov",
            "robot",
            "control",
            "planning",
            "imitation",
            "reward",
            "human feedback",
        ],
    ),
    (
        "Representation, Self-Supervised, and Transfer Learning",
        [
            "self-supervised",
            "contrastive",
            "representation",
            "domain adaptation",
            "few-shot",
            "zero-shot",
            "transfer learning",
            "metric learning",
            "distillation",
            "pre-training",
        ],
    ),
    (
        "Trustworthy, Explainable, and Responsible AI",
        [
            "explainable",
            "interpretability",
            "interpretable",
            "fairness",
            "bias",
            "robust",
            "adversarial",
            "privacy",
            "safety",
            "uncertainty",
            "calibration",
        ],
    ),
    (
        "Graph Learning, Recommendation, and Core Methods",
        [
            "graph neural",
            "recommendation",
            "recommender",
            "optimization",
            "hyperparameter",
            "neural architecture",
            "bayesian",
            "random forest",
            "ensemble",
            "kernel",
            "decision tree",
        ],
    ),
    (
        "AI for Science, Healthcare, and Robotics",
        [
            "protein",
            "molecule",
            "drug",
            "genomic",
            "healthcare",
            "medical",
            "clinical",
            "biomedical",
            "science",
            "robotics",
            "autonomous driving",
            "biology",
        ],
    ),
]

KEYWORD_CONVENTION = [
    (
        "foundation-models",
        "Large language models, foundation models, scaling, prompting, alignment, or retrieval-augmented systems.",
        "2563eb",
    ),
    (
        "generative-ai",
        "Generative adversarial, diffusion, synthetic media, text-to-image, or other model-based generation work.",
        "a855f7",
    ),
    (
        "multimodal",
        "Vision-language, audio-language, video-language, or cross-modal representation learning.",
        "0891b2",
    ),
    (
        "nlp",
        "Natural language processing, language modeling, retrieval, dialogue, summarization, or speech-language work.",
        "f59e0b",
    ),
    (
        "vision",
        "Computer vision, image/video understanding, object detection, segmentation, or visual recognition.",
        "0f766e",
    ),
    (
        "reinforcement-learning",
        "Reinforcement learning, agents, planning, control, robotics, reward modeling, or human feedback.",
        "dc2626",
    ),
    (
        "trustworthy-ai",
        "Explainability, robustness, safety, fairness, uncertainty, privacy, bias, or responsible AI.",
        "be123c",
    ),
    (
        "graph-learning",
        "Graph neural networks, recommender systems, knowledge graphs, graph benchmarks, or graph-based core AI methods.",
        "4f46e5",
    ),
    (
        "ai4science",
        "AI for science, healthcare, biology, molecules, proteins, robotics, autonomous systems, or clinical domains.",
        "16a34a",
    ),
]
KEYWORD_COLORS = {keyword: color for keyword, _, color in KEYWORD_CONVENTION}

LANGUAGES = {
    "en": "English",
    "ko": "한국어",
}

UI_LABELS = {
    "en": {
        "papers": "papers",
        "categories": "categories",
        "overview": "Category Overview",
        "limitations": "Limitations",
        "analysis": "Selected-period analysis",
        "totalSelected": "Total selected papers",
        "categoryCount": "Categories",
        "keyIdea": "Key idea",
        "strengths": "Strengths",
        "paperLimitations": "Limitations",
    },
    "ko": {
        "papers": "편",
        "categories": "개 분류",
        "overview": "분류 개요",
        "limitations": "한계",
        "analysis": "선택 기간 분석",
        "totalSelected": "선정 논문",
        "categoryCount": "분류",
        "keyIdea": "핵심 아이디어",
        "strengths": "장점",
        "paperLimitations": "한계",
    },
}

TAXONOMY_TRENDS = {
    "Foundation Models and Large Language Models": [
        "The main trend is a shift from task-specific NLP systems toward general-purpose foundation models that transfer across tasks through prompting, retrieval, and instruction tuning.",
        "Scaling laws, retrieval augmentation, alignment, and data governance are now central design axes rather than afterthoughts.",
        "Citation-ranked work tends to emphasize architectures, pretraining corpora, benchmark behavior, and broad capability evaluation.",
    ],
    "Generative Models and Synthetic Media": [
        "Generative AI research has moved from GAN-centered image synthesis toward diffusion, score-based modeling, and controllable multimodal generation.",
        "The area increasingly connects generation quality with data curation, safety, copyright, controllability, and evaluation reliability.",
        "Highly cited papers often introduce reusable model families, training objectives, or evaluation protocols that become infrastructure for later systems.",
    ],
    "Vision and Multimodal Learning": [
        "Vision research is increasingly organized around transformer backbones, self-supervised pretraining, segmentation/detection foundation models, and vision-language alignment.",
        "The strongest papers often combine large-scale data, reusable architectures, and transfer to multiple downstream tasks.",
        "Multimodal work is pushing vision beyond single-task recognition toward retrieval, grounding, robotics, and generative interfaces.",
    ],
    "Natural Language Processing and Knowledge": [
        "NLP is moving from supervised task pipelines toward pretrained language models, retrieval-augmented methods, and knowledge-intensive reasoning benchmarks.",
        "Search, retrieval, summarization, dialogue, and domain adaptation are increasingly evaluated as integrated knowledge workflows.",
        "Citation-ranked NLP papers often become shared components, datasets, or baseline methods for later foundation model research.",
    ],
    "Reinforcement Learning and Agents": [
        "Reinforcement learning is converging with human feedback, offline datasets, robotics, planning, and agentic tool-use settings.",
        "A major thread is reducing sample inefficiency while improving robustness under distribution shift and sparse rewards.",
        "The field is increasingly judged by transfer, safety, and real-world interaction rather than benchmark score alone.",
    ],
    "Representation, Self-Supervised, and Transfer Learning": [
        "Self-supervised and contrastive methods are reducing dependence on labeled data while improving transfer across tasks and domains.",
        "Reusable representations, distillation, domain adaptation, and few-shot learning form the connective tissue between specialized AI subfields.",
        "Citation impact is often driven by methods that become default pretraining or adaptation recipes.",
    ],
    "Trustworthy, Explainable, and Responsible AI": [
        "Trustworthy AI work is broadening from post-hoc explanations to robustness, fairness, privacy, calibration, uncertainty, and safety-aware evaluation.",
        "Highly cited papers provide taxonomies, metrics, attacks, or toolkits that make system behavior easier to inspect and compare.",
        "The area is increasingly coupled to foundation model deployment, high-stakes domains, and governance requirements.",
    ],
    "Graph Learning, Recommendation, and Core Methods": [
        "Core AI methods include graph neural networks, recommender systems, optimization, neural architecture search, Bayesian methods, and efficient training recipes.",
        "Many papers in this category become reusable algorithmic infrastructure for applied AI systems.",
        "The dominant trend is stronger inductive bias for non-Euclidean data, sparse interaction data, and efficient hyperparameter or architecture search.",
    ],
    "AI for Science, Healthcare, and Robotics": [
        "Applied AI for science and healthcare is shifting from proof-of-concept prediction toward validated workflows for biology, medicine, molecules, and robotics.",
        "The most visible work couples domain data with deep learning architectures that can transfer into laboratory, clinical, or embodied settings.",
        "Evaluation increasingly needs external validation, prospective testing, reproducible datasets, and domain-specific safety constraints.",
    ],
    "General AI Methods and Systems": [
        "General AI methods consolidate architectures, benchmarks, surveys, datasets, and system-level observations that cut across subfields.",
        "This category often captures high-citation survey or infrastructure work that shapes how later papers define progress.",
        "Citation-ranked views can be especially useful here, but they should be read as a map of influence rather than a complete quality assessment.",
    ],
}

TAXONOMY_LIMITATIONS = {
    "Foundation Models and Large Language Models": [
        "Capability gains are difficult to separate from data scale, benchmark leakage, and evaluation prompt sensitivity.",
        "Alignment and safety claims often need stronger real-world and multilingual validation.",
        "Compute-intensive training can limit reproducibility and concentrate follow-up work around a small number of institutions.",
    ],
    "Generative Models and Synthetic Media": [
        "Image and media quality metrics may not capture factuality, controllability, provenance, or downstream harms.",
        "Training data provenance and copyright constraints can be under-specified in highly cited generation work.",
        "Robust evaluation across cultures, modalities, and adversarial uses remains difficult.",
    ],
    "Vision and Multimodal Learning": [
        "Large-scale benchmark success can overstate robustness under distribution shift, rare classes, and real deployment constraints.",
        "Multimodal alignment may inherit biases and spurious correlations from web-scale data.",
        "High-performing systems often require data and compute resources that are hard for smaller labs to reproduce.",
    ],
    "Natural Language Processing and Knowledge": [
        "Benchmark scores can hide brittle reasoning, retrieval failures, hallucination, and domain transfer issues.",
        "Language coverage is often uneven, with English and high-resource domains overrepresented.",
        "Evaluation can be sensitive to annotation protocols, prompt wording, and changing model APIs.",
    ],
    "Reinforcement Learning and Agents": [
        "Sample efficiency, reward misspecification, simulator bias, and safety under exploration remain persistent barriers.",
        "Benchmark performance may not transfer to physical robots, human-facing tools, or open-ended environments.",
        "Agentic systems need stronger evidence on reliability, recovery from errors, and long-horizon oversight.",
    ],
    "Representation, Self-Supervised, and Transfer Learning": [
        "Transfer claims depend heavily on downstream task choice, data overlap, and evaluation protocol.",
        "Contrastive and self-supervised methods can learn spurious shortcuts when augmentations or negatives are poorly matched.",
        "Representation quality is hard to compare when model size, data scale, and training recipes differ.",
    ],
    "Trustworthy, Explainable, and Responsible AI": [
        "Explanations can be persuasive without being faithful to model internals or decision processes.",
        "Fairness, robustness, privacy, and safety metrics can conflict and require domain-specific tradeoffs.",
        "Responsible AI results often need stronger deployment evidence beyond benchmark or synthetic settings.",
    ],
    "Graph Learning, Recommendation, and Core Methods": [
        "Graph and recommender benchmarks can contain temporal leakage, popularity bias, or unrealistic train/test splits.",
        "Algorithmic gains may be sensitive to hyperparameter budgets and implementation details.",
        "Core methods need careful ablations before broad claims about generality or efficiency are accepted.",
    ],
    "AI for Science, Healthcare, and Robotics": [
        "External validation, prospective testing, and domain expert review are often more important than retrospective benchmark scores.",
        "Clinical, biological, or robotic deployment can fail when data collection protocols differ from training assumptions.",
        "Safety, interpretability, uncertainty, and regulatory evidence remain essential for translation.",
    ],
    "General AI Methods and Systems": [
        "Survey and infrastructure papers can dominate citations while empirical evidence remains distributed across subfields.",
        "Broad claims need careful mapping to specific tasks, datasets, and operational constraints.",
        "Metadata-driven ranking cannot replace expert reading of full papers and experimental details.",
    ],
}

KOREAN_CATEGORY_NAMES = {
    "Foundation Models and Large Language Models": "파운데이션 모델 및 대규모 언어모델",
    "Generative Models and Synthetic Media": "생성 모델 및 합성 미디어",
    "Vision and Multimodal Learning": "비전 및 멀티모달 학습",
    "Natural Language Processing and Knowledge": "자연어처리 및 지식",
    "Reinforcement Learning and Agents": "강화학습 및 에이전트",
    "Representation, Self-Supervised, and Transfer Learning": "표현학습, 자기지도학습 및 전이학습",
    "Trustworthy, Explainable, and Responsible AI": "신뢰가능하고 설명가능한 책임 AI",
    "Graph Learning, Recommendation, and Core Methods": "그래프 학습, 추천 및 핵심 방법론",
    "AI for Science, Healthcare, and Robotics": "과학, 헬스케어 및 로보틱스를 위한 AI",
    "General AI Methods and Systems": "일반 AI 방법론 및 시스템",
}


def norm_text(value):
    return re.sub(r"\s+", " ", value or "").strip()


def safe_slug(value):
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return value[:90] or "paper"


def normalize_title(value):
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def paper_key(paper):
    ext = paper.get("externalIds") or {}
    for key in ("DOI", "ArXiv", "PubMed", "CorpusId"):
        value = ext.get(key)
        if value:
            return f"{key.lower()}:{str(value).lower()}"
    if paper.get("paperId"):
        return f"s2:{paper['paperId']}"
    return f"title:{normalize_title(paper.get('title'))}"


def cache_name(year, query):
    digest = hashlib.sha1(f"{year}:{query}".encode("utf-8")).hexdigest()[:12]
    return CACHE_DIR / f"s2_{year}_{safe_slug(query)}_{digest}.json"


def fetch_year_query(year, query):
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_file = cache_name(year, query)
    if cache_file.exists():
        return json.loads(cache_file.read_text(encoding="utf-8"))

    params = {
        "query": query,
        "year": str(year),
        "fields": S2_FIELDS,
        "sort": "citationCount:desc",
        "limit": str(CANDIDATES_PER_YEAR),
    }
    last_error = None
    for attempt in range(4):
        try:
            response = requests.get(S2_BULK_URL, params=params, timeout=90)
            if response.status_code == 429:
                time.sleep(15 + attempt * 10)
                continue
            response.raise_for_status()
            payload = response.json()
            rows = payload.get("data") or []
            cache_file.write_text(json.dumps(rows, ensure_ascii=False), encoding="utf-8")
            return rows
        except Exception as exc:
            last_error = exc
            time.sleep(5 + attempt * 5)
    raise RuntimeError(f"Semantic Scholar query failed for {year} {query!r}: {last_error}")


def venue_name(paper):
    publication_venue = paper.get("publicationVenue") or {}
    return norm_text(publication_venue.get("name") or paper.get("venue") or "")


def authors_text(paper, max_authors=8):
    authors = paper.get("authors") or []
    names = [norm_text(a.get("name")) for a in authors if norm_text(a.get("name"))]
    if len(names) > max_authors:
        return ", ".join(names[:max_authors]) + ", et al."
    return ", ".join(names)


def primary_link(paper):
    ext = paper.get("externalIds") or {}
    doi = ext.get("DOI")
    if doi:
        doi = str(doi)
        return doi if doi.startswith("http") else f"https://doi.org/{doi}"
    if paper.get("url"):
        return paper["url"]
    pdf = paper.get("openAccessPdf") or {}
    return pdf.get("url") or ""


def title_is_bad(title):
    title_l = normalize_title(title)
    if len(title_l) < 8:
        return True
    return any(re.search(pattern, title_l, re.I) for pattern in BAD_TITLE_PATTERNS)


def relevance_counts(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract} {venue}"
    title_hits = sum(1 for pattern in AI_RELEVANCE_PATTERNS if re.search(pattern, title, re.I))
    text_hits = sum(1 for pattern in AI_RELEVANCE_PATTERNS if re.search(pattern, text, re.I))
    return title_hits, text_hits


def is_relevant(paper):
    title = norm_text(paper.get("title"))
    if not title or title_is_bad(title):
        return False
    publication_types = [str(x).lower() for x in (paper.get("publicationTypes") or [])]
    if any(kind in publication_types for kind in ("editorial", "news", "lettersandcomments")):
        return False
    title_hits, text_hits = relevance_counts(paper)
    fields = " ".join(
        norm_text(item.get("category"))
        for item in (paper.get("s2FieldsOfStudy") or [])
        if item.get("category")
    ).lower()
    if title_hits:
        return True
    if text_hits >= 2:
        return True
    return text_hits >= 1 and "computer science" in fields and "review" in title.lower()


def importance_score(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract} {venue}".lower()
    citations = int(paper.get("citationCount") or 0)
    influential = int(paper.get("influentialCitationCount") or 0)
    title_hits, text_hits = relevance_counts(paper)
    query_hits = len(paper.get("sourceQueries") or [])
    score = math.log1p(citations) * 22.0 + math.log1p(influential) * 14.0
    reasons = [f"citations={citations}", f"influential={influential}"]
    if any(v in venue.lower() for v in IMPORTANT_VENUES):
        score += 10
        reasons.append("recognized venue")
    if re.search(r"\b(review|survey|systematic review|benchmark|dataset)\b", text):
        score += 5
        reasons.append("survey/benchmark signal")
    if re.search(r"\b(open-source|code|dataset|benchmark|corpus)\b", text):
        score += 3
        reasons.append("resource signal")
    if re.search(r"\b(safety|fairness|privacy|robust|explainable|interpretability|clinical|healthcare)\b", text):
        score += 3
        reasons.append("translation/trustworthiness signal")
    if paper.get("openAccessPdf"):
        score += 1
        reasons.append("open PDF metadata")
    if title_hits:
        score += min(8, title_hits * 2)
        reasons.append(f"title AI matches={title_hits}")
    if text_hits:
        score += min(8, text_hits)
        reasons.append(f"AI term matches={text_hits}")
    if query_hits > 1:
        score += min(6, query_hits)
        reasons.append(f"query hits={query_hits}")
    return round(score, 3), "; ".join(reasons)


def text_for(paper):
    return f"{paper.get('title', '')} {paper.get('abstract', '')} {paper.get('venue', '')} {paper.get('methodTags', '')}".lower()


def category_for(paper):
    text = text_for(paper)
    if re.search(r"\b(large language models?|language models?|llm|foundation model|gpt|chatgpt|prompting?|instruction[- ]?tuning|scaling laws?|retrieval-augmented generation|rag)\b", text, re.I):
        return "Foundation Models and Large Language Models"
    if re.search(r"\b(protein|molecule|drug|genomic|healthcare|medical|clinical|biomedical|biology|autonomous driving)\b", text, re.I):
        return "AI for Science, Healthcare, and Robotics"
    if re.search(r"\b(graph neural|graph convolution|open graph benchmark|knowledge graph|recommender|recommendation)\b", text, re.I):
        return "Graph Learning, Recommendation, and Core Methods"
    best = ("General AI Methods and Systems", 0)
    for category, terms in CATEGORIES:
        score = sum(1 for term in terms if term in text)
        if score > best[1]:
            best = (category, score)
    return best[0]


def method_tags(paper):
    text = text_for(paper)
    rules = [
        ("survey/review", r"\b(review|survey|overview)\b"),
        ("benchmark/dataset", r"\b(benchmark|dataset|corpus|leaderboard|open graph)\b"),
        ("transformer", r"\b(transformer|attention|bert|gpt|vit)\b"),
        ("self-supervised", r"\b(self-supervised|contrastive|masked|pretrain|pre-training)\b"),
        ("generative", r"\b(diffusion|generative|gan|vae|score-based|text-to-image)\b"),
        ("graph-learning", r"\b(graph neural|gnn|knowledge graph)\b"),
        ("retrieval", r"\b(retrieval|rag|dense passage|information retrieval)\b"),
        ("trustworthy-ai", r"\b(explainable|interpretability|fairness|robust|privacy|safety|uncertainty)\b"),
        ("ai4science", r"\b(protein|molecule|drug|clinical|medical|healthcare|robot|biology)\b"),
        ("reinforcement-learning", r"\b(reinforcement|policy|agent|reward|human feedback)\b"),
    ]
    tags = [name for name, pattern in rules if re.search(pattern, text, re.I)]
    return tags[:6] or ["metadata-ranked"]


def keyword_tags(paper, category=None):
    text = text_for(paper)
    tags = set()
    if re.search(r"\b(language model|large language|llm|foundation model|gpt|bert|prompt|rag|retrieval-augmented|scaling law)\b", text, re.I):
        tags.add("foundation-models")
    if re.search(r"\b(diffusion|generative|gan|vae|synthetic|text-to-image|image synthesis)\b", text, re.I):
        tags.add("generative-ai")
    if re.search(r"\b(multimodal|multi-modal|vision-language|clip|cross-modal|video-language)\b", text, re.I):
        tags.add("multimodal")
    if re.search(r"\b(natural language|nlp|text|translation|question answering|dialogue|summarization|speech recognition|information retrieval|language model)\b", text, re.I):
        tags.add("nlp")
    if re.search(r"\b(computer vision|image|video|object detection|segmentation|visual|vision transformer|point cloud)\b", text, re.I):
        tags.add("vision")
    if re.search(r"\b(reinforcement|policy|agent|robot|control|planning|reward|human feedback|imitation)\b", text, re.I):
        tags.add("reinforcement-learning")
    if re.search(r"\b(explainable|interpretability|fairness|bias|robust|adversarial|privacy|safety|uncertainty|calibration|responsible)\b", text, re.I):
        tags.add("trustworthy-ai")
    if re.search(r"\b(graph neural|graph convolution|open graph benchmark|knowledge graph|recommender|recommendation|graph machine learning)\b", text, re.I):
        tags.add("graph-learning")
    if re.search(r"\b(protein|molecule|drug|genomic|healthcare|medical|clinical|biomedical|science|robotics|autonomous driving|biology)\b", text, re.I):
        tags.add("ai4science")

    category = category or category_for(paper)
    defaults = {
        "Foundation Models and Large Language Models": "foundation-models",
        "Generative Models and Synthetic Media": "generative-ai",
        "Vision and Multimodal Learning": "vision",
        "Natural Language Processing and Knowledge": "nlp",
        "Reinforcement Learning and Agents": "reinforcement-learning",
        "Trustworthy, Explainable, and Responsible AI": "trustworthy-ai",
        "Graph Learning, Recommendation, and Core Methods": "graph-learning",
        "AI for Science, Healthcare, and Robotics": "ai4science",
    }
    if not tags and category in defaults:
        tags.add(defaults[category])
    ordered = [keyword for keyword, _, _ in KEYWORD_CONVENTION if keyword in tags]
    return ordered or ["foundation-models"]


def first_sentence(value, fallback):
    text = norm_text(value)
    if not text:
        return fallback
    match = re.search(r"(.{40,280}?[.!?])\s", text + " ")
    return match.group(1) if match else text[:280]


def research_limitations(paper, method_tag_list, keyword_list):
    category = paper.get("category") or category_for(paper)
    items = list(TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General AI Methods and Systems"]))
    tags = set(method_tag_list + keyword_list)
    if "survey/review" in tags:
        items.append("Review-level synthesis cannot resolve inconsistent study quality, benchmark leakage, or reproducibility gaps.")
    if "benchmark/dataset" in tags:
        items.append("Benchmark value depends on sustained maintenance, clear licensing, and representative task design.")
    if "trustworthy-ai" in tags:
        items.append("Trustworthiness metrics require domain-specific thresholds and may trade off against accuracy or utility.")
    if "ai4science" in tags:
        items.append("Domain translation needs external validation under laboratory, clinical, or embodied deployment conditions.")
    if "foundation-models" in tags:
        items.append("Large-model conclusions can be sensitive to data scale, prompt design, and hidden training-set overlap.")
    return items[:3]


def localized_paper_fields(paper, language):
    if language == "en":
        return {
            "keyIdea": paper["keyIdea"],
            "strengths": paper["strengths"],
            "limitations": paper["limitations"],
        }
    category_name = KOREAN_CATEGORY_NAMES.get(paper["category"], paper["category"])
    tags = paper.get("methodTags", "").replace("; ", ", ") or "AI 방법론"
    return {
        "keyIdea": f"'{paper['title']}'은(는) {category_name} 분야에서 {tags} 흐름을 보여주는 citation-ranked AI 연구입니다.",
        "strengths": paper["strengths"]
        .replace("high citation signal", "높은 인용 신호")
        .replace("influential citation signal", "영향력 있는 인용 신호")
        .replace("recognized venue", "주요 학술지/학회 신호")
        .replace("open-access PDF metadata", "오픈액세스 PDF 메타데이터")
        .replace("selected by citation count from the audited AI candidate pool", "감사 가능한 AI 후보군에서 인용수 기준으로 선정"),
        "limitations": " ".join(TAXONOMY_LIMITATIONS.get(paper["category"], TAXONOMY_LIMITATIONS["General AI Methods and Systems"])[:2]),
    }


def enrich_paper(paper):
    base = dict(paper)
    category = category_for(base)
    methods = method_tags(base)
    keywords = keyword_tags(base, category)
    citations = int(base.get("citationCount") or 0)
    influential = int(base.get("influentialCitationCount") or 0)
    strengths = []
    if citations >= 100:
        strengths.append(f"high citation signal ({citations:,})")
    if influential >= 10:
        strengths.append(f"influential citation signal ({influential:,})")
    if "recognized venue" in base.get("importanceReasons", ""):
        strengths.append("recognized venue")
    if base.get("openAccessPdf"):
        strengths.append("open-access PDF metadata")
    if not strengths:
        strengths.append("selected by citation count from the audited AI candidate pool")
    base["category"] = category
    base["methodTags"] = "; ".join(methods)
    base["keywordTags"] = "; ".join(keywords)
    base["keyIdea"] = first_sentence(
        base.get("abstract", ""),
        f"Positions {base.get('title', 'this paper')} within {category}.",
    )
    base["strengths"] = "; ".join(strengths[:4])
    base["limitations"] = "; ".join(research_limitations(base, methods, keywords))
    return base


def normalize_paper(paper, year, candidate_rank):
    ext = paper.get("externalIds") or {}
    score, reasons = importance_score(paper)
    pdf = paper.get("openAccessPdf") or {}
    row = {
        "paperId": paper.get("paperId") or "",
        "title": norm_text(paper.get("title")),
        "authors": authors_text(paper),
        "year": int(paper.get("year") or year),
        "venue": venue_name(paper),
        "publicationDate": paper.get("publicationDate") or "",
        "citationCount": int(paper.get("citationCount") or 0),
        "influentialCitationCount": int(paper.get("influentialCitationCount") or 0),
        "abstract": norm_text(paper.get("abstract")),
        "url": primary_link(paper),
        "semanticScholarUrl": paper.get("url") or "",
        "openAccessPdf": pdf.get("url") or "",
        "doi": ext.get("DOI") or "",
        "arxiv": ext.get("ArXiv") or "",
        "pubmed": ext.get("PubMed") or "",
        "corpusId": ext.get("CorpusId") or "",
        "fieldsOfStudy": "; ".join(
            sorted(
                {
                    norm_text(item.get("category"))
                    for item in (paper.get("s2FieldsOfStudy") or [])
                    if norm_text(item.get("category"))
                }
            )
        ),
        "publicationTypes": "; ".join(paper.get("publicationTypes") or []),
        "sourceQueries": "; ".join(sorted(paper.get("sourceQueries") or [])),
        "queryHitCount": len(paper.get("sourceQueries") or []),
        "candidateRank": candidate_rank,
        "importanceScore": score,
        "importanceReasons": reasons,
    }
    return enrich_paper(row)


def collect_papers():
    selected_by_year = {}
    candidates_by_year = {}
    for year in YEARS:
        merged = {}
        for query in QUERIES:
            print(f"[collect] {year} :: {query}", flush=True)
            try:
                papers = fetch_year_query(year, query)
            except Exception as exc:
                print(f"[warn] {year} {query}: {exc}", flush=True)
                continue
            for paper in papers:
                if paper.get("year") != year:
                    continue
                if not is_relevant(paper):
                    continue
                key = paper_key(paper)
                if key not in merged:
                    paper["sourceQueries"] = set()
                    merged[key] = paper
                merged[key]["sourceQueries"].add(query)
                if int(paper.get("citationCount") or 0) > int(merged[key].get("citationCount") or 0):
                    merged[key].update(paper)
            time.sleep(REQUEST_DELAY)

        ranked = sorted(
            merged.values(),
            key=lambda p: (
                int(p.get("citationCount") or 0),
                int(p.get("influentialCitationCount") or 0),
                importance_score(p)[0],
                p.get("title") or "",
            ),
            reverse=True,
        )
        candidate_pool = ranked[:CANDIDATES_PER_YEAR]
        normalized = [normalize_paper(p, year, i + 1) for i, p in enumerate(candidate_pool)]
        candidates_by_year[year] = normalized
        selected_by_year[year] = []
        print(
            f"[collect] {year}: retained {len(normalized):,}/{CANDIDATES_PER_YEAR:,} candidates from {len(ranked):,} relevant records",
            flush=True,
        )

    all_candidates = [p for rows in candidates_by_year.values() for p in rows]
    selected = sorted(
        all_candidates,
        key=lambda p: (
            p["citationCount"],
            p["influentialCitationCount"],
            p["importanceScore"],
            p["title"],
        ),
        reverse=True,
    )[:TARGET_TOTAL]
    for rank, paper in enumerate(selected, 1):
        paper["rank"] = rank
        selected_by_year[paper["year"]].append(paper)
    for rows in selected_by_year.values():
        rows.sort(key=lambda p: p["rank"])
    print(f"[collect] selected {len(selected):,} top-cited papers overall", flush=True)
    return selected, selected_by_year, candidates_by_year


def reuse_existing_candidates():
    path = DATA_DIR / CANDIDATES_JSON
    if not path.exists():
        raise FileNotFoundError(f"Missing existing candidate pool: {path}")
    payload = json.loads(path.read_text(encoding="utf-8"))
    rows = []
    for row in payload.get("candidates") or []:
        normalized = dict(row)
        normalized["year"] = int(normalized.get("year") or 0)
        normalized["citationCount"] = int(normalized.get("citationCount") or 0)
        normalized["influentialCitationCount"] = int(normalized.get("influentialCitationCount") or 0)
        normalized["importanceScore"] = float(normalized.get("importanceScore") or 0)
        rows.append(enrich_paper(normalized))
    candidates_by_year = {}
    for year in YEARS:
        year_rows = [p for p in rows if p["year"] == year]
        year_rows.sort(key=lambda p: int(p.get("candidateRank") or 999999))
        candidates_by_year[year] = year_rows[:CANDIDATES_PER_YEAR]
    all_candidates = [p for values in candidates_by_year.values() for p in values]
    selected = sorted(
        all_candidates,
        key=lambda p: (
            p["citationCount"],
            p["influentialCitationCount"],
            p["importanceScore"],
            p["title"],
        ),
        reverse=True,
    )[:TARGET_TOTAL]
    selected_by_year = {year: [] for year in YEARS}
    for rank, paper in enumerate(selected, 1):
        paper["rank"] = rank
        selected_by_year[paper["year"]].append(paper)
    for values in selected_by_year.values():
        values.sort(key=lambda p: p["rank"])
    print(f"[reuse] regenerated {len(selected):,} selected papers from {len(all_candidates):,} existing candidates", flush=True)
    return selected, selected_by_year, candidates_by_year


def csv_ready(row):
    out = {}
    for key, value in row.items():
        if isinstance(value, (list, dict, set)):
            out[key] = json.dumps(value, ensure_ascii=False)
        else:
            out[key] = value
    return out


def write_csv(path, rows, fields):
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(csv_ready(row) for row in rows)


def write_data(selected, selected_by_year, candidates_by_year):
    DATA_DIR.mkdir(exist_ok=True)
    flat_candidates = [p for rows in candidates_by_year.values() for p in rows]
    metadata = {
        "topic": "AI research",
        "source": "Semantic Scholar Academic Graph bulk search",
        "generated": date.today().isoformat(),
        "years": YEARS,
        "candidate_pool_per_year": CANDIDATES_PER_YEAR,
        "selected_total": TARGET_TOTAL,
        "ranking": "citationCount desc, influentialCitationCount desc",
        "queries": QUERIES,
        "openalex_concept": "https://openalex.org/C154945302",
        "openalex_subfield": "https://openalex.org/subfields/1702",
    }
    (DATA_DIR / PAPERS_JSON).write_text(
        json.dumps({"metadata": metadata, "papers": selected, "byYear": selected_by_year}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (DATA_DIR / CANDIDATES_JSON).write_text(
        json.dumps({"metadata": metadata, "candidates": flat_candidates, "byYear": candidates_by_year}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    fields = [
        "rank",
        "candidateRank",
        "year",
        "title",
        "authors",
        "venue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "importanceScore",
        "category",
        "methodTags",
        "keywordTags",
        "keyIdea",
        "strengths",
        "limitations",
        "url",
        "semanticScholarUrl",
        "openAccessPdf",
        "doi",
        "arxiv",
        "pubmed",
        "corpusId",
        "fieldsOfStudy",
        "publicationTypes",
        "sourceQueries",
        "queryHitCount",
        "importanceReasons",
        "abstract",
    ]
    candidate_fields = [field for field in fields if field != "rank"]
    write_csv(DATA_DIR / PAPERS_CSV, selected, fields)
    write_csv(DATA_DIR / CANDIDATES_CSV, flat_candidates, candidate_fields)
    for year, rows in candidates_by_year.items():
        stem = f"candidates_top1000_{year}"
        (DATA_DIR / f"{stem}.json").write_text(
            json.dumps({"metadata": metadata, "year": year, "candidates": rows}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        write_csv(DATA_DIR / f"{stem}.csv", rows, candidate_fields)
    for year, rows in selected_by_year.items():
        write_csv(DATA_DIR / f"papers_{year}.csv", rows, fields)
    return flat_candidates


def category_groups(rows):
    groups = defaultdict(list)
    for paper in rows:
        groups[paper["category"]].append(paper)
    for papers in groups.values():
        papers.sort(key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["importanceScore"]), reverse=True)
    return groups


def category_stats(rows):
    return Counter(p["category"] for p in rows)


def year_stats(rows):
    stats = {}
    for year in YEARS:
        year_rows = [p for p in rows if p["year"] == year]
        if not year_rows:
            continue
        stats[year] = {
            "count": len(year_rows),
            "citations": sum(p["citationCount"] for p in year_rows),
            "top": max(year_rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"])),
        }
    return stats


def period_key(start, end):
    return f"{start}-{end}"


def period_label(start, end):
    if start == START_YEAR and end == END_YEAR:
        return f"All years ({YEAR_RANGE_TEXT})"
    return str(start) if start == end else f"{start}-{end}"


def all_period_ranges():
    return [(start, end) for start in YEARS for end in range(start, END_YEAR + 1)]


def period_select_ranges():
    full = (START_YEAR, END_YEAR)
    return [full] + [pair for pair in all_period_ranges() if pair != full]


def top_metadata_values(rows, key, limit=3):
    counts = Counter()
    for row in rows:
        for value in str(row.get(key) or "").split("; "):
            if value:
                counts[value] += 1
    return [value for value, _ in counts.most_common(limit)]


def category_display_name(category, language):
    if language == "ko":
        return KOREAN_CATEGORY_NAMES.get(category, category)
    return category


def language_period_analysis(language, category, rows, start, end):
    count = len(rows)
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"]))
    active_year, active_count = Counter(p["year"] for p in rows).most_common(1)[0]
    tags = ", ".join(top_metadata_values(rows, "methodTags")) or "metadata-ranked"
    venues = ", ".join(top_metadata_values(rows, "venue")) or "mixed venues"
    name = category_display_name(category, language)
    if language == "ko":
        return {
            "categoryName": name,
            "overview": [
                f"{start}-{end} 기간의 {name} 분류에는 선정 논문 {count:,}편과 인용 {citations:,}회가 포함됩니다. 가장 활발한 연도는 {active_year}년({active_count:,}편)이며, 대표 상위 논문은 '{top['title']}'({top['citationCount']:,}회 인용)입니다.",
                f"이 기간의 주요 방법 태그는 {tags}이고, 자주 보이는 venue는 {venues}입니다.",
            ],
            "limitations": [
                TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General AI Methods and Systems"])[0],
                "기간별 citation ranking은 최근 논문을 구조적으로 불리하게 만들 수 있으므로 최신성은 별도 전문가 검토가 필요합니다.",
            ],
        }
    return {
        "categoryName": name,
        "overview": [
            f"In {start}-{end}, {name} contains {count:,} selected papers and {citations:,} citations. The most active year is {active_year} ({active_count:,} papers), and the leading citation-ranked paper is \"{top['title']}\" ({top['citationCount']:,} citations).",
            f"Frequent method tags include {tags}, with venue concentration around {venues}.",
        ],
        "limitations": [
            TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General AI Methods and Systems"])[0],
            "Citation-ranked period views structurally disadvantage very recent papers, so novelty needs separate expert review.",
        ],
    }


def build_period_analysis(selected):
    ranges = [
        {"key": period_key(start, end), "label": period_label(start, end), "from": start, "to": end}
        for start, end in period_select_ranges()
    ]
    analysis = {}
    for start, end in all_period_ranges():
        rows = [p for p in selected if start <= p["year"] <= end]
        groups = category_groups(rows)
        entry = {}
        for category, papers in groups.items():
            entry[safe_slug(category)] = {
                language: language_period_analysis(language, category, papers, start, end)
                for language in LANGUAGES
            }
        analysis[period_key(start, end)] = entry
    return {
        "generated": date.today().isoformat(),
        "yearRange": YEAR_RANGE_TEXT,
        "languages": LANGUAGES,
        "uiLabels": UI_LABELS,
        "ranges": ranges,
        "analysis": analysis,
    }


def write_period_analysis(selected):
    payload = build_period_analysis(selected)
    for target in (DATA_DIR / PERIOD_ANALYSIS_JSON, DOCS_DIR / "data" / PERIOD_ANALYSIS_JSON):
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(payload, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")


def write_taxonomy_dataset(selected):
    rows = []
    for category, papers in sorted(category_groups(selected).items()):
        overview = " ".join(TAXONOMY_TRENDS.get(category, []))
        limitations = " ".join(TAXONOMY_LIMITATIONS.get(category, []))
        for idx, paper in enumerate(papers, 1):
            row = dict(paper)
            row["taxonomyRank"] = idx
            row["categoryOverview"] = overview
            row["categoryLimitations"] = limitations
            rows.append(row)
    fields = [
        "category",
        "taxonomyRank",
        "rank",
        "candidateRank",
        "year",
        "title",
        "authors",
        "venue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "importanceScore",
        "categoryOverview",
        "categoryLimitations",
        "methodTags",
        "keywordTags",
        "keyIdea",
        "strengths",
        "limitations",
        "url",
        "semanticScholarUrl",
        "openAccessPdf",
        "doi",
        "arxiv",
        "pubmed",
        "fieldsOfStudy",
        "sourceQueries",
        "importanceReasons",
    ]
    write_csv(DATA_DIR / TAXONOMY_CSV, rows, fields)


def shields_keyword_badge(keyword):
    color = KEYWORD_COLORS.get(keyword, "64748b")
    badge = keyword.replace("-", "--")
    return f"https://img.shields.io/badge/keyword-{badge}-{color}"


def readme_keyword_badges(keywords):
    return " ".join(
        f'<img alt="{html.escape(keyword)}" src="{shields_keyword_badge(keyword)}">'
        for keyword in keywords
    )


def readme_keyword_convention_lines():
    lines = []
    for keyword, description, _ in KEYWORD_CONVENTION:
        lines.append(
            f"- ![{keyword}]({shields_keyword_badge(keyword)}) **{keyword}**: {description}"
        )
    return lines


def md_link(label, url):
    return f"[{label}]({url})" if url else label


def readme_taxonomy_table(rows, max_rows=10):
    out = [
        '<table width="100%">',
        "<colgroup>",
        '<col width="5%">',
        '<col width="22%">',
        '<col width="12%">',
        '<col width="12%">',
        '<col width="25%">',
        '<col width="12%">',
        '<col width="12%">',
        "</colgroup>",
        "<thead><tr>",
        '<th align="right">Rank</th><th>Paper</th><th>Meta</th><th>Keywords</th><th>Key idea</th><th>Strengths</th><th>Limitations</th>',
        "</tr></thead><tbody>",
    ]
    for paper in rows[:max_rows]:
        keywords = [x for x in paper["keywordTags"].split("; ") if x]
        out.extend(
            [
                "<tr>",
                f'<td align="right">{paper["rank"]}</td>',
                f'<td>{md_link(html.escape(paper["title"]), html.escape(paper["url"]))}<br><sub>{html.escape(paper["authors"] or "Unknown authors")}</sub></td>',
                f'<td>{paper["year"]}<br>{html.escape(paper["venue"] or "Unknown venue")}<br>{paper["citationCount"]:,} citations</td>',
                f"<td>{readme_keyword_badges(keywords)}</td>",
                f"<td>{html.escape(paper['keyIdea'])}</td>",
                f"<td>{html.escape(paper['strengths'])}</td>",
                f"<td>{html.escape(paper['limitations'])}</td>",
                "</tr>",
            ]
        )
    if len(rows) > max_rows:
        out.append(
            f'<tr><td colspan="7"><em>{len(rows) - max_rows:,} additional selected papers in this category are available in the dataset and website.</em></td></tr>'
        )
    out.extend(["</tbody></table>"])
    return "\n".join(out)


def write_readme(selected, candidates):
    stats = category_stats(selected)
    years = year_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    candidate_counts = Counter(p["year"] for p in candidates)
    groups = category_groups(selected)
    lines = [
        "# Awesome AI",
        "",
        "[![Awesome](https://awesome.re/badge-flat.svg)](https://awesome.re)",
        "[![Open Interactive Website](https://img.shields.io/badge/Open%20Interactive%20Website-GitHub%20Pages-0f766e)](https://honggi82.github.io/awesome-ai/)",
        "",
        "A taxonomy-first, citation-ranked map of AI research from 2020 through 2026.",
        "",
        f"Generated on {date.today().isoformat()} from free public Semantic Scholar metadata. This edition investigates up to {CANDIDATES_PER_YEAR:,} AI-related candidate papers per year for {YEAR_RANGE_TEXT}, keeps an audited candidate pool of {len(candidates):,} records, selects the top {TARGET_TOTAL:,} papers overall by citation count, and reorganizes them by AI research taxonomy.",
        "",
        "## Project Links",
        "",
        "- Open Interactive Website: https://honggi82.github.io/awesome-ai/",
        f"- Selected dataset: `data/{PAPERS_CSV}`",
        f"- Taxonomy dataset with paper-level ideas, strengths, and limitations: `data/{TAXONOMY_CSV}`",
        f"- Precomputed period and language analysis: `data/{PERIOD_ANALYSIS_JSON}`",
        f"- Candidate Pool: `data/{CANDIDATES_CSV}`",
        "- English review draft: `paper/review_en.html`, `paper/review_en.docx`",
        "- Korean review draft: `paper/review_ko.html`",
        "",
        "## Keywords Convention",
        "",
        "These badges define the AI keyword tags used to read and extend this collection.",
        "",
        *readme_keyword_convention_lines(),
        "",
        "## Taxonomy Overview",
        "",
        f"- **Total selected papers**: {len(selected):,} papers",
        f"- **Candidate pool audited**: {len(candidates):,} papers ({', '.join(f'{year}: {candidate_counts[year]:,}' for year in YEARS)})",
        f"- **Citation count in selected set**: {total_cites:,}",
    ]
    for category, count in stats.most_common():
        lines.append(f"- **{category}**: {count:,} papers")
    lines.extend(["", "## Taxonomy Collections", ""])
    for category, count in stats.most_common():
        rows = groups[category]
        citations = sum(p["citationCount"] for p in rows)
        covered_years = sorted({p["year"] for p in rows})
        lines.extend(
            [
                f"### {category}",
                "",
                f"- Papers selected: **{count:,}**",
                f"- Years covered: **{covered_years[0]}-{covered_years[-1]}**",
                f"- Citation count in selected set: **{citations:,}**",
                "- Category Overview (main research trends):",
            ]
        )
        lines.extend(f"  - {item}" for item in TAXONOMY_TRENDS[category])
        lines.append("- Limitations:")
        lines.extend(f"  - {item}" for item in TAXONOMY_LIMITATIONS[category])
        lines.extend(
            [
                "",
                f"<details>",
                f"<summary><strong>Show representative papers for {category}</strong></summary>",
                "",
                readme_taxonomy_table(rows),
                "",
                "</details>",
                "",
            ]
        )
    lines.extend(["## Yearly Coverage", ""])
    lines.append("| Year | Candidate papers audited | Selected top-100 papers | Citations in selected set | Top selected paper |")
    lines.append("|---:|---:|---:|---:|---|")
    for year in YEARS:
        stat = years.get(year)
        top = md_link(html.escape(stat["top"]["title"]), html.escape(stat["top"]["url"])) if stat else "-"
        lines.append(
            f"| {year} | {candidate_counts[year]:,} | {stat['count'] if stat else 0:,} | {stat['citations'] if stat else 0:,} | {top} |"
        )
    lines.extend(
        [
            "",
            "## Methodology",
            "",
            "The collection uses Semantic Scholar Academic Graph bulk search. Queries cover broad AI, machine learning, deep learning, foundation models, language, vision, reinforcement learning, generative models, graph learning, multimodal learning, trustworthy AI, and AI-for-science themes. For each year from 2020 through 2026, results are filtered to the publication year, screened with explicit AI relevance expressions in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most 1,000 candidates by citation count. The final awesome list selects the top 100 papers overall by citation count from those audited yearly pools; influential citation count and a deterministic metadata importance score are retained as tie-breakers and audit signals.",
            "",
            "The taxonomy, key ideas, strengths, limitations, method tags, and keyword tags are generated deterministically from public metadata and rule-based domain conventions. No paid API, paid LLM, paid translation, or paid compute was used.",
            "",
            "## Caveats",
            "",
            "- This is a metadata-driven citation map, not a full systematic review of every PDF.",
            f"- Citation counts favor older papers; {END_YEAR} should be interpreted as a partial and still-moving year.",
            "- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.",
            "- Citation ranking measures influence and visibility; it does not directly measure methodological quality, safety, or reproducibility.",
            "",
            "## License",
            "",
            "CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.",
        ]
    )
    (ROOT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def chart_color(index):
    colors = ["#2563eb", "#0f766e", "#a855f7", "#f59e0b", "#dc2626", "#0891b2", "#be123c", "#16a34a", "#64748b"]
    return colors[index % len(colors)]


def write_svg(path, content):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def svg_text(value):
    return html.escape(str(value), quote=True)


def render_category_chart(rows, start, end):
    data = category_stats(rows).most_common()
    width = 1000
    row_h = 44
    height = max(260, 96 + row_h * max(1, len(data)))
    max_value = max([value for _, value in data] or [1])
    bars = []
    if not data:
        bars.append('<text x="500" y="150" text-anchor="middle" fill="#64748b" font-size="24">No selected papers</text>')
    for i, (category, value) in enumerate(data):
        y = 72 + i * row_h
        bar_w = int((value / max_value) * 520)
        bars.append(f'<text x="32" y="{y + 24}" fill="#172033" font-size="15">{svg_text(category[:52])}</text>')
        bars.append(f'<rect x="420" y="{y + 4}" width="{bar_w}" height="26" rx="5" fill="{chart_color(i)}"/>')
        bars.append(f'<text x="{430 + bar_w}" y="{y + 24}" fill="#172033" font-size="15" font-weight="700">{value}</text>')
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="img" aria-label="Category distribution {start}-{end}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="32" y="38" fill="#0f172a" font-size="24" font-weight="800">AI Paper Taxonomy, {period_label(start, end)}</text>
<text x="32" y="60" fill="#64748b" font-size="14">Selected papers by category</text>
{''.join(bars)}
</svg>"""


def render_citation_chart(rows, start, end):
    width = 1000
    height = 420
    stats = year_stats(rows)
    years = list(range(start, end + 1))
    values = [stats.get(year, {}).get("citations", 0) for year in years]
    max_value = max(values or [1]) or 1
    plot_x = 80
    plot_y = 70
    plot_w = 850
    plot_h = 270
    gap = 14
    bar_w = max(24, int((plot_w - gap * (len(years) - 1)) / max(1, len(years))))
    bars = []
    for i, (year, value) in enumerate(zip(years, values)):
        h = int((value / max_value) * plot_h) if value else 0
        x = plot_x + i * (bar_w + gap)
        y = plot_y + plot_h - h
        bars.append(f'<rect x="{x}" y="{y}" width="{bar_w}" height="{h}" rx="5" fill="#7c3aed"/>')
        bars.append(f'<text x="{x + bar_w / 2:.1f}" y="{plot_y + plot_h + 28}" text-anchor="middle" fill="#475569" font-size="14">{year}</text>')
        if value:
            bars.append(f'<text x="{x + bar_w / 2:.1f}" y="{max(58, y - 8)}" text-anchor="middle" fill="#172033" font-size="13" font-weight="700">{value:,}</text>')
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="img" aria-label="Yearly citation mass {start}-{end}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="32" y="38" fill="#0f172a" font-size="24" font-weight="800">Yearly Citation Mass, {period_label(start, end)}</text>
<text x="32" y="60" fill="#64748b" font-size="14">Citation counts among visible selected papers</text>
<line x1="{plot_x}" y1="{plot_y + plot_h}" x2="{plot_x + plot_w}" y2="{plot_y + plot_h}" stroke="#cbd5e1" stroke-width="2"/>
{''.join(bars)}
</svg>"""


def write_charts(selected):
    period_dir = DOCS_DIR / "assets" / "periods"
    if period_dir.exists():
        for stale in period_dir.glob("*.svg"):
            stale.unlink()
    for start, end in all_period_ranges():
        rows = [p for p in selected if start <= p["year"] <= end]
        write_svg(period_dir / f"category_distribution_{start}_{end}.svg", render_category_chart(rows, start, end))
        write_svg(period_dir / f"yearly_citations_{start}_{end}.svg", render_citation_chart(rows, start, end))
    assets = DOCS_DIR / "assets"
    shutil.copyfile(period_dir / f"category_distribution_{START_YEAR}_{END_YEAR}.svg", assets / "category_distribution.svg")
    shutil.copyfile(period_dir / f"yearly_citations_{START_YEAR}_{END_YEAR}.svg", assets / "yearly_citations.svg")


def taxonomy_icon_svg(category):
    colors = {
        "Foundation Models and Large Language Models": ("#2563eb", "#93c5fd"),
        "Generative Models and Synthetic Media": ("#a855f7", "#f0abfc"),
        "Vision and Multimodal Learning": ("#0f766e", "#67e8f9"),
        "Natural Language Processing and Knowledge": ("#f59e0b", "#fde68a"),
        "Reinforcement Learning and Agents": ("#dc2626", "#fca5a5"),
        "Representation, Self-Supervised, and Transfer Learning": ("#0891b2", "#7dd3fc"),
        "Trustworthy, Explainable, and Responsible AI": ("#be123c", "#f9a8d4"),
        "Graph Learning, Recommendation, and Core Methods": ("#4f46e5", "#c4b5fd"),
        "AI for Science, Healthcare, and Robotics": ("#16a34a", "#86efac"),
        "General AI Methods and Systems": ("#334155", "#cbd5e1"),
    }
    accent, soft = colors.get(category, colors["General AI Methods and Systems"])
    label = category.split(" and ")[0][:28]
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="320" height="200" viewBox="0 0 320 200" role="img" aria-label="{svg_text(category)}">
<rect x="0" y="0" width="320" height="200" rx="18" fill="#f8fafc"/>
<circle cx="82" cy="78" r="34" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<circle cx="164" cy="58" r="24" fill="#fff" stroke="{accent}" stroke-width="5"/>
<circle cx="222" cy="120" r="36" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<path d="M112 78 C132 64 138 60 142 58" stroke="{accent}" stroke-width="6" fill="none" stroke-linecap="round"/>
<path d="M182 72 C200 88 206 100 212 108" stroke="{accent}" stroke-width="6" fill="none" stroke-linecap="round"/>
<path d="M92 112 C118 148 176 158 214 140" stroke="{accent}" stroke-width="4" fill="none" stroke-linecap="round" stroke-dasharray="10 8"/>
<text x="160" y="178" text-anchor="middle" fill="#172033" font-size="18" font-weight="800">{svg_text(label)}</text>
</svg>"""


def write_taxonomy_icons(selected):
    icon_dir = DOCS_DIR / "assets" / "taxonomy"
    icon_dir.mkdir(parents=True, exist_ok=True)
    for category in category_stats(selected):
        write_svg(icon_dir / f"{safe_slug(category)}.svg", taxonomy_icon_svg(category))


def html_attrs(paper):
    attrs = {
        "data-year": paper["year"],
        "data-citations": paper["citationCount"],
        "data-keywords": " ".join(paper["keywordTags"].split("; ")),
        "data-title": paper["title"],
        "data-venue": paper["venue"],
    }
    for language in LANGUAGES:
        fields = localized_paper_fields(paper, language)
        suffix = language
        attrs[f"data-key-idea-{suffix}"] = fields["keyIdea"]
        attrs[f"data-strengths-{suffix}"] = fields["strengths"]
        attrs[f"data-paper-limitations-{suffix}"] = fields["limitations"]
    return " ".join(f'{key}="{html.escape(str(value), quote=True)}"' for key, value in attrs.items())


def keyword_chips_html(keyword_string):
    tags = [tag for tag in keyword_string.split("; ") if tag]
    return "".join(
        f'<span class="keyword-chip paper-keyword" style="--chip-color:#{KEYWORD_COLORS.get(tag, "64748b")}">{html.escape(tag)}</span>'
        for tag in tags
    )


def paper_card(paper, taxonomy_rank):
    semantic = f'<a href="{html.escape(paper["semanticScholarUrl"])}">Semantic Scholar</a>' if paper.get("semanticScholarUrl") else ""
    pdf = f'<a href="{html.escape(paper["openAccessPdf"])}">PDF</a>' if paper.get("openAccessPdf") else ""
    doi = f'<a href="{html.escape(paper["url"])}">Paper</a>' if paper.get("url") else ""
    links = " ".join(x for x in (doi, semantic, pdf) if x)
    return f"""
      <article class="paper-card" {html_attrs(paper)}>
        <div class="paper-rank">#{paper['rank']}</div>
        <div>
          <h3><a href="{html.escape(paper['url'] or paper['semanticScholarUrl'])}">{html.escape(paper['title'])}</a></h3>
          <p class="authors">{html.escape(paper['authors'] or 'Unknown authors')}</p>
          <div class="meta">
            <span>{paper['year']}</span>
            <span>{html.escape(paper['venue'] or 'Unknown venue')}</span>
            <span>{paper['citationCount']:,} citations</span>
            <span>{paper['influentialCitationCount']:,} influential</span>
            <span>score {paper['importanceScore']}</span>
            <span>taxonomy #{taxonomy_rank}</span>
          </div>
          <div class="paper-keywords" aria-label="Keywords">{keyword_chips_html(paper['keywordTags'])}</div>
          <div class="assessment">
            <p><strong class="paper-key-idea-label">Key idea:</strong> <span class="paper-key-idea-text">{html.escape(paper['keyIdea'])}</span></p>
            <p><strong class="paper-strengths-label">Strengths:</strong> <span class="paper-strengths-text">{html.escape(paper['strengths'])}</span></p>
            <p><strong class="paper-limitations-label">Limitations:</strong> <span class="paper-limitations-text">{html.escape(paper['limitations'])}</span></p>
          </div>
          <p class="links">{links}</p>
        </div>
      </article>
"""


def taxonomy_section(category, rows):
    slug = safe_slug(category)
    citations = sum(p["citationCount"] for p in rows)
    years = sorted({p["year"] for p in rows})
    top = rows[0]
    cards = "".join(paper_card(paper, idx) for idx, paper in enumerate(rows, 1))
    trends = "".join(f"<li>{html.escape(item)}</li>" for item in TAXONOMY_TRENDS[category])
    limitations = "".join(f"<li>{html.escape(item)}</li>" for item in TAXONOMY_LIMITATIONS[category])
    return f"""
    <section class="taxonomy-section" data-category="{slug}">
      <details>
        <summary>
          <img class="summary-thumb" src="assets/taxonomy/{slug}.svg" alt="">
          <span class="summary-title">{html.escape(category)}</span>
          <span class="category-count">{len(rows):,} papers</span>
          <span class="category-years">{years[0]}-{years[-1]}</span>
          <span class="category-citations">{citations:,} citations</span>
        </summary>
        <div class="section-intro">
          <figure class="section-visual"><img src="assets/taxonomy/{slug}.svg" alt="{html.escape(category)} illustration"></figure>
          <p><strong>Top paper:</strong> <span class="top-paper">{html.escape(top['title'])}</span></p>
          <div class="insight-grid">
            <div class="insight-box">
              <strong class="overview-heading">Category Overview</strong>
              <ul class="category-overview-list">{trends}</ul>
            </div>
            <div class="insight-box limitation-box">
              <strong class="limitation-heading">Limitations</strong>
              <ul class="category-limitations-list">{limitations}</ul>
            </div>
          </div>
        </div>
        <div class="paper-list">{cards}</div>
      </details>
    </section>
"""


def site_keyword_convention_html():
    return "\n".join(
        f"<button class='keyword-item' type='button' data-keyword='{html.escape(keyword)}' aria-pressed='false'><span class='keyword-chip' style='--chip-color:#{color}'>{html.escape(keyword)}</span><span>{html.escape(description)}</span></button>"
        for keyword, description, color in KEYWORD_CONVENTION
    )


def write_site(selected):
    DOCS_DIR.mkdir(exist_ok=True)
    (DOCS_DIR / "data").mkdir(exist_ok=True)
    (DOCS_DIR / "paper").mkdir(exist_ok=True)
    groups = category_groups(selected)
    stats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    start_year_options = "\n".join(f'<option value="{year}"{" selected" if year == START_YEAR else ""}>{year}</option>' for year in YEARS)
    end_year_options = "\n".join(f'<option value="{year}"{" selected" if year == END_YEAR else ""}>{year}</option>' for year in YEARS)
    period_options = "\n".join(
        f'<option value="{period_key(start, end)}" data-from="{start}" data-to="{end}"{" selected" if start == START_YEAR and end == END_YEAR else ""}>{period_label(start, end)}</option>'
        for start, end in period_select_ranges()
    )
    language_options = "\n".join(
        f'<option value="{code}"{" selected" if code == "en" else ""}>{html.escape(label)}</option>'
        for code, label in LANGUAGES.items()
    )
    sections = "".join(taxonomy_section(category, groups[category]) for category, _ in stats.most_common())
    keyword_html = site_keyword_convention_html()
    year_script = f"""
  <script>
    (() => {{
      const periodSelect = document.getElementById("periodPreset");
      const languageSelect = document.getElementById("languageSelect");
      const startSelect = document.getElementById("startYear");
      const endSelect = document.getElementById("endYear");
      const resetButton = document.getElementById("resetYears");
      const rangeStatus = document.getElementById("rangeStatus");
      const statPapers = document.getElementById("statPapers");
      const statYears = document.getElementById("statYears");
      const statCitations = document.getElementById("statCitations");
      const statCategories = document.getElementById("statCategories");
      const taxonomyTotalSummary = document.getElementById("taxonomyTotalSummary");
      const keywordFilterStatus = document.getElementById("keywordFilterStatus");
      const categoryChart = document.getElementById("categoryDistributionChart");
      const citationChart = document.getElementById("yearlyCitationsChart");
      const categoryChartCaption = document.getElementById("categoryChartCaption");
      const citationChartCaption = document.getElementById("citationChartCaption");
      const defaultStart = startSelect.value;
      const defaultEnd = endSelect.value;
      const defaultLanguage = languageSelect.value;
      const periodOptions = Array.from(periodSelect.options);
      const validYears = Array.from(startSelect.options).map(option => option.value);
      const keywordGrid = document.querySelector(".keyword-grid");
      const keywordButtons = Array.from(document.querySelectorAll(".keyword-item[data-keyword]"));
      let precomputed = null;

      function formatNumber(value) {{ return Number(value).toLocaleString("en-US"); }}
      function rangeKey(start, end) {{ return `${{start}}-${{end}}`; }}
      function rangeLabel(start, end) {{ return start === end ? String(start) : `${{start}}-${{end}}`; }}
      function chartPath(kind, start, end) {{ return `assets/periods/${{kind}}_${{start}}_${{end}}.svg`; }}
      function selectedKeywords() {{
        return keywordButtons.filter(button => button.getAttribute("aria-pressed") === "true").map(button => button.dataset.keyword);
      }}
      function setKeywordPressed(button, pressed) {{
        button.setAttribute("aria-pressed", pressed ? "true" : "false");
        button.classList.toggle("is-selected", pressed);
      }}
      function keywordMatches(card, selected) {{
        if (!selected.length) return true;
        const cardKeywords = (card.dataset.keywords || "").split(" ").filter(Boolean);
        return selected.some(keyword => cardKeywords.includes(keyword));
      }}
      function yearRangeText(years) {{
        if (!years.length) return "No years";
        const sorted = [...new Set(years)].sort((a, b) => a - b);
        return sorted[0] === sorted[sorted.length - 1] ? String(sorted[0]) : `${{sorted[0]}}-${{sorted[sorted.length - 1]}}`;
      }}
      function selectedRangeValue(start, end) {{
        const match = periodOptions.find(option => option.dataset.from === String(start) && option.dataset.to === String(end));
        return match ? match.value : "custom";
      }}
      function updatePeriodSelect(start, end) {{ periodSelect.value = selectedRangeValue(start, end); }}
      function labels() {{
        const fallback = {{
          papers: "papers", categories: "categories", overview: "Category Overview", limitations: "Limitations",
          totalSelected: "Total selected papers", categoryCount: "Categories",
          keyIdea: "Key idea", strengths: "Strengths", paperLimitations: "Limitations"
        }};
        return {{...fallback, ...(precomputed?.uiLabels?.en || {{}}), ...(precomputed?.uiLabels?.[languageSelect.value] || {{}})}};
      }}
      function setList(target, items) {{
        if (!target || !items) return;
        target.innerHTML = "";
        items.forEach(item => {{
          const li = document.createElement("li");
          li.textContent = item;
          target.appendChild(li);
        }});
      }}
      function updateCharts(start, end) {{
        const label = rangeLabel(start, end);
        if (categoryChart) {{
          categoryChart.src = chartPath("category_distribution", start, end);
          categoryChart.alt = `Category distribution chart for ${{label}}`;
        }}
        if (citationChart) {{
          citationChart.src = chartPath("yearly_citations", start, end);
          citationChart.alt = `Yearly citation chart for ${{label}}`;
        }}
        if (categoryChartCaption) categoryChartCaption.textContent = `Category distribution (${{label}})`;
        if (citationChartCaption) citationChartCaption.textContent = `Yearly citation mass (${{label}})`;
      }}
      function localizedCardText(card, field, language) {{
        const suffix = language.charAt(0).toUpperCase() + language.slice(1);
        return card.dataset[`${{field}}${{suffix}}`] || card.dataset[`${{field}}En`] || "";
      }}
      function applyPaperLocalization(card, copy) {{
        const language = languageSelect.value;
        [
          ["keyIdea", ".paper-key-idea-label", ".paper-key-idea-text", copy.keyIdea],
          ["strengths", ".paper-strengths-label", ".paper-strengths-text", copy.strengths],
          ["paperLimitations", ".paper-limitations-label", ".paper-limitations-text", copy.paperLimitations]
        ].forEach(([field, labelSelector, textSelector, label]) => {{
          const labelNode = card.querySelector(labelSelector);
          const textNode = card.querySelector(textSelector);
          if (labelNode) labelNode.textContent = `${{label}}:`;
          if (textNode) textNode.textContent = localizedCardText(card, field, language);
        }});
      }}
      function applyPrecomputedAnalysis(section, start, end) {{
        const entry = precomputed?.analysis?.[rangeKey(start, end)]?.[section.dataset.category];
        const analysis = entry?.[languageSelect.value] || entry?.en;
        if (!analysis) return;
        const copy = labels();
        const title = section.querySelector(".summary-title");
        if (title) title.textContent = analysis.categoryName;
        const overviewHeading = section.querySelector(".overview-heading");
        const limitationHeading = section.querySelector(".limitation-heading");
        if (overviewHeading) overviewHeading.textContent = copy.overview;
        if (limitationHeading) limitationHeading.textContent = copy.limitations;
        setList(section.querySelector(".category-overview-list"), analysis.overview);
        setList(section.querySelector(".category-limitations-list"), analysis.limitations);
      }}
      function syncUrl(start, end) {{
        const url = new URL(window.location.href);
        const language = languageSelect.value;
        if (language === defaultLanguage) url.searchParams.delete("lang"); else url.searchParams.set("lang", language);
        if (String(start) === defaultStart && String(end) === defaultEnd) {{
          url.searchParams.delete("period"); url.searchParams.delete("from"); url.searchParams.delete("to");
        }} else {{
          const value = selectedRangeValue(start, end);
          if (value !== "custom") {{ url.searchParams.set("period", value); url.searchParams.delete("from"); url.searchParams.delete("to"); }}
          else {{ url.searchParams.delete("period"); url.searchParams.set("from", start); url.searchParams.set("to", end); }}
        }}
        const keywords = selectedKeywords();
        if (keywords.length) url.searchParams.set("keywords", keywords.join(",")); else url.searchParams.delete("keywords");
        window.history.replaceState(null, "", url);
      }}
      function setFromUrl() {{
        const params = new URLSearchParams(window.location.search);
        const lang = params.get("lang");
        if (lang && Array.from(languageSelect.options).some(option => option.value === lang)) languageSelect.value = lang;
        const requestedKeywords = (params.get("keywords") || "").split(",").filter(Boolean);
        const keywordValues = keywordButtons.map(button => button.dataset.keyword);
        const activeKeyword = requestedKeywords.find(keyword => keywordValues.includes(keyword)) || "";
        keywordButtons.forEach(button => setKeywordPressed(button, button.dataset.keyword === activeKeyword));
        const period = params.get("period");
        if (period) {{
          const option = periodOptions.find(item => item.value === period && item.dataset.from && item.dataset.to);
          if (option) {{ periodSelect.value = period; startSelect.value = option.dataset.from; endSelect.value = option.dataset.to; return; }}
        }}
        const from = params.get("from");
        const to = params.get("to");
        if (validYears.includes(from)) startSelect.value = from;
        if (validYears.includes(to)) endSelect.value = to;
        updatePeriodSelect(startSelect.value, endSelect.value);
      }}
      function updateKeywordFilterStatus(selected, totalPapers, copy) {{
        if (!keywordFilterStatus) return;
        const keyword = selected.length ? selected[0] : "all";
        keywordFilterStatus.textContent = `Selected keyword: ${{keyword}} | Matching papers: ${{formatNumber(totalPapers)}} ${{copy.papers}}`;
      }}
      function applyYearFilter(sync = true) {{
        let start = Number(startSelect.value);
        let end = Number(endSelect.value);
        if (start > end) {{ const previous = start; start = end; end = previous; startSelect.value = String(start); endSelect.value = String(end); }}
        const copy = labels();
        const activeKeywords = selectedKeywords();
        let totalPapers = 0;
        let totalCitations = 0;
        let activeCategories = 0;
        const activeYears = [];
        document.querySelectorAll(".taxonomy-section").forEach(section => {{
          let sectionCount = 0;
          let sectionCitations = 0;
          const sectionYears = [];
          section.querySelectorAll(".paper-card").forEach(card => {{
            const year = Number(card.dataset.year);
            const citations = Number(card.dataset.citations || 0);
            const visible = year >= start && year <= end && keywordMatches(card, activeKeywords);
            applyPaperLocalization(card, copy);
            card.hidden = !visible;
            if (visible) {{ sectionCount += 1; sectionCitations += citations; sectionYears.push(year); activeYears.push(year); }}
          }});
          const hasPapers = sectionCount > 0;
          section.hidden = !hasPapers;
          if (!hasPapers) return;
          activeCategories += 1;
          totalPapers += sectionCount;
          totalCitations += sectionCitations;
          section.querySelector(".category-count").textContent = `${{formatNumber(sectionCount)}} ${{copy.papers}}`;
          section.querySelector(".category-years").textContent = yearRangeText(sectionYears);
          section.querySelector(".category-citations").textContent = `${{formatNumber(sectionCitations)}} citations`;
          const topPaper = section.querySelector(".paper-card:not([hidden]) h3");
          const topPaperTarget = section.querySelector(".top-paper");
          if (topPaper && topPaperTarget) topPaperTarget.textContent = topPaper.textContent.trim();
          applyPrecomputedAnalysis(section, start, end);
        }});
        statPapers.textContent = formatNumber(totalPapers);
        statYears.textContent = formatNumber(new Set(activeYears).size);
        statCitations.textContent = formatNumber(totalCitations);
        statCategories.textContent = formatNumber(activeCategories);
        updatePeriodSelect(start, end);
        updateCharts(start, end);
        taxonomyTotalSummary.innerHTML = `<strong>${{copy.totalSelected}}:</strong> ${{formatNumber(totalPapers)}} ${{copy.papers}}; <strong>${{copy.categoryCount}}:</strong> ${{formatNumber(activeCategories)}} ${{copy.categories}}.`;
        updateKeywordFilterStatus(activeKeywords, totalPapers, copy);
        const keywordText = activeKeywords.length ? ` | ${{activeKeywords[0]}}` : "";
        rangeStatus.textContent = `${{start}}-${{end}} | ${{formatNumber(totalPapers)}} ${{copy.papers}} | ${{formatNumber(activeCategories)}} ${{copy.categories}}${{keywordText}}`;
        if (sync) syncUrl(start, end);
      }}
      setFromUrl();
      applyYearFilter(false);
      fetch("data/{PERIOD_ANALYSIS_JSON}")
        .then(response => response.json())
        .then(data => {{ precomputed = data; applyYearFilter(false); }})
        .catch(() => {{ rangeStatus.textContent = "Precomputed analysis could not be loaded."; }});
      periodSelect.addEventListener("change", () => {{
        const option = periodSelect.selectedOptions[0];
        if (option && option.dataset.from && option.dataset.to) {{ startSelect.value = option.dataset.from; endSelect.value = option.dataset.to; }}
        applyYearFilter(true);
      }});
      languageSelect.addEventListener("change", () => applyYearFilter(true));
      startSelect.addEventListener("change", () => applyYearFilter(true));
      endSelect.addEventListener("change", () => applyYearFilter(true));
      if (keywordGrid) {{
        keywordGrid.addEventListener("click", event => {{
          const button = event.target.closest(".keyword-item[data-keyword]");
          if (!button || !keywordGrid.contains(button)) return;
          event.preventDefault();
          const pressed = button.getAttribute("aria-pressed") === "true";
          keywordButtons.forEach(item => setKeywordPressed(item, !pressed && item === button));
          applyYearFilter(true);
        }});
      }}
      resetButton.addEventListener("click", () => {{
        startSelect.value = defaultStart;
        endSelect.value = defaultEnd;
        periodSelect.value = `${{defaultStart}}-${{defaultEnd}}`;
        keywordButtons.forEach(button => setKeywordPressed(button, false));
        applyYearFilter(true);
      }});
    }})();
  </script>
"""
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:,">
  <title>Awesome AI</title>
  <style>
    :root {{ color-scheme: light; --ink:#172033; --muted:#5b6678; --line:#d9dee8; --accent:#2563eb; --accent2:#0f766e; --bg:#f7f9fc; --panel:#ffffff; }}
    body {{ margin:0; font-family: Inter, Segoe UI, Arial, sans-serif; color:var(--ink); background:var(--bg); }}
    header {{ padding:54px 7vw 34px; background:linear-gradient(120deg,#eff6ff,#ecfdf5); border-bottom:1px solid var(--line); }}
    h1 {{ font-size:48px; margin:0 0 12px; letter-spacing:0; }}
    h2 {{ margin-top:36px; }}
    p {{ line-height:1.65; color:var(--muted); }}
    main {{ padding:28px 7vw 72px; }}
    nav a {{ display:inline-block; margin:0 12px 10px 0; color:#0f5f97; font-weight:700; }}
    a {{ color:#0f5f97; text-decoration:none; }}
    a:hover {{ text-decoration:underline; }}
    .stats {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:12px; margin:24px 0; }}
    .stat {{ background:white; border:1px solid var(--line); border-radius:8px; padding:16px; }}
    .stat strong {{ display:block; font-size:28px; color:var(--accent); }}
    .stat span {{ display:block; margin-top:8px; color:var(--muted); }}
    .filters {{ display:flex; flex-wrap:wrap; align-items:end; gap:12px; margin:24px 0; padding:14px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .filter-field {{ display:grid; gap:6px; }}
    .wide-field {{ min-width:min(100%,280px); }}
    .filter-field label {{ font-weight:700; font-size:13px; color:#344255; }}
    select {{ min-width:104px; height:38px; border:1px solid var(--line); border-radius:8px; background:white; color:var(--ink); padding:0 10px; font:inherit; }}
    #periodPreset {{ min-width:280px; }}
    button {{ min-height:38px; border:1px solid #bfc8d8; border-radius:8px; background:#f6f8fb; color:var(--ink); padding:0 14px; font-weight:700; cursor:pointer; }}
    button:hover {{ background:#eef2f7; }}
    #rangeStatus {{ color:var(--muted); font-weight:700; min-height:38px; display:inline-flex; align-items:center; }}
    .keyword-section {{ margin:28px 0; }}
    .keyword-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(240px,1fr)); gap:10px; }}
    .keyword-item {{ display:flex; gap:10px; align-items:flex-start; height:auto; min-height:54px; padding:12px; background:white; border:1px solid var(--line); border-radius:8px; color:var(--muted); line-height:1.45; text-align:left; font:inherit; cursor:pointer; }}
    .keyword-item[aria-pressed="true"], .keyword-item.is-selected {{ border-color:var(--accent); box-shadow:0 0 0 2px rgba(37,99,235,0.16); color:var(--ink); }}
    .keyword-chip {{ flex:0 0 auto; min-width:96px; text-align:center; background:var(--chip-color); color:white; border-radius:999px; padding:4px 9px; font-size:13px; font-weight:800; }}
    .keyword-filter-status {{ margin:10px 0 0; font-weight:700; color:var(--accent2); }}
    .figures {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(280px,1fr)); gap:16px; margin:24px 0; }}
    .chart-figure {{ margin:0; }}
    .chart-figure figcaption {{ margin-top:8px; color:var(--muted); font-size:13px; font-weight:700; }}
    .figures img {{ width:100%; aspect-ratio:16 / 9; object-fit:contain; background:white; border:1px solid var(--line); border-radius:8px; display:block; }}
    .taxonomy-section {{ margin-top:16px; }}
    details {{ background:var(--panel); border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    summary {{ cursor:pointer; display:grid; grid-template-columns:64px minmax(260px,1fr) repeat(3,minmax(110px,auto)); gap:12px; align-items:center; padding:14px 18px; font-weight:700; }}
    .summary-thumb {{ width:56px; height:40px; object-fit:cover; border:1px solid var(--line); border-radius:6px; background:#f8fafc; }}
    .summary-title {{ color:var(--accent); }}
    .section-intro {{ padding:0 18px 14px; border-top:1px solid var(--line); }}
    .section-visual {{ margin:14px 0 4px; }}
    .section-visual img {{ width:min(320px,100%); aspect-ratio:16 / 11; object-fit:contain; border:1px solid var(--line); border-radius:8px; background:#f8fafc; display:block; }}
    .insight-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:12px; margin-top:12px; }}
    .insight-box {{ padding:12px 14px; background:#f4f8ff; border:1px solid #cfe0ff; border-radius:8px; }}
    .limitation-box {{ background:#fff8f1; border-color:#ead7c1; }}
    .insight-box ul {{ margin:8px 0 0; padding-left:20px; color:var(--muted); line-height:1.55; }}
    .paper-list {{ display:grid; gap:12px; padding:16px; background:#f9fbfd; }}
    .paper-card {{ display:grid; grid-template-columns:56px 1fr; gap:14px; padding:16px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .paper-rank {{ font-weight:800; color:var(--accent2); }}
    .paper-card h3 {{ margin:0 0 6px; font-size:18px; line-height:1.35; }}
    .authors {{ margin:0 0 8px; }}
    .meta {{ display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 10px; }}
    .meta span {{ display:inline-block; background:#eef2f7; border:1px solid #dce3ee; border-radius:999px; padding:5px 9px; color:#344255; font-size:13px; }}
    .paper-keywords {{ display:flex; flex-wrap:wrap; gap:6px; margin:0 0 10px; }}
    .paper-keyword {{ min-width:0; font-size:12px; padding:3px 8px; }}
    .assessment {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:8px; }}
    .links a {{ margin-right:12px; font-weight:700; }}
    @media (max-width:760px) {{
      h1 {{ font-size:36px; }}
      summary {{ grid-template-columns:1fr; }}
      .paper-card {{ grid-template-columns:1fr; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>Awesome AI</h1>
    <p>A taxonomy-first, citation-ranked map of AI research from {START_YEAR} through {END_YEAR}. Each year investigates up to {CANDIDATES_PER_YEAR:,} candidate papers; the final collection selects the top {TARGET_TOTAL:,} papers overall by citation count.</p>
    <nav>
      <a href="https://github.com/honggi82/awesome-ai">README</a>
      <a href="data/{PAPERS_CSV}">CSV Dataset</a>
      <a href="data/{TAXONOMY_CSV}">Taxonomy CSV</a>
      <a href="data/{PERIOD_ANALYSIS_JSON}">Period Analysis JSON</a>
      <a href="#keywords-convention">Keywords Convention</a>
      <a href="data/{CANDIDATES_CSV}">Candidate Pool</a>
      <a href="paper/review_en.html">Review Paper</a>
      <a href="paper/review_ko.html">Korean Review</a>
    </nav>
  </header>
  <main>
    <div class="stats">
      <div class="stat"><strong id="statPapers">{len(selected)}</strong><span>selected papers</span></div>
      <div class="stat"><strong id="statYears">{len(year_stats(selected))}</strong><span>years represented</span></div>
      <div class="stat"><strong id="statCitations">{total_cites:,}</strong><span>citation count total</span></div>
      <div class="stat"><strong id="statCategories">{len(stats)}</strong><span>topic categories</span></div>
    </div>
    <form class="filters" id="yearFilter">
      <div class="filter-field wide-field">
        <label for="periodPreset">Period</label>
        <select id="periodPreset" name="period">{period_options}</select>
      </div>
      <div class="filter-field">
        <label for="languageSelect">Language</label>
        <select id="languageSelect" name="lang">{language_options}</select>
      </div>
      <div class="filter-field">
        <label for="startYear">Start year</label>
        <select id="startYear" name="from">{start_year_options}</select>
      </div>
      <div class="filter-field">
        <label for="endYear">End year</label>
        <select id="endYear" name="to">{end_year_options}</select>
      </div>
      <button type="button" id="resetYears">Reset</button>
      <span id="rangeStatus"></span>
    </form>
    <section class="keyword-section" id="keywords-convention">
      <h2>Keywords Convention</h2>
      <p>These clickable keyword tags define the AI-specific convention used to scan, filter, and extend this collection.</p>
      <div class="keyword-grid">{keyword_html}</div>
      <p class="keyword-filter-status" id="keywordFilterStatus">Selected keyword: all | Matching papers: {len(selected):,} papers</p>
    </section>
    <h2>Taxonomy</h2>
    <p id="taxonomyTotalSummary"><strong>Total selected papers:</strong> {len(selected):,} papers; <strong>Categories:</strong> {len(stats)} categories.</p>
    <p>Each taxonomy section lists papers with publication year, venue, citation count, influential citations, score, keywords, key idea, strengths, research-focused limitations, and paper links. Sections are collapsed by default.</p>
    <div class="figures">
      <figure class="chart-figure">
        <img id="categoryDistributionChart" src="assets/periods/category_distribution_{START_YEAR}_{END_YEAR}.svg" alt="Category distribution chart for {YEAR_RANGE_TEXT}">
        <figcaption id="categoryChartCaption">Category distribution ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
      <figure class="chart-figure">
        <img id="yearlyCitationsChart" src="assets/periods/yearly_citations_{START_YEAR}_{END_YEAR}.svg" alt="Yearly citation chart for {YEAR_RANGE_TEXT}">
        <figcaption id="citationChartCaption">Yearly citation mass ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
    </div>
    {sections}
  </main>
{year_script}
</body>
</html>
"""
    (DOCS_DIR / "index.html").write_text(html_doc, encoding="utf-8")
    (DOCS_DIR / ".nojekyll").write_text("", encoding="utf-8")


def reference_line(paper):
    venue = paper["venue"] or "Unknown venue"
    return f"{paper['authors'] or 'Unknown authors'}. ({paper['year']}). {paper['title']}. {venue}. {paper['url'] or paper['semanticScholarUrl']}"


def review_sections(selected, korean=False):
    stats = year_stats(selected)
    cats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    top_cited = sorted(selected, key=lambda p: p["citationCount"], reverse=True)[:12]
    top_scored = sorted(selected, key=lambda p: p["importanceScore"], reverse=True)[:12]
    peak_year = max(stats, key=lambda y: stats[y]["citations"])
    leading_cat, leading_count = cats.most_common(1)[0]
    if korean:
        title = f"{YEAR_RANGE_TEXT} AI 연구 동향: 공개 메타데이터 기반 citation-ranked 리뷰"
        abstract = (
            f"이 리뷰 초안은 {START_YEAR}년부터 {END_YEAR}년까지 AI 연구를 연도별 최대 {CANDIDATES_PER_YEAR:,}편의 후보 논문으로 조사하고, "
            f"그 후보군 중 citation이 높은 {TARGET_TOTAL:,}편을 선정해 taxonomy-first 방식으로 분석한다. "
            "선정과 분류는 Semantic Scholar 공개 메타데이터, 명시적 AI 관련성 필터, DOI/arXiv/PubMed/CorpusId/paperId 중복 제거, 인용수 정렬을 사용했다."
        )
        methods = (
            "각 연도에 대해 AI, machine learning, deep learning, foundation model, LLM, NLP, computer vision, reinforcement learning, generative AI, graph learning, multimodal learning, trustworthy AI, AI for science 관련 질의를 보냈다. "
            "제목/초록/venue에서 AI 관련 표현이 확인되는 record만 유지하고, 연도별 최대 1,000편을 citation count 기준 후보군으로 저장했다. 최종 목록은 전체 후보군에서 citation count 상위 100편이다."
        )
        findings = [
            f"선정 논문 {len(selected):,}편은 총 {total_cites:,}회의 인용을 포함하며, citation mass가 가장 큰 연도는 {peak_year}년이다.",
            f"가장 큰 분류는 {KOREAN_CATEGORY_NAMES.get(leading_cat, leading_cat)}({leading_count}편)이다.",
            "2020-2026 구간은 비전 트랜스포머, LLM, RAG, self-supervised learning, diffusion/generative AI, trustworthy AI, AI for science가 서로 연결되는 흐름을 보인다.",
            f"{END_YEAR}년 논문은 아직 인용 누적 시간이 짧으므로 최신성과 영향력은 분리해서 읽어야 한다.",
        ]
        caveat = "이 문서는 PDF 전문 기반 systematic review가 아니라 공개 메타데이터 기반의 citation map이다. 인용수는 영향력을 보여주지만 방법론적 품질, 안전성, 재현성을 직접 보장하지 않는다."
        conclusion = "AI 연구는 대규모 모델과 데이터 중심 방법론을 축으로 확장하면서, 동시에 안전성, 설명가능성, domain translation, 재현성 검증을 더 강하게 요구받고 있다."
    else:
        title = f"AI Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Citation Map"
        abstract = (
            f"This draft review maps AI research from {START_YEAR} through {END_YEAR}, investigating up to {CANDIDATES_PER_YEAR:,} candidate papers per year "
            f"from free public Semantic Scholar metadata and selecting the top {TARGET_TOTAL:,} papers overall by citation count. "
            "The resulting collection is organized by research taxonomy and enriched with deterministic key ideas, strengths, limitations, and AI-specific keyword tags."
        )
        methods = (
            "For each year, broad AI-oriented queries were sent to Semantic Scholar Academic Graph bulk search. Records were retained when title, abstract, or venue metadata matched explicit AI relevance expressions, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most 1,000 candidates per year by citation count. The final 100 papers were selected across the full period by citation count, with influential citation count and metadata importance score retained as tie-breakers and audit signals."
        )
        findings = [
            f"The {len(selected):,} selected papers account for {total_cites:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The largest category is {leading_cat} ({leading_count} papers), reflecting the influence of large-scale model and representation work.",
            "The period connects vision transformers, LLMs, retrieval augmentation, self-supervised learning, diffusion/generative AI, trustworthy AI, and AI-for-science workflows.",
            f"Papers from {END_YEAR} are structurally citation-disadvantaged because the year is partial and citation accumulation is still immature.",
        ]
        caveat = "This is a metadata-driven citation map rather than a full systematic review of every PDF. Citation count is a useful influence signal, but it is not a direct measure of methodological quality, safety, or reproducibility."
        conclusion = "AI research is expanding around large-scale models and data-centric methods while simultaneously increasing pressure for safety, interpretability, domain validation, and reproducible evaluation."
    category_lines = [f"{cat}: {count}" for cat, count in cats.most_common()]
    year_lines = [
        f"{year}: {stats[year]['count']} selected papers, {stats[year]['citations']:,} citations, top selected paper: {stats[year]['top']['title']}"
        for year in YEARS
        if year in stats
    ]
    return {
        "title": title,
        "abstract": abstract,
        "methods": methods,
        "findings": findings,
        "category_lines": category_lines,
        "year_lines": year_lines,
        "top_cited": top_cited,
        "top_scored": top_scored,
        "caveat": caveat,
        "conclusion": conclusion,
    }


def html_ranked_table(rows, metric):
    label = "Citations" if metric == "citations" else "Importance"
    out = [f"<table><thead><tr><th>Year</th><th>Rank</th><th>Paper</th><th>{label}</th><th>Category</th></tr></thead><tbody>"]
    for paper in rows:
        value = paper["citationCount"] if metric == "citations" else paper["importanceScore"]
        link = f'<a href="{html.escape(paper["url"] or paper["semanticScholarUrl"])}">{html.escape(paper["title"])}</a>'
        out.append(f"<tr><td>{paper['year']}</td><td>{paper['rank']}</td><td>{link}</td><td>{value}</td><td>{html.escape(paper['category'])}</td></tr>")
    out.append("</tbody></table>")
    return "\n".join(out)


def write_review_html(selected, korean=False):
    content = review_sections(selected, korean=korean)
    lang = "ko" if korean else "en"
    name = "review_ko.html" if korean else "review_en.html"
    html_doc = f"""<!doctype html>
<html lang="{lang}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(content['title'])}</title>
  <style>
    body {{ font-family: Georgia, 'Noto Serif KR', serif; max-width: 920px; margin: 40px auto; padding: 0 22px; line-height: 1.72; color:#172033; }}
    h1 {{ line-height:1.15; }}
    h2 {{ margin-top:34px; }}
    .abstract {{ background:#f5f7fb; border-left:4px solid #2563eb; padding:14px 18px; }}
    table {{ width:100%; border-collapse:collapse; margin:16px 0; }}
    th, td {{ border-bottom:1px solid #d9dee8; padding:8px; vertical-align:top; text-align:left; }}
    th {{ background:#f4f6fa; }}
  </style>
</head>
<body>
  <h1>{html.escape(content['title'])}</h1>
  <p><strong>Generated:</strong> {date.today().isoformat()} &middot; <strong>Dataset:</strong> {len(selected):,} selected papers</p>
  <h2>Abstract</h2>
  <p class="abstract">{html.escape(content['abstract'])}</p>
  <h2>1. Scope and Methods</h2>
  <p>{html.escape(content['methods'])}</p>
  <h2>2. Key Findings</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['findings'])}</ul>
  <h2>3. Taxonomy</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['category_lines'])}</ul>
  <h2>4. Year-by-Year Coverage</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['year_lines'])}</ul>
  <h2>5. Top Papers by Citation Count</h2>
  {html_ranked_table(content['top_cited'], 'citations')}
  <h2>6. Top Papers by Metadata Importance Score</h2>
  {html_ranked_table(content['top_scored'], 'importance')}
  <h2>7. Limitations</h2>
  <p>{html.escape(content['caveat'])}</p>
  <h2>8. Conclusion</h2>
  <p>{html.escape(content['conclusion'])}</p>
  <h2>Selected References</h2>
  <ol>{''.join(f'<li>{html.escape(reference_line(paper))}</li>' for paper in content['top_cited'])}</ol>
</body>
</html>
"""
    (PAPER_DIR / name).write_text(html_doc, encoding="utf-8")


def write_review_docx(selected):
    content = review_sections(selected, korean=False)
    doc = Document()
    doc.add_heading(content["title"], level=0)
    doc.add_paragraph(f"Generated: {date.today().isoformat()} | Dataset: {len(selected):,} selected papers")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(content["abstract"])
    doc.add_heading("1. Scope and Methods", level=1)
    doc.add_paragraph(content["methods"])
    doc.add_heading("2. Key Findings", level=1)
    for item in content["findings"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("3. Taxonomy", level=1)
    for item in content["category_lines"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("4. Year-by-Year Coverage", level=1)
    for item in content["year_lines"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("5. Top Papers by Citation Count", level=1)
    for paper in content["top_cited"]:
        doc.add_paragraph(f"{paper['year']} #{paper['rank']}: {paper['title']} ({paper['citationCount']:,} citations)", style="List Number")
    doc.add_heading("6. Limitations", level=1)
    doc.add_paragraph(content["caveat"])
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(content["conclusion"])
    doc.save(PAPER_DIR / "review_en.docx")


def write_curation_method_html(markdown_text):
    body = []
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            body.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("- "):
            body.append(f"<li>{html.escape(line[2:])}</li>")
        elif line.strip():
            body.append(f"<p>{html.escape(line)}</p>")
    html_doc = f"""<!doctype html>
<html lang="en">
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>Curation Method</title></head>
<body>{''.join(body)}</body>
</html>
"""
    (PAPER_DIR / "curation_method.html").write_text(html_doc, encoding="utf-8")


def write_project_files(selected, candidates):
    citation = f"""cff-version: 1.2.0
title: "Awesome AI: A Metadata-Driven Citation Map of AI Research, {YEAR_RANGE_TEXT}"
message: "If you use this curated dataset or review draft, please cite this repository."
type: dataset
authors:
  - name: "Honggi"
repository-code: "https://github.com/honggi82/awesome-ai"
date-released: "{date.today().isoformat()}"
license: "CC-BY-4.0"
keywords:
  - "artificial intelligence"
  - "machine learning"
  - "deep learning"
  - "foundation models"
  - "bibliometrics"
"""
    (ROOT / "CITATION.cff").write_text(citation, encoding="utf-8")
    (ROOT / "LICENSE").write_text("CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.\n", encoding="utf-8")
    (ROOT / ".gitignore").write_text("__pycache__/\n*.pyc\n.tools/\ndata/cache/\n.playwright-cli/\noutput/playwright/\n", encoding="utf-8")
    publish = r"""@echo off
setlocal
cd /d "%~dp0"

set "GH_EXE=%~dp0.tools\gh\bin\gh.exe"
if not exist "%GH_EXE%" if exist "%~dp0..\awesome-BCI\.tools\gh\bin\gh.exe" set "GH_EXE=%~dp0..\awesome-BCI\.tools\gh\bin\gh.exe"
if not exist "%GH_EXE%" set "GH_EXE=gh"

"%GH_EXE%" auth status
if errorlevel 1 (
  echo.
  echo GitHub login is required. Run:
  echo   "%GH_EXE%" auth login --hostname github.com --web --scopes repo
  exit /b 1
)

"%GH_EXE%" repo view honggi82/awesome-ai >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" repo create honggi82/awesome-ai --public --description "Awesome AI: metadata-driven AI paper curation, 2020-2026" --source . --remote origin --push
) else (
  git remote set-url origin https://github.com/honggi82/awesome-ai.git
  git push -u origin main
)
if errorlevel 1 exit /b %errorlevel%

"%GH_EXE%" api repos/honggi82/awesome-ai/pages -X POST -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" api repos/honggi82/awesome-ai/pages -X PUT -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
)

echo.
echo Done: https://github.com/honggi82/awesome-ai
echo Pages: https://honggi82.github.io/awesome-ai/
"""
    (ROOT / "publish_to_github.bat").write_text(publish, encoding="utf-8")
    method = f"""# Curation Method

## Scope

- Topic: AI research across machine learning, deep learning, foundation models, NLP, vision, reinforcement learning, generative AI, trustworthy AI, graph learning, multimodal learning, and AI for science.
- Period: {YEAR_RANGE_TEXT}.
- Candidate target: up to {CANDIDATES_PER_YEAR:,} papers per year.
- Final selection: top {TARGET_TOTAL:,} papers overall by citation count from the audited yearly candidate pools.

## Data Source

Metadata comes from the free public Semantic Scholar Academic Graph bulk search endpoint. OpenAlex was used only as a public concept/subfield reference for the AI topic (`Artificial intelligence`, `C154945302`; Artificial Intelligence subfield `1702`).

## Ranking

Records are filtered to the requested publication year, screened for explicit AI relevance in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, and normalized title, then ranked by citation count. Influential citation count and a deterministic metadata importance score are retained as audit fields.

## Enrichment

Taxonomy, key ideas, strengths, limitations, method tags, and keyword convention tags are generated with deterministic rules from public metadata. No paid API, paid LLM, paid translation, or paid compute was used.
"""
    (PAPER_DIR / "curation_method.md").write_text(method, encoding="utf-8")
    write_curation_method_html(method)


def copy_public_assets():
    for filename in (PAPERS_CSV, TAXONOMY_CSV, CANDIDATES_CSV, PERIOD_ANALYSIS_JSON):
        shutil.copyfile(DATA_DIR / filename, DOCS_DIR / "data" / filename)
    shutil.copyfile(PAPER_DIR / "review_en.html", DOCS_DIR / "paper" / "review_en.html")
    shutil.copyfile(PAPER_DIR / "review_ko.html", DOCS_DIR / "paper" / "review_ko.html")
    shutil.copyfile(PAPER_DIR / "curation_method.html", DOCS_DIR / "paper" / "curation_method.html")


def main():
    for path in (DATA_DIR, DOCS_DIR, PAPER_DIR):
        path.mkdir(exist_ok=True)
    if "--reuse-candidates" in sys.argv:
        selected, selected_by_year, candidates_by_year = reuse_existing_candidates()
    else:
        selected, selected_by_year, candidates_by_year = collect_papers()
    candidates = write_data(selected, selected_by_year, candidates_by_year)
    write_taxonomy_dataset(selected)
    write_period_analysis(selected)
    write_readme(selected, candidates)
    write_taxonomy_icons(selected)
    write_charts(selected)
    write_site(selected)
    write_review_html(selected, korean=False)
    write_review_html(selected, korean=True)
    write_review_docx(selected)
    write_project_files(selected, candidates)
    copy_public_assets()
    print(f"[done] generated {len(selected):,} selected papers from {len(candidates):,} candidates", flush=True)


if __name__ == "__main__":
    main()
